"""
routers/teacher.py — 교직원 전용 기능
  POST /conflict/   : 문서 업로드 → 기존 규정과 충돌 분석
  POST /export/pdf  : 선택 답변 → PDF 다운로드
  POST /export/docx : 선택 답변 → DOCX 다운로드
"""
import os, io, re, json
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from pdfminer.high_level import extract_text
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import psycopg2

router = APIRouter(tags=["teacher"])

GEN_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

# ── PDF 폰트: 서버 시작 시 1회 초기화 ──────────────────────────────
_PDF_FONT = "Helvetica"

def _init_pdf_font():
    global _PDF_FONT
    # 이미 한글 폰트 등록됐으면 스킵
    try:
        pdfmetrics.getFont("KR")
        _PDF_FONT = "KR"
        return
    except Exception:
        pass

    # 폰트 후보 (Windows → Linux → Mac 순)
    candidates = [
        "C:/Windows/Fonts/malgun.ttf",      # 맑은 고딕 (Windows 10/11)
        "C:/Windows/Fonts/malgunbd.ttf",     # 맑은 고딕 Bold
        "C:/Windows/Fonts/batang.ttc",       # 바탕체 (ttc 제외)
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
        "/System/Library/Fonts/AppleGothic.ttf",
    ]

    for fp in candidates:
        if not os.path.exists(fp):
            continue
        if fp.endswith(".ttc"):
            continue  # TTC 컬렉션은 TTFont 단독 등록 불가 → 스킵
        try:
            pdfmetrics.registerFont(TTFont("KR", fp))
            _PDF_FONT = "KR"
            return
        except Exception:
            continue

_init_pdf_font()  # 모듈 임포트 시 1회 실행


def _db_url():
    url = os.getenv("DATABASE_URL")
    if not url:
        raise HTTPException(500, "DATABASE_URL이 .env에 설정되지 않았습니다.")
    return url

def set_ai_client(client):
    pass  # 호환성 유지용

def _client():
    from groq import Groq
    key = os.getenv("GROQ_API_KEY")
    if not key:
        raise HTTPException(500, "GROQ_API_KEY가 .env에 없습니다.")
    return Groq(api_key=key)


# ── 텍스트 추출 ────────────────────────────────────────────────────
def _extract(file: UploadFile) -> str:
    file.file.seek(0)
    data  = file.file.read()
    fname = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()

    if fname.endswith(".pdf") or "pdf" in ctype:
        try:
            text = extract_text(io.BytesIO(data))
            if text and text.strip():
                return text.strip()
        except Exception:
            pass
        raise HTTPException(400, "PDF 텍스트 추출 실패. 텍스트 기반 PDF만 지원합니다.")

    if fname.endswith(".docx") or "wordprocessingml" in ctype:
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if fname.endswith(".hwpx"):
        import zipfile, tempfile
        from bs4 import BeautifulSoup as _BS4
        texts = []
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                section_files = sorted([f for f in z.namelist()
                    if re.match(r'Contents/[Ss]ection\d+\.xml', f)])
                if not section_files:
                    section_files = [f for f in z.namelist()
                        if f.endswith('.xml') and 'section' in f.lower()]
                for sf in section_files:
                    try:
                        soup = _BS4(z.read(sf), "xml")
                    except Exception:
                        soup = _BS4(z.read(sf), "html.parser")
                    for tag in soup.find_all(re.compile(r'(?:hp:)?t$')):
                        if tag.get_text().strip():
                            texts.append(tag.get_text())
        except Exception as e:
            raise HTTPException(400, f"HWPX 추출 실패: {e}")
        return "\n".join(texts)

    if fname.endswith(".hwp"):
        import tempfile, subprocess, sys, os as _os
        tmp = _os.path.join(tempfile.gettempdir(), f"hsu_hwp_{_os.getpid()}.hwp")
        try:
            with open(tmp, "wb") as f:
                f.write(data)
            scripts_dir = _os.path.join(_os.path.dirname(sys.executable), "Scripts")
            for cmd in [
                [_os.path.join(scripts_dir, "hwp5txt.exe"), tmp],
                ["hwp5txt", tmp],
                [sys.executable, "-m", "hwp5.hwp5txt", tmp],
            ]:
                try:
                    result = subprocess.run(cmd, capture_output=True, timeout=60,
                                            encoding="utf-8", errors="replace")
                    if result.returncode == 0 and result.stdout.strip():
                        return result.stdout.strip()
                except (FileNotFoundError, OSError):
                    continue
                except Exception:
                    continue
            raise HTTPException(400, "HWP 추출 실패. hwp5txt 명령어를 확인하세요.")
        finally:
            try: _os.unlink(tmp)
            except Exception: pass

    if fname.endswith(".json"):
        try:
            items = json.loads(data.decode("utf-8"))
            if isinstance(items, list):
                return "\n".join(" ".join(str(v) for v in item.values() if v)
                                 for item in items if isinstance(item, dict))
            return json.dumps(items, ensure_ascii=False)
        except Exception as e:
            raise HTTPException(400, f"JSON 파싱 실패: {e}")

    for enc in ["utf-8", "cp949", "euc-kr"]:
        try:
            return data.decode(enc).strip()
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore").strip()


# ── POST /conflict/ ───────────────────────────────────────────────
class ConflictReport(BaseModel):
    summary: str
    conflicts: list[dict]
    recommendations: str
    filename: str


# 한국어 일반 불용어 — 검색 키워드에서 제외
_STOPWORDS = {
    "규정", "조항", "조의", "제정", "개정", "시행", "본문", "내용", "사항",
    "경우", "다음", "다른", "이때", "기타", "관련", "필요", "가능", "있다",
    "한다", "한성대", "한성대학교", "대학교", "대학", "학교", "학생", "직원",
    "교직원", "위원", "위원회", "기관", "부서", "팀", "센터", "본부",
    "이상", "이하", "이내", "미만", "초과", "약간", "각종", "일부", "전부",
    "또는", "그리고", "다만", "단", "이와", "이를", "이러한", "그러한",
    "이라", "이라고", "라고", "에서", "에는", "에게", "으로", "로서",
    "하는", "하여", "하지", "한다고", "되며", "되어", "이라", "이라는",
    "총칙", "부칙", "별표", "별지", "양식",
}


def _is_garbage_pdf_text(text: str) -> bool:
    """PDF 추출이 깨져서 의미없는 텍스트인지 판정"""
    if len(text) < 30:
        return True
    # 한글이 거의 없으면 깨진 거 (영문/숫자/심볼만)
    hangul = len(re.findall(r'[가-힣]', text))
    if hangul / max(len(text), 1) < 0.15:
        return True
    # 같은 짧은 토큰이 5번 이상 반복되면 깨진 거 (nn nn nn... 패턴)
    tokens = text.split()
    if len(tokens) >= 5:
        from collections import Counter
        cnt = Counter(t for t in tokens if 1 <= len(t) <= 3)
        if cnt and cnt.most_common(1)[0][1] >= len(tokens) * 0.3:
            return True
    return False


def _extract_keywords(text: str, limit: int = 15) -> list[str]:
    """업로드 문서에서 의미있는 검색 키워드 추출.
    한글 2-6글자 명사 위주, 빈도순, 불용어 제외."""
    # 한글 토큰만 (영문/숫자 제외)
    tokens = re.findall(r'[가-힣]{2,8}', text)
    from collections import Counter
    cnt = Counter()
    for tok in tokens:
        if tok in _STOPWORDS:
            continue
        if len(tok) < 2:
            continue
        cnt[tok] += 1
    # 빈도 2회 이상만, 빈도순으로
    candidates = [w for w, c in cnt.most_common(50) if c >= 2]
    # 너무 일반적인 한 글자 어휘 빠진 단어도 일부 포함 (제목·머리글에서 한 번만 나오는 핵심어)
    once = [w for w, c in cnt.most_common(80) if c == 1 and len(w) >= 3]
    return (candidates + once)[:limit]


def _find_target_regulation(doc_text: str, keywords: list[str]) -> str | None:
    """업로드 문서가 어느 기존 규정에 해당하는지 찾는다.
    - 문서 첫 600자에서 '○○ 규정' 패턴 추출 후 DB title과 매칭
    - 못 찾으면 키워드 매칭이 가장 많은 규정 선택"""
    head = doc_text[:600]
    # '○○○ 규정' 패턴
    title_m = re.search(r'([가-힣\s·\(\)]{4,30}?(?:규정|규칙|학칙|지침|세칙|기준|운영규정))', head)
    candidate_title = title_m.group(1).strip() if title_m else ""

    try:
        conn = psycopg2.connect(_db_url())
        cur = conn.cursor()

        if candidate_title:
            # 제목 직접 매칭 (공백 제거 후 부분일치)
            stripped = candidate_title.replace(" ", "")
            cur.execute("""
                SELECT rule_title, COUNT(*) AS cnt FROM rule_chunks
                WHERE REPLACE(rule_title, ' ', '') ILIKE %s
                  AND (url IS NULL OR url NOT LIKE 'upload://%%')
                GROUP BY rule_title ORDER BY cnt DESC LIMIT 1
            """, (f"%{stripped}%",))
            row = cur.fetchone()
            if row:
                conn.close()
                return row[0]

        # fallback: 키워드 점수로 가장 많이 매칭되는 규정 선택
        if keywords:
            # 키워드별 1회씩 SELECT 후 title별 카운트
            from collections import Counter
            title_score = Counter()
            for kw in keywords[:10]:
                cur.execute("""
                    SELECT DISTINCT rule_title FROM rule_chunks
                    WHERE content LIKE %s
                      AND (url IS NULL OR url NOT LIKE 'upload://%%')
                    LIMIT 20
                """, (f"%{kw}%",))
                for (t,) in cur.fetchall():
                    title_score[t] += 1
            conn.close()
            if title_score:
                top, score = title_score.most_common(1)[0]
                # 최소 키워드 3개 이상 겹치는 규정만 인정
                if score >= 3:
                    return top
        else:
            conn.close()
    except Exception:
        pass
    return None


@router.post("/conflict/", response_model=ConflictReport)
async def analyze_conflict(file: UploadFile = File(...)):
    doc_text = _extract(file)
    if not doc_text:
        raise HTTPException(400, "텍스트 추출 실패")

    # ① PDF 추출 깨짐 검사
    if _is_garbage_pdf_text(doc_text):
        return ConflictReport(
            summary="문서에서 의미있는 텍스트를 추출하지 못했습니다. "
                    "스캔본 PDF이거나 텍스트가 이미지로 저장되어 있을 수 있습니다. "
                    "HWPX, DOCX, 또는 텍스트 기반 PDF로 다시 업로드해 주세요.",
            conflicts=[],
            recommendations="문서를 텍스트가 추출 가능한 형식(HWPX/DOCX/텍스트 PDF)으로 변환해 다시 시도하세요.",
            filename=file.filename or "문서",
        )

    # ② 핵심 키워드 추출
    keywords = _extract_keywords(doc_text)
    if not keywords:
        return ConflictReport(
            summary="문서에서 분석에 사용할 의미있는 키워드를 추출하지 못했습니다.",
            conflicts=[],
            recommendations="규정 본문이 충분히 포함된 문서인지 확인해 주세요.",
            filename=file.filename or "문서",
        )

    # ③ 가장 관련 깊은 기존 규정 1개 식별
    target_title = _find_target_regulation(doc_text, keywords)

    # ④ 비교용 기존 규정 본문 수집
    target_articles = []  # [(article, content), ...] — 핵심 규정의 전체 조항
    other_chunks   = []   # 그 외 키워드 매칭 보조 청크

    try:
        conn = psycopg2.connect(_db_url())
        cur = conn.cursor()

        if target_title:
            # 핵심 규정의 모든 본문 조항을 가져온다 (개정이력 제외)
            cur.execute("""
                SELECT article, content FROM rule_chunks
                WHERE rule_title = %s AND article != '개정이력'
                  AND (url IS NULL OR url NOT LIKE 'upload://%%')
                ORDER BY article
                LIMIT 40
            """, (target_title,))
            for art, ctn in cur.fetchall():
                target_articles.append((art, ctn))

        # 보조 — 키워드로 다른 규정에서 관련 청크 일부 (다른 규정과의 충돌도 본다)
        seen = set((target_title, a) for a, _ in target_articles)
        for kw in keywords[:8]:
            cur.execute("""
                SELECT rule_title, article, content FROM rule_chunks
                WHERE content LIKE %s AND article != '개정이력'
                  AND (url IS NULL OR url NOT LIKE 'upload://%%')
                LIMIT 4
            """, (f"%{kw}%",))
            for t, a, c in cur.fetchall():
                if (t, a) in seen:
                    continue
                # 핵심 규정과 다른 규정만 보조로 (핵심 규정은 이미 위에서 다 가져옴)
                if target_title and t == target_title:
                    continue
                other_chunks.append({"title": t, "article": a, "content": c})
                seen.add((t, a))
                if len(other_chunks) >= 8:
                    break
            if len(other_chunks) >= 8:
                break
        conn.close()
    except Exception as e:
        raise HTTPException(500, f"DB 오류: {e}")

    if not target_articles and not other_chunks:
        return ConflictReport(
            summary="이 문서와 비교할 만한 기존 규정을 DB에서 찾지 못했습니다. "
                    "완전히 새로운 영역의 규정이라면 충돌 없이 신규 등록이 가능합니다.",
            conflicts=[],
            recommendations="유관 부서와 협의해 신규 규정으로 등록하는 것을 검토해 주세요.",
            filename=file.filename or "문서",
        )

    # ⑤ 비교 컨텍스트 구성
    rules_ctx_parts = []
    if target_articles:
        rules_ctx_parts.append(f"━━ 핵심 비교 대상: 「{target_title}」 ━━")
        for art, ctn in target_articles[:30]:
            # 청크 본문은 1500자까지
            rules_ctx_parts.append(f"[{art}]\n{(ctn or '').strip()[:1500]}")
    if other_chunks:
        rules_ctx_parts.append("━━ 그 외 관련 규정 일부 (상호 충돌 검토용) ━━")
        for c in other_chunks:
            rules_ctx_parts.append(
                f"[{c['title']} {c['article']}]\n{(c['content'] or '').strip()[:600]}"
            )
    rules_ctx = "\n\n".join(rules_ctx_parts)

    # 업로드 문서는 8000자까지 (앞부분이 보통 핵심)
    doc_for_ai = doc_text[:8000]

    prompt = f"""당신은 한성대학교 규정 검토 전문가입니다.
아래 [업로드 문서]가 [기존 규정]과 어떻게 충돌하거나 불일치하는지 **구체적으로** 분석하세요.

[업로드 문서]
{doc_for_ai}

[기존 규정]
{rules_ctx}

━━ 분석 규칙 (반드시 지킬 것) ━━
1. 추상적·뻔한 표현 절대 금지:
   - 금지 예: "기존 규정과의 충돌 가능성", "현행 규정과 실무 간 괴리",
     "일관성이 부족함", "수정이 필요해 보임", "구체적으로 명시되지 않음"
   - 이런 식의 두루뭉술한 지적은 0점이며 빈 배열을 반환하는 것보다 못함.

2. 충돌이라고 보고하려면 반드시:
   - 업로드 문서의 어느 조/항 ("제N조 N항")에 어떤 내용이 있는지,
   - 기존 규정의 어느 조/항에 어떤 내용이 있는지,
   - 두 내용이 어떻게 부딪치는지 (숫자가 다른지, 조건이 다른지, 기간이 다른지, 자격이 다른지 등)
   - 위 세 가지를 모두 명확히 인용·비교하세요.

3. 좋은 충돌 보고서 예시:
   - "업로드 문서 제12조 1항은 '재택근무 신청은 월 4회 이내'로 규정하나,
      기존 「교직원 복무 규정」 제12조는 '재택근무 신청은 월 2회 이내'로 정하고 있어 횟수가 충돌함"
   - "업로드 문서 제18조는 연가일수를 '연 25일'로 정하나,
      기존 규정 제18조는 '연 21일'이며 근속연수별 가산 규정이 누락됨"

4. 단순히 '업로드 문서에 N년이 언급됐다 / 어떤 문구가 있다 / 없다' 같은 표면적 차이는
   충돌이 아닙니다. 그건 보고하지 마세요.

5. 업로드 문서가 기존 규정의 '개정안'이라면, 변경하려는 부분이 명백히 다른 상위 규정이나
   다른 조항과 부딪칠 때만 충돌로 봅니다. 단순 개정 자체는 충돌이 아닙니다.

6. 충돌이 정말 없으면 conflicts는 빈 배열 []로 두고 summary에서 '검토 결과 충돌 없음'이라고
   솔직하게 답하세요. 억지로 만들지 마세요.

7. severity 기준:
   - "높음": 상위 규정 위반, 법령 저촉, 학생/교직원 권리 침해 가능성
   - "보통": 다른 규정과 숫자·조건이 명백히 다름, 절차상 모순
   - "낮음": 용어 불일치, 경미한 표현 차이

반드시 아래 JSON 형식으로만 응답하세요 (다른 설명·머리말 없이):
{{
  "summary": "업로드 문서가 「OOO 규정」과 비교했을 때 어떤 점에서 충돌하는지 2-4문장으로 구체적으로. 충돌이 없으면 '검토 결과 충돌 사항이 발견되지 않았습니다'.",
  "conflicts": [
    {{
      "regulation": "기존 규정명",
      "article": "기존 규정의 조항 (예: 제12조)",
      "issue": "업로드 문서 제N조의 [원문 일부]는 ~~~로 규정하나, 기존 규정 제N조는 [원문 일부]로 정하고 있어 ~~~ 점에서 충돌함",
      "severity": "높음|보통|낮음"
    }}
  ],
  "recommendations": "어느 조항을 어떻게 수정하면 충돌이 해소되는지 구체적으로 한 단락"
}}"""

    try:
        resp = _client().chat.completions.create(
            model=GEN_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2200,
            temperature=0.1,
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r'^```json\s*', '', raw)
        raw = re.sub(r'^```\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        # JSON 블록만 추출 (혹시 앞뒤 설명이 묻어왔으면)
        m = re.search(r'\{[\s\S]*\}', raw)
        if m:
            raw = m.group(0)
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        parsed = {
            "summary": "AI 응답 파싱에 실패했습니다. 다시 시도해 주세요.",
            "conflicts": [],
            "recommendations": "",
        }
    except Exception as e:
        raise HTTPException(500, f"AI 분석 실패: {e}")

    # ⑥ AI 결과 후처리 — 두루뭉술한 issue 필터
    _BAD_PHRASES = [
        "충돌 가능성", "충돌가능성",
        "괴리", "일관성이 부족", "일관성 부족",
        "구체적으로 명시되지", "구체적으로 명시되어 있지",
        "수정이 필요", "검토가 필요",
        "언급되어 있지만", "언급되었지만",
    ]
    cleaned_conflicts = []
    for c in (parsed.get("conflicts") or []):
        if not isinstance(c, dict):
            continue
        issue = (c.get("issue") or "").strip()
        if len(issue) < 30:
            continue  # 너무 짧은 지적은 버림
        # 두루뭉술 키워드만 있고 구체적 비교가 없으면 버림
        is_bad = any(bp in issue for bp in _BAD_PHRASES)
        has_comparison = ("→" in issue or "그러나" in issue or "그러" in issue
                          or "반면" in issue or "기존" in issue or "현행" in issue
                          or "충돌" in issue)
        if is_bad and not has_comparison:
            continue
        sev = c.get("severity", "보통")
        if sev not in ("높음", "보통", "낮음"):
            sev = "보통"
        cleaned_conflicts.append({
            "regulation": (c.get("regulation") or target_title or "").strip()[:80],
            "article":    (c.get("article") or "").strip()[:60],
            "issue":      issue,
            "severity":   sev,
        })

    return ConflictReport(
        summary         = (parsed.get("summary") or "").strip(),
        conflicts       = cleaned_conflicts,
        recommendations = (parsed.get("recommendations") or "").strip(),
        filename        = file.filename or "문서",
    )


# ── POST /export/pdf ──────────────────────────────────────────────
class ExportReq(BaseModel):
    title:    str = "규정 답변 리포트"
    content:  str
    sources:  list[dict] = []
    question: str = ""

@router.post("/export/pdf")
def export_pdf(req: ExportReq):
    buf = io.BytesIO()
    fn  = _PDF_FONT  # "KR" or "Helvetica"

    def ps(name, size, color="#1A2340", sa=6, leading=None):
        return ParagraphStyle(
            name, fontName=fn, fontSize=size, spaceAfter=sa,
            textColor=colors.HexColor(color),
            leading=leading or size * 1.65
        )

    def safe(t: str) -> str:
        return (t or "").replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('\n','<br/>')

    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm)

    story = []

    # 헤더
    story.append(Paragraph("한성대학교 규정 마스터 AI", ps("meta", 9, "#6B7A99", sa=4)))
    story.append(Paragraph(safe(req.title or "규정 답변 리포트"), ps("title", 16, "#003087", sa=10)))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#003087"), spaceAfter=14))

    # 질문
    if req.question:
        story.append(Paragraph("질문", ps("qlabel", 10, "#003087", sa=4)))
        story.append(Paragraph(safe(req.question), ps("qtxt", 11, "#1A2340", sa=12)))

    # 답변
    story.append(Paragraph("답변", ps("alabel", 10, "#003087", sa=6)))
    for line in (req.content or "").split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 4))
            continue
        # 출처 라인은 색 달리
        if stripped.startswith("(출처"):
            story.append(Paragraph(safe(stripped), ps("src", 9, "#0050B3", sa=3)))
        else:
            story.append(Paragraph(safe(stripped), ps("body", 11, "#1A2340", sa=4)))

    # 참조 조항
    if req.sources:
        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#D1DCF0"), spaceAfter=8))
        story.append(Paragraph("참조 조항", ps("srclabel", 10, "#003087", sa=6)))
        for s in req.sources:
            score = int(float(s.get("score", 0)) * 100)
            line  = f"• {s.get('title','')}  {s.get('article','')}  (유사도 {score}%)"
            story.append(Paragraph(safe(line), ps("srcitem", 9, "#6B7A99", sa=3)))

    doc.build(story)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="hansung_report.pdf"'})


# ── POST /export/docx ─────────────────────────────────────────────
def _set_font(run, name="맑은 고딕", size=None):
    """run에 한글 폰트 직접 지정"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)

def _add_para(doc, text, font="맑은 고딕", size=10.5, bold=False, color=None, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    _set_font(run, font, size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

@router.post("/export/docx")
def export_docx(req: ExportReq):
    doc = DocxDocument()

    # 기본 Normal 스타일에 한글 폰트 적용
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    # 제목
    h = doc.add_heading(req.title or "규정 답변 리포트", level=1)
    if h.runs:
        _set_font(h.runs[0], "맑은 고딕", 16)
        h.runs[0].font.color.rgb = RGBColor(0, 48, 135)

    _add_para(doc, "한성대학교 규정 마스터 AI", size=9, color=(107, 122, 153))
    doc.add_paragraph()

    # 질문
    if req.question:
        _add_para(doc, "[ 질문 ]", size=11, bold=True, color=(0, 48, 135))
        _add_para(doc, req.question, size=11)
        doc.add_paragraph()

    # 답변
    _add_para(doc, "[ 답변 ]", size=11, bold=True, color=(0, 48, 135))
    for line in (req.content or "").split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        color = (0, 80, 179) if stripped.startswith("(출처") else None
        size  = 9 if stripped.startswith("(출처") else 11
        _add_para(doc, stripped, size=size, color=color)

    # 참조 조항
    if req.sources:
        doc.add_paragraph()
        _add_para(doc, "[ 참조 조항 ]", size=11, bold=True, color=(0, 48, 135))
        for s in req.sources:
            score = int(float(s.get("score", 0)) * 100)
            text  = f"• {s.get('title','')}  {s.get('article','')}  (유사도 {score}%)"
            _add_para(doc, text, size=9, color=(107, 122, 153))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="hansung_report.docx"'})