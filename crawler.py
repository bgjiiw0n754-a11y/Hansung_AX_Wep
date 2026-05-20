"""
한성대학교 규정관리시스템 — 전체 개정 이력 크롤러
=====================================================
구조:
  SEQ 하나 = 규정 하나 (예: "학생 집단활동 안전관리 규정")
  SEQ_HISTORY 여러 개 = 그 규정의 개정 버전들

출력: hansung_rules_history.json
[
  {
    "seq": 101,
    "title": "학생 집단활동 안전관리 규정",
    "department": "학생복지팀",
    "url_latest": "https://rule.hansung.ac.kr/...",
    "versions": [
      {
        "seq_history": "456",
        "revision_date": "2025-09-23",   ← 이 버전의 개정/제정 날짜
        "revision_type": "개정",          ← 제정 / 개정 / 일부개정 / 전부개정
        "is_latest": true,
        "content": "...",                 ← 해당 버전 본문 전체
        "attachments": [{"filename": "..."}]
      },
      {
        "seq_history": "234",
        "revision_date": "2021-10-21",
        "revision_type": "개정",
        "is_latest": false,
        "content": "..."
      },
      {
        "seq_history": "101",
        "revision_date": "2016-09-30",
        "revision_type": "제정",
        "is_latest": false,
        "content": "..."
      }
    ]
  },
  ...
]

→ 이 JSON이면 "언제 개정됐나요?", "몇 번 바뀌었나요?",
  "어떤 부분이 달라졌나요?" 질문에 모두 답 가능.

[필요 패키지]
pip install requests beautifulsoup4 pdfplumber python-docx pyhwp lxml
"""

import requests
import json
import re
import time
import zipfile
import tempfile
import subprocess
import io
from pathlib import Path
from bs4 import BeautifulSoup
from datetime import datetime

BASE   = "https://rule.hansung.ac.kr"
OUTPUT = "hansung_rules_history.json"

# ══════════════════════════════════════════════════════════════════
# ⚠️ 쿠키 — 크롤링 전 반드시 최신으로 교체할 것!
#   세션이 만료되면 모든 요청이 빈 페이지/에러로 돌아옴.
#   교체 방법:
#     1) 브라우저로 rule.hansung.ac.kr 로그인
#     2) F12 → Network 탭 → 아무 요청 클릭 → Request Headers의 'cookie' 값 복사
#     3) 아래 COOKIE_STRING 에 통째로 붙여넣기
# ══════════════════════════════════════════════════════════════════
COOKIE_STRING = "_ga_NL5KXSTFMR=GS2.1.s1755828457$o1$g1$t1755828471$j46$l0$h0; _ga_H9TXVM0LGL=GS2.1.s1778759247$o7$g1$t1778759502$j60$l0$h0; _ga_K9JR8L935E=GS2.1.s1779082447$o2$g0$t1779082450$j57$l0$h0; _ga=GA1.1.1100685177.1755828368; JSESSIONID=41D751DFCD764834E7F9E3451FBA019E; _ga_2G1Q8E8S12=GS2.1.s1779210254$o97$g1$t1779211595$j60$l0$h0; TS01e66883=014902a0e1f86b75834dd7ecbee9d19d25ed78fe4c21b0bdffe96e823e731ffa0a47f732ad233a7e5e7e743484453e8c9dcf776306"

HEADERS = {
    "User-Agent":      "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/148.0.0.0 Safari/537.36",
    "Accept":          "text/html,application/xhtml+xml,*/*;q=0.9",
    "Accept-Language": "ko-KR,ko;q=0.9",
    "Referer":         "https://rule.hansung.ac.kr/lmxsrv/main/main.do",
    "Cookie":          COOKIE_STRING,
}

session = requests.Session()
session.headers.update(HEADERS)


# ─────────────────────────────────────────────────────────────────
# 날짜 파싱 헬퍼
# ─────────────────────────────────────────────────────────────────

def parse_date(text: str) -> str:
    """
    다양한 날짜 표기 → "YYYY-MM-DD"
    예: "2025.09.23", "2025-09-23", "2025년 9월 23일", "25.9.23" 등
    """
    if not text:
        return ""
    # YYYY.MM.DD or YYYY-MM-DD
    m = re.search(r'(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})', text)
    if m:
        return f"{m.group(1)}-{m.group(2).zfill(2)}-{m.group(3).zfill(2)}"
    # YY.MM.DD (두 자리 연도)
    m = re.search(r'(\d{2})[.\-/](\d{1,2})[.\-/](\d{1,2})', text)
    if m:
        yy = int(m.group(1))
        year = 2000 + yy if yy < 50 else 1900 + yy
        return f"{year}-{m.group(2).zfill(2)}-{m.group(3).zfill(2)}"
    # 한글: YYYY년 M월 D일
    m = re.search(r'(\d{4})년\s*(\d{1,2})월\s*(\d{1,2})일', text)
    if m:
        return f"{m.group(1)}-{m.group(2).zfill(2)}-{m.group(3).zfill(2)}"
    return ""

def parse_revision_type(text: str) -> str:
    """라벨에서 개정 유형 추출"""
    for kw in ["전부개정", "일부개정", "타법개정", "폐지", "개정", "제정"]:
        if kw in text:
            return kw
    return "개정"


# ─────────────────────────────────────────────────────────────────
# 첨부파일 추출
# ─────────────────────────────────────────────────────────────────

def extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(p.extract_text() for p in pdf.pages if p.extract_text())
    except Exception as e:
        print(f"    ⚠️  PDF: {e}"); return ""

def extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if r: parts.append(r)
        return "\n".join(parts)
    except Exception as e:
        print(f"    ⚠️  DOCX: {e}"); return ""

def extract_hwpx(data: bytes) -> str:
    try:
        texts = []
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            secs = sorted([f for f in z.namelist() if re.match(r'Contents/[Ss]ection\d+\.xml', f)])
            if not secs:
                secs = [f for f in z.namelist() if f.endswith('.xml') and 'section' in f.lower()]
            for fn in secs:
                try:    soup = BeautifulSoup(z.read(fn), "xml")
                except: soup = BeautifulSoup(z.read(fn), "html.parser")
                for tag in soup.find_all(re.compile(r'(?:hp:)?t$')):
                    if tag.get_text().strip(): texts.append(tag.get_text())
        return "\n".join(texts)
    except Exception as e:
        print(f"    ⚠️  HWPX: {e}"); return ""

def extract_hwp(data: bytes) -> str:
    try:
        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as f:
            f.write(data); tmp = f.name
        res = subprocess.run(["hwp5txt", tmp], capture_output=True,
                             timeout=30, encoding="utf-8", errors="replace")
        Path(tmp).unlink(missing_ok=True)
        return res.stdout if res.returncode == 0 else ""
    except Exception as e:
        print(f"    ⚠️  HWP: {e}"); return ""

def _decode_filename(raw: str) -> str:
    raw = (raw or "").strip().strip('"\'')
    if not raw:
        return raw
    # 1) %XX URL 인코딩이면 먼저 디코딩 (예: %EA%B5%90 → 교)
    if re.search(r'%[0-9A-Fa-f]{2}', raw):
        from urllib.parse import unquote
        for enc in ("utf-8", "euc-kr", "cp949"):
            try:
                dec = unquote(raw, encoding=enc, errors="strict")
                if dec and "%" not in dec:
                    return dec
            except Exception:
                continue
        try:
            return unquote(raw, encoding="utf-8", errors="replace")
        except Exception:
            pass
    # 2) latin-1 잘못 디코딩된 경우 복원
    for enc in ("utf-8", "euc-kr", "cp949"):
        try:
            return raw.encode("latin-1").decode(enc)
        except Exception:
            continue
    return raw

def get_attachments(view_soup: BeautifulSoup) -> list[dict]:
    result = []
    seen_seq = set()
    for a in view_soup.find_all("a", href=True):
        m = re.match(r"javascript:fileDown\(['\"](\d+)['\"],\s*['\"](\w+)['\"]\)", a["href"])
        if not m: continue
        fseq, ftype = m.group(1), m.group(2)
        if ftype == "oriPdf" and fseq in seen_seq: continue
        try:
            r = session.post(f"{BASE}/lmxsrv/fileDown.do",
                             data={"FILE_SEQ": fseq, "FILE_TYPE": ftype}, timeout=30)
            if r.status_code != 200 or len(r.content) < 100: continue
            ct = r.headers.get("Content-Type", "")
            if "html" in ct.lower(): continue
            cd = r.headers.get("Content-Disposition", "")
            cd_m = re.search(r"filename\*?=(?:UTF-8'')?([^;\n]+)", cd, re.I)
            fname = _decode_filename(cd_m.group(1)) if cd_m else f"file_{fseq}.hwp"
            ext = Path(fname).suffix.lower()
            text = ""
            if ext == ".pdf":   text = extract_pdf(r.content)
            elif ext == ".docx": text = extract_docx(r.content)
            elif ext == ".hwpx": text = extract_hwpx(r.content)
            elif ext == ".hwp":  text = extract_hwp(r.content)
            if text.strip():
                print(f"    📎 {fname} ({len(text)}자)")
                result.append({"filename": fname, "text": text, "file_seq": fseq})
                seen_seq.add(fseq)
        except Exception as e:
            print(f"    ⚠️  첨부 실패 {fseq}: {e}")
    return result


# ─────────────────────────────────────────────────────────────────
# 개정이력 테이블 파싱
# ─────────────────────────────────────────────────────────────────

def parse_revision_history_table(view_soup: BeautifulSoup) -> list[dict]:
    """
    규정 상세 페이지의 개정이력 테이블 파싱
    보통 <table> 안에 날짜 / 개정유형 / 주요내용 컬럼으로 구성
    """
    history = []
    for table in view_soup.find_all("table"):
        headers = [th.get_text(strip=True) for th in table.find_all("th")]
        # 날짜/개정 관련 헤더가 있는 테이블만 처리
        header_text = " ".join(headers)
        if not any(kw in header_text for kw in ["개정", "제정", "시행", "날짜", "일자", "연혁"]):
            continue
        for tr in table.find_all("tr")[1:]:  # 첫 행은 헤더
            cells = [td.get_text(strip=True) for td in tr.find_all("td")]
            if not cells or len(cells) < 2:
                continue
            # 날짜 포함된 셀 찾기
            date_str = ""
            rev_type = ""
            summary  = ""
            for i, cell in enumerate(cells):
                if not date_str and parse_date(cell):
                    date_str = parse_date(cell)
                elif not rev_type and any(kw in cell for kw in ["개정", "제정", "폐지"]):
                    rev_type = cell
                elif cell and len(cell) > 3:
                    summary = cell
            if date_str:
                history.append({
                    "date":    date_str,
                    "type":    rev_type or parse_revision_type(date_str),
                    "summary": summary
                })
    return history


# ─────────────────────────────────────────────────────────────────
# 히스토리 목록 파싱 (SELECT 드롭다운)
# ─────────────────────────────────────────────────────────────────

def get_all_histories(seq: int) -> list[dict]:
    """
    lawDetail.do?SEQ=X → histroySeq 드롭다운 전체 파싱
    반환: [{"hist": "456", "date": "2025-09-23", "type": "개정", "label": "...", "is_latest": True}, ...]
    인덱스 0 = 최신, 마지막 = 최초 제정
    """
    try:
        r = session.get(
            f"{BASE}/lmxsrv/law/lawDetail.do",
            params={"SEQ": seq, "PAGE_MODE": ""},
            allow_redirects=True, timeout=12
        )
        if r.status_code >= 400 or "errorPage" in r.url or len(r.text) < 500:
            return []

        soup = BeautifulSoup(r.text, "html.parser")
        sel_el = soup.find("select", {"id": "histroySeq"})

        if not sel_el:
            # 드롭다운 없음 → 단일 버전
            m = re.search(r'SEQ_HISTORY=(\d+)', r.url + r.text)
            if m:
                # 본문에서 날짜 추출 시도
                date_m = re.search(r'제정\s*(\d{4}[.\-]\d{1,2}[.\-]\d{1,2})', r.text)
                date_str = parse_date(date_m.group(1)) if date_m else ""
                return [{"hist": m.group(1), "date": date_str,
                         "type": "제정", "label": date_str or "제정", "is_latest": True}]
            return []

        options = sel_el.find_all("option")
        histories = []
        for i, opt in enumerate(options):
            val   = opt.get("value", "").strip()
            label = opt.get_text(strip=True)
            if not val: continue
            hist_m = re.search(r'(\d+)', val)
            if not hist_m: continue
            hist     = hist_m.group(1)
            date_str = parse_date(label)
            rev_type = parse_revision_type(label)
            # 제정이 label에 없으면 마지막 버전은 "제정"으로 추정
            if i == len(options) - 1 and not date_str:
                rev_type = "제정"
            histories.append({
                "hist":     hist,
                "date":     date_str,
                "type":     rev_type,
                "label":    label,
                "is_latest": (i == 0)
            })

        return histories

    except Exception as e:
        print(f"  ⚠️  SEQ={seq} 히스토리 목록: {e}")
        return []


# ─────────────────────────────────────────────────────────────────
# 본문 파싱 — lawFullContent.do 페이지를 구조대로 파싱
# ─────────────────────────────────────────────────────────────────

def parse_law_content(html: str) -> str:
    """
    lawFullContent.do 페이지에서 규정 본문만 정확히 추출.
    - 조항: <div class="JO"> 안의 머리글(<td>) + 본문(<div class="none/hang/ho/mok">)
    - 부칙: <div class="addenda"> + 바로 뒤 <div id="...">
    - 'doc_btn'(조항 연혁/인쇄 버튼), <caption>조항</caption> 등 UI 노이즈는 자동 제외
    """
    soup = BeautifulSoup(html, "html.parser")
    root = soup.find("div", id="lawcontent") or soup.find("body") or soup
    lines: list[str] = []

    # 법령명
    lawname = root.find("div", class_="lawname")
    if lawname:
        nm = lawname.get_text(strip=True)
        if nm:
            lines.append(nm)

    def clean(s: str) -> str:
        return re.sub(r'\s+', ' ', (s or "")).strip()

    # 본문 블록을 문서 순서대로 순회
    for el in root.find_all("div", recursive=True):
        cls = el.get("class") or []

        # ── 조 (제 N 조) ──
        if "JO" in cls:
            # 머리글: article > table > 마지막 td (doc_btn 들어있는 td는 건너뜀)
            head = ""
            art = el.find("div", class_="article")
            if art:
                tds = art.find_all("td")
                for td in tds:
                    if td.find("ul", class_="doc_btn"):
                        continue
                    txt = clean(td.get_text())
                    if txt:
                        head = txt
            if head:
                lines.append(head)
            # 본문: JO 바로 아래의 항/호/목 div (article 내부는 제외)
            for sub in el.find_all("div", recursive=False):
                scls = sub.get("class") or []
                if "article" in scls:
                    continue
                txt = clean(sub.get_text(" "))
                if txt:
                    lines.append(txt)

        # ── 부칙 ──
        elif "addenda" in cls:
            title = clean(el.get_text())            # "부   칙" 또는 "부 칙 (2025.9.23.)"
            if title:
                lines.append(title)
            # 부칙 본문 = addenda 바로 다음 형제 div
            sib = el.find_next_sibling("div")
            if sib and not (sib.get("class") and "addenda" in sib.get("class")):
                txt = clean(sib.get_text(" "))
                if txt:
                    lines.append(txt)

    return "\n".join(lines).strip()


# ─────────────────────────────────────────────────────────────────
# 버전 1개 크롤링
# ─────────────────────────────────────────────────────────────────

def crawl_version(seq: int, hist_info: dict, fetch_attachments: bool) -> dict | None:
    hist = hist_info["hist"]
    try:
        r_view = session.get(
            f"{BASE}/lmxsrv/law/lawFullView.do?SEQ={seq}&SEQ_HISTORY={hist}",
            timeout=15, allow_redirects=True
        )
        if r_view.status_code >= 400 or "errorPage" in r_view.url:
            return None

        view_soup = BeautifulSoup(r_view.text, "html.parser")

        # 제목
        title = ""
        stit = view_soup.select_one(".Stit")
        if stit:
            title = re.sub(r'^\d+-\d+-\S+\s*', '', stit.get_text(strip=True)).strip()

        # 담당부서
        dept = ""
        dtit = view_soup.select_one(".Dtit")
        if dtit:
            dept = dtit.get_text(strip=True).replace("담당부서 :", "").strip()

        # 본문 — 구조 기반 파싱 (조/부칙만, UI 노이즈 제외)
        rc = session.get(
            f"{BASE}/lmxsrv/law/lawFullContent.do?SEQ={seq}&SEQ_HISTORY={hist}",
            timeout=15
        )
        if rc.status_code != 200:
            return None
        content = parse_law_content(rc.text)

        # 개정이력 테이블 (최신 버전에서만 파싱)
        rev_table = []
        if hist_info["is_latest"]:
            rev_table = parse_revision_history_table(view_soup)

        # 첨부파일 (최신 버전만) — 본문에 합치지 않고 별도 보관
        attachments = []
        attachment_text = ""
        if fetch_attachments and hist_info["is_latest"]:
            attachments = get_attachments(view_soup)
            if attachments:
                attachment_text = "\n\n".join(
                    f"[첨부: {a['filename']}]\n{a['text']}" for a in attachments
                ).strip()

        if len(content) < 20 or not title or len(title) < 2:
            return None

        # 날짜가 비어있으면 본문에서 추출 시도
        date_str = hist_info.get("date", "")
        if not date_str:
            # "제정 2016.09.30" 또는 "(2016.09.30.)" 패턴
            dm = re.search(r'(?:제정|개정|시행)[^\d]*(\d{4}[.\-]\d{1,2}[.\-]\d{1,2})', content)
            if not dm:
                dm = re.search(r'\((\d{4}[.\-]\d{1,2}[.\-]\d{1,2})\.?\)', content)
            if dm:
                date_str = parse_date(dm.group(1))

        tag = "✅최신" if hist_info["is_latest"] else f"📜{date_str or hist}"
        attach_cnt = len(attachments)
        print(f"  {tag} SEQ={seq} HIST={hist} {title[:35]}"
              + (f" +첨부{attach_cnt}" if attach_cnt else ""))

        ver = {
            "seq_history":    hist,
            "revision_date":  date_str,
            "revision_type":  hist_info.get("type", "개정"),
            "revision_label": hist_info.get("label", ""),
            "is_latest":      hist_info["is_latest"],
            "content":        content,
            "url":            f"{BASE}/lmxsrv/law/lawFullView.do?SEQ={seq}&SEQ_HISTORY={hist}",
            "attachments":    [{"filename": _decode_filename(a["filename"])}
                               for a in attachments],
        }
        if attachment_text:
            ver["attachment_text"] = attachment_text
        if rev_table:
            ver["revision_history_table"] = rev_table

        return {"title": title, "department": dept, "version": ver}

    except Exception as e:
        print(f"  ❌ SEQ={seq} HIST={hist}: {e}")
        return None


# ─────────────────────────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────────────────────────

def main():
    # 기존 결과 로드 (이어서 실행 지원)
    results: list[dict] = []
    done_seq_set: set[int] = set()

    if Path(OUTPUT).exists():
        results = json.loads(Path(OUTPUT).read_text(encoding="utf-8"))
        # 이미 최신버전까지 완료된 SEQ는 건너뜀
        done_seq_set = {r["seq"] for r in results}
        print(f"기존 {len(results)}개 규정 로드, 이어서 진행\n")

    # SEQ 1~1500 순회
    for seq in range(1, 1501):
        if seq in done_seq_set:
            continue

        histories = get_all_histories(seq)
        if not histories:
            time.sleep(0.3)
            continue

        ver_count = len(histories)
        print(f"\n▶ SEQ={seq}  ({ver_count}개 버전)")

        # 규정 객체 초기화
        regulation = {
            "seq":          seq,
            "title":        "",
            "department":   "",
            "url_latest":   "",
            "version_count": ver_count,
            "versions":     [],
            "revision_history_table": []   # 개정이력 테이블 (최신에서 파싱)
        }

        for idx, h in enumerate(histories):
            # 최신 버전만 첨부파일 수집
            fetch_att = h["is_latest"]
            result = crawl_version(seq, h, fetch_att)

            if result is None:
                time.sleep(0.4)
                continue

            # 제목/부서는 최신에서
            if h["is_latest"]:
                regulation["title"]      = result["title"]
                regulation["department"] = result["department"]
                regulation["url_latest"] = result["version"]["url"]
                if "revision_history_table" in result["version"]:
                    regulation["revision_history_table"] = result["version"].pop("revision_history_table")

            regulation["versions"].append(result["version"])
            time.sleep(0.4)

        if not regulation["title"] and regulation["versions"]:
            regulation["title"] = regulation["versions"][0].get("title", "")

        if regulation["versions"]:
            results.append(regulation)
            done_seq_set.add(seq)

        # 30 SEQ마다 중간저장
        if seq % 30 == 0:
            _save(results)
            print(f"\n  💾 중간저장 (규정 {len(results)}개)\n")

    _save(results)

    # 통계
    total_versions = sum(len(r.get("versions", [])) for r in results)
    multi_ver = [r for r in results if len(r.get("versions", [])) > 1]
    print(f"\n🎉 완료!")
    print(f"   규정 수:         {len(results)}개")
    print(f"   전체 버전 수:    {total_versions}개")
    print(f"   개정 이력 있음:  {len(multi_ver)}개 규정")
    print(f"   출력 파일:       {OUTPUT}")


def _save(results):
    Path(OUTPUT).write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()