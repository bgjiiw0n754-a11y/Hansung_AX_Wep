import os
import jwt
import difflib
import shutil
import time
import requests
from datetime import datetime, timedelta
from dotenv import load_dotenv

# .env를 server.py 파일과 같은 디렉토리에서 명시적으로 로드
# (cwd가 달라도 동작하도록)
_ENV_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
load_dotenv(_ENV_PATH)
print(f"[ENV] .env 로드 시도: {_ENV_PATH}  존재: {os.path.exists(_ENV_PATH)}")

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
import psycopg2
from groq import Groq
from anthropic import Anthropic
import voyageai
from routers.teacher import router as teacher_router, set_ai_client
import re as _re
_re2 = _re   # _re2 는 _re의 별칭 (전역에서 사용 가능)
import json as _json
from pathlib import Path as _Path
from collections import defaultdict as _defaultdict
import secrets as _secrets

# ── 환경 변수 ─────────────────────────────────────────────────────
GROQ_KEY      = os.getenv("GROQ_API_KEY", "")            # fallback / Groq 일부 기능
ANTHROPIC_KEY = os.getenv("ANTHROPIC_API_KEY", "")       # 메인 LLM
UPSTAGE_KEY   = os.getenv("UPSTAGE_API_KEY", "")         # 임베딩
VOYAGE_KEY    = os.getenv("VOYAGE_API_KEY", "")          # rerank
DB_URL        = os.getenv("DATABASE_URL")

_raw_key   = os.getenv("SECRET_KEY", "")
SECRET_KEY = _raw_key if len(_raw_key) >= 32 else _secrets.token_hex(32)
if len(_raw_key) < 32:
    print("⚠️  SECRET_KEY가 짧거나 없습니다.")

ADMIN_ID = os.getenv("ADMIN_ID", "admin")
ADMIN_PW = os.getenv("ADMIN_PW", "1234")

if not ANTHROPIC_KEY:
    raise RuntimeError("❌ ANTHROPIC_API_KEY 미설정")
if not UPSTAGE_KEY:
    raise RuntimeError("❌ UPSTAGE_API_KEY 미설정")
if not VOYAGE_KEY:
    raise RuntimeError("❌ VOYAGE_API_KEY 미설정")
if not DB_URL:
    raise RuntimeError("❌ DATABASE_URL 미설정")

# 모델 식별자
UPSTAGE_EMB_PASSAGE_MODEL = "solar-embedding-1-large-passage"   # 적재용 (build_db에서 사용)
UPSTAGE_EMB_QUERY_MODEL   = "solar-embedding-1-large-query"     # 검색용 (질문 임베딩)
UPSTAGE_EMB_URL           = "https://api.upstage.ai/v1/solar/embeddings"
VOYAGE_RERANK_MODEL       = "rerank-2"
CLAUDE_MODEL              = "claude-sonnet-4-5"

# 백업/fallback용 Groq 모델
GROQ_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

# 검색 파라미터
TOP_K_VECTOR  = 30   # 임베딩으로 후보 30개 추리고
TOP_K_RERANK  = 8    # Voyage로 그 중 진짜 관련 8개만 골라 Claude에 넘김

# 클라이언트 초기화
anthropic_client = Anthropic(api_key=ANTHROPIC_KEY)
voyage_client    = voyageai.Client(api_key=VOYAGE_KEY)
groq_client      = Groq(api_key=GROQ_KEY) if GROQ_KEY else None

print(f"✅ Anthropic ({CLAUDE_MODEL})")
print(f"✅ Upstage embeddings ({UPSTAGE_EMB_QUERY_MODEL})")
print(f"✅ Voyage rerank ({VOYAGE_RERANK_MODEL})")
if groq_client:
    print(f"✅ Groq fallback ({GROQ_MODEL})")


# ── Upstage 임베딩 헬퍼 ───────────────────────────────────────────
def upstage_embed_query(text: str, retries: int = 2) -> list[float]:
    """질문을 4096차원 벡터로 변환 (검색용 query 모델)."""
    headers = {"Authorization": f"Bearer {UPSTAGE_KEY}", "Content-Type": "application/json"}
    payload = {"model": UPSTAGE_EMB_QUERY_MODEL, "input": text[:8000]}
    for attempt in range(retries + 1):
        try:
            r = requests.post(UPSTAGE_EMB_URL, headers=headers, json=payload, timeout=30)
            if r.status_code == 429:
                time.sleep((attempt + 1) * 2)
                continue
            r.raise_for_status()
            return r.json()["data"][0]["embedding"]
        except Exception as e:
            if attempt >= retries:
                raise RuntimeError(f"Upstage 임베딩 실패: {e}")
            time.sleep((attempt + 1) * 1)
    return []


def upstage_embed_passage(texts: list[str], retries: int = 2) -> list[list[float]]:
    """본문 여러 개를 한 번에 임베딩 (적재용 passage 모델, server에서도 업로드 규정 적재 시 사용)."""
    headers = {"Authorization": f"Bearer {UPSTAGE_KEY}", "Content-Type": "application/json"}
    payload = {"model": UPSTAGE_EMB_PASSAGE_MODEL, "input": [t[:8000] for t in texts]}
    for attempt in range(retries + 1):
        try:
            r = requests.post(UPSTAGE_EMB_URL, headers=headers, json=payload, timeout=60)
            if r.status_code == 429:
                time.sleep((attempt + 1) * 2)
                continue
            r.raise_for_status()
            return [d["embedding"] for d in r.json()["data"]]
        except Exception as e:
            if attempt >= retries:
                raise RuntimeError(f"Upstage 임베딩 실패: {e}")
            time.sleep((attempt + 1) * 1)
    return []


# ── Voyage rerank 헬퍼 ────────────────────────────────────────────
def voyage_rerank(query: str, documents: list[str], top_k: int = TOP_K_RERANK) -> list[tuple[int, float]]:
    """후보 문서를 질문과의 관련도순으로 재정렬.
    반환: [(원본 인덱스, 관련도 점수)] — top_k개"""
    try:
        result = voyage_client.rerank(
            query=query,
            documents=documents,
            model=VOYAGE_RERANK_MODEL,
            top_k=top_k,
        )
        return [(r.index, r.relevance_score) for r in result.results]
    except Exception as e:
        print(f"⚠️ Voyage rerank 실패, 상위 {top_k}개 그대로 사용: {e}")
        # fallback — 그냥 앞에서부터
        return [(i, 1.0) for i in range(min(top_k, len(documents)))]


# ── Claude 호출 헬퍼 ──────────────────────────────────────────────
def claude_chat(messages: list[dict], system: str = "", max_tokens: int = 2500,
                temperature: float = 0.3, model: str = None) -> str:
    """Claude messages API 호출. messages는 [{role, content}] 형식.
    fallback: Anthropic 실패 시 Groq로 백업."""
    try:
        kwargs = {
            "model": model or CLAUDE_MODEL,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": messages,
        }
        if system:
            kwargs["system"] = system
        resp = anthropic_client.messages.create(**kwargs)
        # content는 [TextBlock] 형태
        return "".join(b.text for b in resp.content if hasattr(b, "text"))
    except Exception as e:
        print(f"⚠️ Claude 실패, Groq fallback: {e}")
        if not groq_client:
            raise
        # Groq 형식으로 변환
        groq_msgs = []
        if system:
            groq_msgs.append({"role": "system", "content": system})
        groq_msgs.extend(messages)
        resp = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=groq_msgs,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        return resp.choices[0].message.content or ""


def claude_stream(messages: list[dict], system: str = "", max_tokens: int = 1500,
                  temperature: float = 0.3, model: str = None):
    """Claude messages API 스트리밍 — yield text chunks. SSE/StreamingResponse용."""
    kwargs = {
        "model": model or CLAUDE_MODEL,
        "max_tokens": max_tokens,
        "temperature": temperature,
        "messages": messages,
    }
    if system:
        kwargs["system"] = system
    with anthropic_client.messages.stream(**kwargs) as stream:
        for text in stream.text_stream:
            yield text


app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"]
)
security = HTTPBearer()

app.include_router(teacher_router)
# teacher 라우터는 Anthropic 클라이언트와 헬퍼를 받음 (충돌분석에서 사용)
set_ai_client({
    "anthropic": anthropic_client,
    "claude_model": CLAUDE_MODEL,
    "claude_chat": claude_chat,
    "groq": groq_client,
})

def _load_dept_phones():
    """dept_phones.json에서 부서-직통번호 매핑을 로드. 없으면 빈 사전."""
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "dept_phones.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = _json.load(f)
        # 주석/메타 키(_로 시작) 제외
        return {k: v for k, v in data.items()
                if not k.startswith("_") and isinstance(v, str) and v.strip()}
    except Exception as e:
        print(f"[dept_phones] 로드 실패: {e} — 빈 사전으로 시작")
        return {}

DEPT_PHONE = _load_dept_phones()
DEFAULT_PHONE = "02-760-4114"   # 대학 대표번호

# 규정 title → 담당부서 매핑 캐시 (DB의 department 컬럼이 비어있을 때 fallback)
_TITLE_TO_DEPT_CACHE = {"data": None}

def _get_title_to_dept_map() -> dict:
    """hansung_rules_history.json의 최상위 department 정보를 title→부서 매핑으로 캐시.
    DB의 department 컬럼은 versions[]를 따라가 비어있는 경우가 많아 fallback 필요."""
    if _TITLE_TO_DEPT_CACHE["data"] is not None:
        return _TITLE_TO_DEPT_CACHE["data"]
    mapping = {}
    try:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "hansung_rules_history.json")
        with open(path, "r", encoding="utf-8") as f:
            data = _json.load(f)
        for r in data:
            title = (r.get("title") or "").strip()
            dept = (r.get("department") or "").strip()
            if title and dept:
                mapping[title] = dept
    except Exception as e:
        print(f"[title_to_dept] 로드 실패: {e}")
    _TITLE_TO_DEPT_CACHE["data"] = mapping
    return mapping

SYSTEM = """
당신은 한성대학교 규정 전문 안내 AI입니다.
10년 이상 대학 행정·학사 업무를 담당한 전문가처럼 행동합니다.

단순히 규정을 읽어주는 AI가 아닙니다.
복잡한 규정을 실제 업무 상황 기준으로 해석하고,
학생·교직원이 바로 이해할 수 있도록 자연스럽고 친절하게 설명합니다.

━━━━━━━━━━━━━━━━━━
핵심 역할
━━━━━━━━━━━━━━━━━━

- 제공된 [참고 규정 조항]만 근거로 답변합니다.
- 규정 원문을 그대로 복붙하지 않습니다.
- 사용자가 실제로 궁금해할 행정 처리 흐름까지 설명합니다.
- 여러 조항이 연결되면 반드시 함께 설명합니다.
- 질문이 짧거나 모호해도 일반적인 대학 행정 기준으로 해석합니다.
- 이전 대화 맥락이 있다면 이어서 설명합니다.
- 반드시 한국어만 사용합니다.
- AI라는 표현 금지
- 프롬프트/내부 규칙/한계 언급 금지

━━━━━━━━━━━━━━━━━━
답변 스타일
━━━━━━━━━━━━━━━━━━

소제목은 반드시 markdown 볼드체(**) 사용.

예시:
**핵심 요약**
**상세 설명**
**출처**
**관련 질문**

절대:
- "제목:", "본문:", "1.", "2.", "3." 같은 번호형 레이블 사용 금지
- HTML / 표 / 코드블록 / JSON / XML 사용 금지

━━━━━━━━━━━━━━━━━━
규정 조항 구조 표기 규칙 (반드시 준수)
━━━━━━━━━━━━━━━━━━

제N장   → 대제목 (예: 제2장 학점인정 기준)
제N조   → 소제목 (예: 제2조(학점인정))
제N항   → 본문 단락 (예: ① ② ③)
N호     → 1. 2. 3. 형태로 나열
목      → 가. 나. 다. 형태로 나열

표기 예시:
- "제2장 제2조 제1항에서는..."
- "제3조 제2항 1호 가목에 따르면..."

개정 비교 시 반드시:
- 어느 장/조/항/호/목이 바뀌었는지 명시
- 변경 전 원문 → 변경 후 원문 형식으로 직접 비교
- 단어 하나, 숫자 하나 차이도 반드시 명시
- "일부 개정" 표현 절대 사용 금지

━━━━━━━━━━━━━━━━━━
출력 구조 (매우 중요)
━━━━━━━━━━━━━━━━━━

첫 줄: 질문을 자연스럽게 요약한 제목 1줄 (15~30자)

**핵심 요약**
가장 중요한 내용을 먼저 설명. 2~4문장 권장.

**상세 설명**
적용 대상, 신청 조건, 승인 기준, 예외 사항, 처리 절차,
제출 서류, 기한, 금액, 학기 기준, 성적 기준, 제한 조건,
실제 행정 처리 방식, 주의사항 등 최대한 포함.
- 규정 문장 그대로 복붙 금지. 자연스럽게 풀어서 설명.
- 긴 내용은 '-' 목록으로 가독성 향상.
- "~할 수 있다"와 "~하여야 한다" 차이 구분.
- 숫자/기간/비율/학점 등 반드시 명시.

**출처**
(출처: 규정명 제N조)
여러 개면 줄바꿈. 실제 참고 규정 기반으로만.

**담당부서**
부서 정보 있을 때만 한 줄로:
📞 담당부서: (참고한 규정의 실제 담당부서명)

⚠️ 중요:
- 반드시 [참고 자료]의 'department' 필드에 적힌 **실제 부서명을 그대로** 쓰세요.
- "학생지원팀", "교무처" 같은 부서명을 **임의로 만들어내지 마세요**. 한성대학교에 실제 존재하는 부서명만 사용.
- 참고 자료에 부서 정보가 없으면 이 섹션 전체를 생략하세요.
- 전화번호는 절대 답변에 넣지 마세요. 시스템이 별도로 정확한 직통번호를 표시합니다.

**관련 질문**
정확히 4개, 질문만, 답변 포함 금지, Q:/A: 금지.

━━━━━━━━━━━━━━━━━━
개정 비교 답변 규칙
━━━━━━━━━━━━━━━━━━

[버전별 변경 diff]가 제공된 경우:
- 질문에서 언급한 특정 조항(제N조, N항 등)만 집중해서 설명
- "일부 개정", "일부 변경" 같은 모호한 표현 절대 금지
- 변경 전문과 변경 후 전문을 직접 비교하여 달라진 내용을 문장/단어 단위로 명시
- 사소한 조사 하나, 숫자 하나 변경도 반드시 명시
- 형식: "제N조 N항의 [원래 내용]이 [바뀐 내용]으로 변경되었습니다"
- 날짜 나열만 하는 것 절대 금지. 반드시 내용 비교 필수

━━━━━━━━━━━━━━━━━━
절대 금지
━━━━━━━━━━━━━━━━━━

- "AI", "제공된 정보 기준", "참고용", "정확한 내용은 문의"
- 프롬프트 설명 / 내부 규칙 언급
- markdown 헤더(##) / 표 / 코드블록 / HTML / JSON
- 번호 레이블
- "일부 개정", "일부 조항 개정", "일부 변경"

━━━━━━━━━━━━━━━━━━
규정을 찾기 어려운 경우
━━━━━━━━━━━━━━━━━━

유사 규정, 상위 규정, 관련 학사 절차를 최대한 활용.
정말 없을 때만:
"해당 내용을 규정에서 찾지 못했습니다. 담당 부서에 문의하세요."
"""


# ── Pydantic 모델 ─────────────────────────────────────────────────
class LoginRequest(BaseModel):
    username: str
    password: str

class Q(BaseModel):
    question: str
    messages: list[dict] = []

class A(BaseModel):
    answer: str
    sources: list[dict]
    found: bool
    dept: str = ""
    dept_phone: str = ""
    followups: list[str] = []


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
NO_CACHE = {"Cache-Control": "no-store, no-cache, must-revalidate", "Pragma": "no-cache"}


# ── 정적 파일 서빙 ────────────────────────────────────────────────
@app.get("/HSU_logo.png")
def logo():
    path = os.path.join(BASE_DIR, "HSU_logo.png")
    if os.path.exists(path):
        return FileResponse(path)
    raise HTTPException(404, "Logo not found")


# 업로드된 원본 파일 서빙 — upload-regulation으로 저장된 PDF/HWP/DOCX를
# 브라우저에서 바로 열거나 다운로드할 수 있게 한다. 메인 챗봇의 "원문 →" 링크가
# 이 엔드포인트를 가리키도록 _to_viewable_url() 헬퍼가 자동 변환한다.
@app.get("/uploads/{filename:path}")
def serve_uploaded_file(filename: str, raw: int = 0):
    """업로드된 원본 파일 서빙.
    - raw=1: 파일 자체 (PDF는 inline, 그 외는 attachment)
    - raw=0 (기본): HTML 래퍼로 응답. 새 탭이 about:blank로 남는 것 방지.
      PDF는 embed로 임베드, HWP/DOCX 등은 자동 다운로드 페이지."""
    safe = os.path.basename(filename)
    path = os.path.join(BASE_DIR, "uploads", safe)
    if not os.path.exists(path):
        raise HTTPException(404, f"업로드된 파일을 찾을 수 없습니다: {safe}")

    import mimetypes
    from urllib.parse import quote
    content_type, _ = mimetypes.guess_type(safe)

    # raw=1 → 실제 파일 반환
    if raw == 1:
        headers = {}
        if content_type == "application/pdf":
            headers["Content-Disposition"] = f'inline; filename*=UTF-8\'\'{quote(safe)}'
        else:
            headers["Content-Disposition"] = f'attachment; filename*=UTF-8\'\'{quote(safe)}'
        return FileResponse(path, media_type=content_type, headers=headers)

    # 기본: HTML 래퍼 (브라우저 새 탭이 빈 채로 남는 것 방지)
    from fastapi.responses import HTMLResponse
    raw_url = f"/uploads/{quote(safe)}?raw=1"
    ext = safe.rsplit(".", 1)[-1].lower() if "." in safe else ""

    if content_type == "application/pdf":
        html = f"""<!DOCTYPE html><html lang="ko"><head><meta charset="utf-8">
<title>{safe}</title>
<style>html,body{{margin:0;padding:0;height:100%;background:#525659;font-family:'Noto Sans KR',sans-serif}}
.bar{{background:#fff;padding:8px 16px;border-bottom:1px solid #ddd;display:flex;align-items:center;gap:12px;font-size:13px}}
.bar a{{color:#1A5FE0;text-decoration:none;font-weight:600}}
.bar a:hover{{text-decoration:underline}}
.bar .name{{color:#333;font-weight:700;margin-right:auto;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
embed,iframe{{width:100vw;height:calc(100vh - 41px);border:none;display:block}}</style>
</head><body>
<div class="bar"><span class="name">📄 {safe}</span>
<a href="{raw_url}" download="{safe}">⬇ 다운로드</a></div>
<embed src="{raw_url}" type="application/pdf"/>
</body></html>"""
        return HTMLResponse(html)

    # PDF가 아니면 즉시 다운로드 트리거하는 페이지
    html = f"""<!DOCTYPE html><html lang="ko"><head><meta charset="utf-8">
<title>{safe}</title>
<style>body{{font-family:'Noto Sans KR',sans-serif;padding:40px;text-align:center;color:#333}}
.box{{max-width:480px;margin:60px auto;padding:32px;border:1px solid #DDE3EE;border-radius:14px;background:#fff;box-shadow:0 4px 16px rgba(0,48,135,0.06)}}
h2{{color:#003087;margin-top:0}} p{{color:#666;line-height:1.7}}
a{{display:inline-block;margin-top:12px;padding:12px 28px;background:#003087;color:#fff;border-radius:10px;text-decoration:none;font-weight:700}}
a:hover{{background:#1A5FE0}}</style>
</head><body><div class="box">
<h2>📎 {safe}</h2>
<p>이 형식({ext})은 브라우저에서 직접 볼 수 없어 다운로드합니다.</p>
<a href="{raw_url}" download="{safe}">⬇ 파일 다운로드</a>
</div>
<script>setTimeout(function(){{location.href="{raw_url}"}},800)</script>
</body></html>"""
    return HTMLResponse(html)


def _to_viewable_url(url, request=None):
    """DB의 url을 브라우저에서 열 수 있는 형식으로 변환.
    - 'upload://파일명' → '/uploads/파일명' (또는 절대 URL — request 있으면)
    - 그 외 정규 URL은 그대로 (rule.hansung.ac.kr 등)"""
    if not url or not isinstance(url, str):
        return url
    if url.startswith("upload://"):
        from urllib.parse import quote
        filename = url.replace("upload://", "", 1)
        path = "/uploads/" + quote(filename)
        # request가 주어지면 절대 URL로 (클라이언트가 base URL을 못 잡을 때 안전)
        if request is not None:
            try:
                base = str(request.base_url).rstrip("/")
                return base + path
            except Exception:
                pass
        return path
    return url


@app.get("/")
def root():
    path = os.path.join(BASE_DIR, "index.html")
    if os.path.exists(path):
        return FileResponse(path, headers=NO_CACHE)
    raise HTTPException(404,
        f"index.html을 찾을 수 없습니다.\n"
        f"경로: {path}\n"
        f"GitHub의 index.html을 이 폴더에 받아 주세요: "
        f"https://github.com/sumiiniee/Hansung_AX/blob/main/index.html")

@app.get("/login-page")
def login_page():
    path = os.path.join(BASE_DIR, "login.html")
    if os.path.exists(path):
        return FileResponse(path, headers=NO_CACHE)
    raise HTTPException(404,
        f"login.html을 찾을 수 없습니다.\n경로: {path}")

@app.get("/upload")
def upload_page():
    path = os.path.join(BASE_DIR, "upload.html")
    if os.path.exists(path):
        return FileResponse(path, headers=NO_CACHE)
    raise HTTPException(404, "upload.html not found")

@app.get("/health")
def health():
    try:
        conn = psycopg2.connect(DB_URL)
        cur  = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM rule_chunks;")
        count = cur.fetchone()[0]
        conn.close()
        return {"ok": True, "chunks": count}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── 로그인 ────────────────────────────────────────────────────────
@app.post("/login")
def login(data: LoginRequest):
    if data.username != ADMIN_ID or data.password != ADMIN_PW:
        raise HTTPException(status_code=401, detail="아이디 또는 비밀번호가 올바르지 않습니다.")
    token = jwt.encode(
        {"sub": data.username, "exp": datetime.utcnow() + timedelta(hours=12)},
        SECRET_KEY, algorithm="HS256"
    )
    return {"success": True, "token": token}


# ── JWT 검증 ──────────────────────────────────────────────────────
def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=["HS256"])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="토큰이 만료되었습니다.")
    except Exception:
        raise HTTPException(status_code=401, detail="인증에 실패했습니다.")


# ── 토큰 유효성 확인 (관리자 페이지 진입 시 호출) ──────────────────
@app.get("/auth/check")
def auth_check(payload: dict = Depends(verify_token)):
    """토큰이 유효한지만 확인. 유효하면 200, 아니면 401."""
    return {"ok": True, "sub": payload.get("sub")}


# ── 텍스트 추출 헬퍼 ──────────────────────────────────────────────
def _extract_hwp_tables_two_column(data: bytes) -> list:
    """HWP를 HTML로 변환한 뒤, "현행규정 | 개정(안)" 양식의 두 컬럼 표를 모두 찾아서
    좌/우 셀 텍스트를 분리해 반환한다. 한 회의록에 여러 안건(여러 표)이 있을 수 있다.
    반환: [{"heading_before": "표 직전 텍스트", "current": "좌측 모든 셀", "new": "우측 모든 셀"}, ...]
    실패 시 빈 리스트."""
    import tempfile, subprocess, sys, os as _os, time as _time
    if not data: return []
    uniq = f"{_os.getpid()}_{int(_time.time()*1000)}"
    tmp_hwp = _os.path.join(tempfile.gettempdir(), f"hsu_mt_{uniq}.hwp")
    tmp_dir = _os.path.join(tempfile.gettempdir(), f"hsu_mt_html_{uniq}")
    try:
        with open(tmp_hwp, "wb") as f:
            f.write(data)
        _os.makedirs(tmp_dir, exist_ok=True)
        scripts_dir = _os.path.join(_os.path.dirname(sys.executable), "Scripts")
        cmds = [
            [_os.path.join(scripts_dir, "hwp5html.exe"), "--output", tmp_dir, tmp_hwp],
            [_os.path.join(scripts_dir, "hwp5html"), "--output", tmp_dir, tmp_hwp],
            ["hwp5html", "--output", tmp_dir, tmp_hwp],
            [sys.executable, "-m", "hwp5.hwp5html", "--output", tmp_dir, tmp_hwp],
            [sys.executable, "-m", "pyhwp.hwp5html", "--output", tmp_dir, tmp_hwp],
        ]
        html_text = ""
        for cmd in cmds:
            try:
                r = subprocess.run(cmd, capture_output=True, timeout=60)
                if r.returncode == 0:
                    # output dir에서 (x)html 파일 찾기
                    for fn in _os.listdir(tmp_dir):
                        if fn.lower().endswith((".xhtml", ".html", ".htm")):
                            try:
                                with open(_os.path.join(tmp_dir, fn), "r", encoding="utf-8") as fh:
                                    html_text = fh.read()
                                break
                            except Exception:
                                continue
                    if html_text:
                        break
            except FileNotFoundError:
                continue
            except Exception:
                continue
        if not html_text:
            return []

        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_text, "html.parser")
        results = []
        for table in soup.find_all("table"):
            rows = table.find_all("tr")
            if len(rows) < 2:
                continue
            # 첫 행이 "현행규정 | 개정(안)" 헤더인지 확인
            head_cells = rows[0].find_all(["td", "th"])
            if len(head_cells) < 2:
                continue
            left_head = head_cells[0].get_text(strip=True)
            right_head = head_cells[1].get_text(strip=True)
            is_diff_table = (
                ("현행" in left_head) and
                (("개정" in right_head) or ("안" in right_head))
            )
            if not is_diff_table:
                continue
            # 좌/우 셀 본문 합치기 (헤더 행 제외)
            cur_parts, new_parts = [], []
            for row in rows[1:]:
                cells = row.find_all(["td", "th"])
                if len(cells) < 2:
                    continue
                cur_parts.append(cells[0].get_text(separator="\n", strip=True))
                new_parts.append(cells[1].get_text(separator="\n", strip=True))
            cur_joined = "\n".join([p for p in cur_parts if p.strip()])
            new_joined = "\n".join([p for p in new_parts if p.strip()])
            if not (cur_joined and new_joined):
                continue
            # 표 직전 텍스트(예: "20. [5-0-2] 인문과학연구원 규정 개정(안)") 추출
            heading_before = ""
            try:
                prev = table.find_previous(string=True)
                # 표 위로 약 600자만 (안건 번호 + 사유)
                chunks = []
                cur_node = table.find_previous()
                while cur_node and len("\n".join(chunks)) < 600:
                    if hasattr(cur_node, "get_text"):
                        t = cur_node.get_text(separator=" ", strip=True)
                        if t:
                            chunks.insert(0, t)
                    cur_node = cur_node.find_previous()
                heading_before = "\n".join(chunks)[-600:]
            except Exception:
                pass
            results.append({
                "heading_before": heading_before,
                "current": cur_joined,
                "new": new_joined,
            })
        return results
    finally:
        # 임시 파일 정리
        try: _os.remove(tmp_hwp)
        except Exception: pass
        try:
            for fn in _os.listdir(tmp_dir):
                try: _os.remove(_os.path.join(tmp_dir, fn))
                except Exception: pass
            _os.rmdir(tmp_dir)
        except Exception: pass


def _extract_meeting_tables_from_text(text: str) -> list:
    """텍스트 기반 회의록 파서 (HWP HTML 파싱이 실패했을 때 폴백).
    핵심 휴리스틱: '제 N 조'가 두 번 등장하면 첫 번째 = 현행, 두 번째 = 개정안.
    안건이 여러 개면 안건마다 따로 분리.
    반환 형식은 _extract_hwp_tables_two_column 와 동일."""
    if not text or len(text) < 50:
        return []
    results = []

    # 안건 시작 패턴: "XX. [코드] ... 개정(안)" 또는 "XX. 규정명 개정(안)"
    agenda_pat = _re.compile(
        r'(?:^|\n)\s*(\d{1,3})\s*\.\s*(?:\[[\d\-]+\]\s*)?[^\n]*?개정\s*\(?\s*안\s*\)?',
        _re.MULTILINE
    )
    matches = list(agenda_pat.finditer(text))
    if not matches:
        sections = [(0, len(text))]
    else:
        sections = []
        for i, m in enumerate(matches):
            sections.append((m.start(),
                             matches[i+1].start() if i+1 < len(matches) else len(text)))

    article_pat = _re.compile(r'제\s*(\d+)\s*조(?:\s*의\s*\d+)?')
    for sec_start, sec_end in sections:
        section = text[sec_start:sec_end]
        articles = list(article_pat.finditer(section))
        if len(articles) < 2:
            continue
        # 첫 등장과 동일한 조 번호의 두 번째 등장 = 개정안 시작
        first = articles[0]
        first_num = first.group(1)
        second = None
        for a in articles[1:]:
            if a.group(1) == first_num:
                second = a
                break
        if not second:
            continue
        current_text = section[first.start():second.start()].rstrip()
        new_text = section[second.start():].rstrip()
        # 너무 짧으면 (잘못 분리) 스킵
        if len(current_text) < 20 or len(new_text) < 20:
            continue
        heading_before = section[:first.start()].strip()[-500:]
        results.append({
            "heading_before": heading_before,
            "current": current_text,
            "new": new_text,
        })
    return results


def _extract_text(file: UploadFile) -> str:
    """업로드 파일에서 텍스트 추출. 지원: PDF, DOCX, TXT, JSON, HWP, HWPX, DOC.
    각 형식별로 여러 폴백 경로를 시도하며, 모두 실패하면 어떤 단계에서 어떻게
    실패했는지 자세한 에러를 반환한다."""
    import io
    file.file.seek(0)
    data  = file.file.read()
    fname = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()

    if not data:
        raise HTTPException(400, "빈 파일입니다.")

    # ── PDF ─────────────────────────────────────────────────────
    if fname.endswith(".pdf") or "pdf" in ctype:
        errors = []
        # 1) pdfminer
        try:
            from pdfminer.high_level import extract_text as _pdfminer_extract
            text = _pdfminer_extract(io.BytesIO(data))
            if text and text.strip():
                return text.strip()
            errors.append("pdfminer: 빈 텍스트")
        except Exception as e:
            errors.append(f"pdfminer: {e}")
        # 2) pdfplumber 폴백
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                parts = []
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
                if parts:
                    return "\n".join(parts).strip()
            errors.append("pdfplumber: 빈 텍스트")
        except Exception as e:
            errors.append(f"pdfplumber: {e}")
        raise HTTPException(400, f"PDF 텍스트 추출 실패. 시도: {' | '.join(errors)}")

    # ── DOCX ────────────────────────────────────────────────────
    if fname.endswith(".docx") or "wordprocessingml" in ctype:
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # 표 안 텍스트도 가져옴
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in parts:
                            parts.append(t)
            if not parts:
                raise HTTPException(400, "DOCX 본문이 비어있습니다.")
            return "\n".join(parts)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, f"DOCX 추출 실패: {e}")

    # ── HWPX (한글 2014+ 표준, ZIP 기반) ─────────────────────────
    if fname.endswith(".hwpx"):
        import zipfile
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                names = z.namelist()
                # section XML 파일들
                section_files = sorted([f for f in names
                    if _re.match(r'Contents/[Ss]ection\d+\.xml', f)])
                if not section_files:
                    section_files = [f for f in names
                        if f.endswith('.xml') and 'section' in f.lower()]
                if not section_files:
                    raise HTTPException(400,
                        f"HWPX 안에 section XML이 없습니다. (포함 파일: {names[:5]}...)")
                from bs4 import BeautifulSoup as _BS
                texts = []
                for sf in section_files:
                    raw = z.read(sf)
                    try:    soup = _BS(raw, "xml")
                    except: soup = _BS(raw, "html.parser")
                    # 한글 hwpx 본문은 보통 hp:t 태그
                    for tag in soup.find_all(_re.compile(r'(?:hp:)?t$')):
                        t = tag.get_text()
                        if t.strip():
                            texts.append(t)
                if not texts:
                    raise HTTPException(400, "HWPX 본문 텍스트가 비어있습니다.")
                return "\n".join(texts)
        except HTTPException:
            raise
        except zipfile.BadZipFile:
            raise HTTPException(400, "HWPX 파일이 손상되었거나 ZIP 형식이 아닙니다.")
        except Exception as e:
            raise HTTPException(400, f"HWPX 추출 실패: {e}")

    # ── HWP (한글 2010 이전, OLE 기반) ──────────────────────────
    if fname.endswith(".hwp"):
        import tempfile, subprocess, sys, os as _os, time as _time
        errors = []

        # 임시 파일 생성 (충돌 방지 — pid + 타임스탬프)
        uniq = f"{_os.getpid()}_{int(_time.time()*1000)}"
        tmp = _os.path.join(tempfile.gettempdir(), f"hsu_hwp_{uniq}.hwp")
        try:
            with open(tmp, "wb") as f:
                f.write(data)

            scripts_dir = _os.path.join(_os.path.dirname(sys.executable), "Scripts")

            # 0) hwp5proc xml — XML로 변환 (표 안 내용까지 모두 포함됨)
            # hwp5txt는 표를 '<표>' 한 단어로만 출력하므로, 표가 있는 회의록에서는
            # 정보가 통째로 사라진다. hwp5proc xml은 표 안의 모든 텍스트를 보존한다.
            xml_cmds = [
                [_os.path.join(scripts_dir, "hwp5proc.exe"), "xml", tmp],
                [_os.path.join(scripts_dir, "hwp5proc"), "xml", tmp],
                ["hwp5proc", "xml", tmp],
                [sys.executable, "-m", "hwp5.hwp5proc", "xml", tmp],
                [sys.executable, "-m", "pyhwp.hwp5proc", "xml", tmp],
            ]
            for cmd in xml_cmds:
                try:
                    result = subprocess.run(cmd, capture_output=True, timeout=60)
                    if result.returncode == 0 and result.stdout:
                        try:
                            from bs4 import BeautifulSoup as _BS
                            # hwp5proc xml의 출력에는 표 셀까지 모든 텍스트가 포함됨
                            try:
                                soup = _BS(result.stdout, "xml")
                            except Exception:
                                soup = _BS(result.stdout, "html.parser")
                            # 한글 문서의 텍스트는 <Text> / <t> 태그 또는 그 자손에
                            # 들어있음. 가장 안전한 방법: 모든 텍스트 노드를 순서대로
                            # 모은 뒤 줄바꿈으로 합친다 (문서 순서 = XML 순서).
                            parts = []
                            for txt_tag in soup.find_all(_re.compile(r'(?:^|:)(Text|text|t|Char)$')):
                                s = txt_tag.get_text()
                                if s and s.strip():
                                    parts.append(s.strip())
                            # 위 셀렉터가 비면 폴백: root.get_text()
                            if not parts:
                                whole = soup.get_text(separator="\n", strip=True) if soup else ""
                                if whole and whole.strip():
                                    return whole.strip()
                            else:
                                merged = "\n".join(parts)
                                if merged.strip():
                                    return merged.strip()
                            errors.append(f"{cmd[0]} xml: 텍스트 노드 없음")
                        except Exception as e:
                            errors.append(f"{cmd[0]} xml-parse: {e}")
                    elif result.returncode != 0:
                        errors.append(f"{cmd[0]} xml: returncode={result.returncode}")
                except FileNotFoundError:
                    continue
                except subprocess.TimeoutExpired:
                    errors.append(f"{cmd[0]} xml: timeout (60s)")
                except Exception as e:
                    errors.append(f"{cmd[0]} xml: {e}")

            # 1) hwp5txt 외부 명령 (폴백 — 표는 '<표>'로 치환되지만 그래도 시도)
            cmds_to_try = [
                [_os.path.join(scripts_dir, "hwp5txt.exe"), tmp],
                [_os.path.join(scripts_dir, "hwp5txt"), tmp],
                ["hwp5txt", tmp],
                [sys.executable, "-m", "hwp5.hwp5txt", tmp],
                [sys.executable, "-m", "pyhwp.hwp5txt", tmp],
            ]
            for cmd in cmds_to_try:
                try:
                    result = subprocess.run(cmd, capture_output=True, timeout=60,
                                            encoding="utf-8", errors="replace")
                    if result.returncode == 0 and result.stdout and result.stdout.strip():
                        return result.stdout.strip()
                    if result.returncode != 0:
                        errors.append(f"{cmd[0]}: returncode={result.returncode} stderr={(result.stderr or '')[:200]}")
                except FileNotFoundError:
                    continue
                except subprocess.TimeoutExpired:
                    errors.append(f"{cmd[0]}: timeout (60s)")
                except Exception as e:
                    errors.append(f"{cmd[0]}: {e}")

            # 2) pyhwp 라이브러리 직접 호출 (외부 명령 다 실패 시)
            try:
                from hwp5.dataio import ParseError
                from hwp5.xmlmodel import Hwp5File
                hwp = Hwp5File(tmp)
                texts = []
                for section in hwp.bodytext.sections:
                    for para in section.paragraphs:
                        t = para.text if hasattr(para, 'text') else str(para)
                        if t and t.strip():
                            texts.append(t.strip())
                if texts:
                    return "\n".join(texts)
                errors.append("pyhwp xmlmodel: 빈 텍스트")
            except Exception as e:
                errors.append(f"pyhwp xmlmodel: {e}")

            # 3) olefile로 PrvText 스트림 직접 추출 (마지막 폴백 — 미리보기 텍스트)
            try:
                import olefile
                ole = olefile.OleFileIO(tmp)
                if ole.exists('PrvText'):
                    raw = ole.openstream('PrvText').read()
                    # PrvText는 UTF-16LE 인코딩
                    text = raw.decode('utf-16-le', errors='ignore').strip()
                    if text:
                        ole.close()
                        # 미리보기는 본문 일부만 — 경고
                        print("⚠️ HWP 본문 추출 실패 → PrvText(미리보기) 사용")
                        return text
                ole.close()
                errors.append("olefile: PrvText 스트림 없음")
            except Exception as e:
                errors.append(f"olefile: {e}")

            # 모두 실패 — 상세 에러
            detail = " | ".join(errors[:5]) if errors else "원인 불명"
            raise HTTPException(400,
                f"HWP 추출 실패. 시도한 방법이 모두 실패했습니다.\n"
                f"세부 원인: {detail}\n"
                f"해결: 한글 프로그램에서 파일을 열어 '다른 이름으로 저장' → HWPX 또는 DOCX로 변환 후 다시 업로드해 주세요.")
        finally:
            try: _os.unlink(tmp)
            except: pass

    # ── DOC (구버전 워드, 한국 학교에 종종 있음) ────────────────
    if fname.endswith(".doc"):
        raise HTTPException(400,
            "구버전 DOC 형식은 지원하지 않습니다. Word에서 '다른 이름으로 저장' → DOCX로 변환 후 업로드해 주세요.")

    # ── JSON ────────────────────────────────────────────────────
    if fname.endswith(".json"):
        try:
            items = _json.loads(data.decode("utf-8"))
            if isinstance(items, list):
                parts = []
                for item in items:
                    if isinstance(item, dict):
                        parts.append(" ".join(str(v) for v in item.values() if v))
                return "\n".join(parts)
            return _json.dumps(items, ensure_ascii=False)
        except Exception as e:
            raise HTTPException(400, f"JSON 파싱 실패: {e}")

    for enc in ["utf-8", "cp949", "euc-kr"]:
        try: return data.decode(enc).strip()
        except: continue
    return data.decode("utf-8", errors="ignore").strip()


def _chunk_text(text: str, max_len: int = 400) -> list[str]:
    paragraphs = [p.strip() for p in _re.split(r'\n{2,}', text) if p.strip()]
    chunks = []
    buf = ""
    for para in paragraphs:
        if len(buf) + len(para) < max_len:
            buf = (buf + "\n" + para).strip() if buf else para
        else:
            if buf: chunks.append(buf)
            if len(para) > max_len:
                sentences = _re.split(r'(?<=[.!?])\s+', para)
                buf = ""
                for s in sentences:
                    if len(buf) + len(s) < max_len:
                        buf = (buf + " " + s).strip() if buf else s
                    else:
                        if buf: chunks.append(buf)
                        buf = s
            else:
                buf = para
    if buf: chunks.append(buf)
    return [c for c in chunks if len(c) >= 20]


def _extract_article_title(chunk: str, idx: int) -> str:
    m = _re.search(r'(제\s*\d+\s*조(?:의\s*\d+)?\s*(?:\([^)]{1,30}\))?)', chunk)
    if m:
        return m.group(1).strip()[:60]
    first = chunk.split('\n')[0].strip()
    if 3 < len(first) <= 50:
        return first
    return f"청크 {idx + 1}"


# ── 규정 파일 업로드 ──────────────────────────────────────────────
@app.post("/upload-regulation")
async def upload_regulation(file: UploadFile = File(...), payload: dict = Depends(verify_token)):
    filename  = file.filename or "unknown"
    upload_tag = f"upload://{filename}"

    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM rule_chunks WHERE url = %s", (upload_tag,))
        if cur.fetchone()[0] > 0:
            conn.close()
            raise HTTPException(409, f'"{filename}"은(는) 이미 등록된 규정입니다.')
        conn.close()
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"DB 확인 실패: {e}")

    try:    text = _extract_text(file)
    except HTTPException: raise
    except Exception as e: raise HTTPException(400, f"파일 읽기 실패: {e}")

    if not text or len(text.strip()) < 10:
        raise HTTPException(400, "텍스트 추출 불가")

    chunks = _chunk_text(text)
    if not chunks: raise HTTPException(400, "청크 분할 실패")

    try:    embeddings = upstage_embed_passage(chunks)
    except Exception as e: raise HTTPException(500, f"임베딩 실패: {e}")

    CHAP_MAP_UP = {"1":"학교법인","2":"학칙","3":"학사행정","4":"부속기관",
                   "5":"부설기관","6":"위원회","7":"산학협력단","8":"학생군사교육단"}
    try:
        raw_clf = claude_chat(
            messages=[{"role":"user","content":f"""한성대학교 규정 체계에서 다음 문서가 속하는 편 번호(1~8)만 답하세요.
1편:학교법인 2편:학칙 3편:학사행정 4편:부속기관
5편:부설기관 6편:위원회 7편:산학협력단 8편:학생군사교육단
규정명: {filename}
내용: {text[:300]}
숫자 하나만 답하세요:"""}],
            max_tokens=4, temperature=0,
        )
        m_clf    = _re.search(r'[1-8]', raw_clf or "")
        chap_num = m_clf.group(0) if m_clf else "3"
    except: chap_num = "3"

    dept_tag   = f"업로드:{chap_num}"
    upload_tag = f"upload://{filename}"
    inserted   = 0
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        import time as _time
        base_id = int(_time.time() * 1000)
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            cur.execute("""
                INSERT INTO rule_chunks (id,rule_title,article,department,url,content,embedding)
                VALUES (%s,%s,%s,%s,%s,%s,%s::vector)
            """, (str(base_id+i), filename, _extract_article_title(chunk,i),
                  dept_tag, upload_tag, chunk, str(emb)))
            inserted += 1
        conn.commit(); conn.close()
    except Exception as e: raise HTTPException(500, f"DB 저장 실패: {e}")

    os.makedirs(os.path.join(BASE_DIR, "uploads"), exist_ok=True)
    try:
        file.file.seek(0)
        with open(os.path.join(BASE_DIR, "uploads", filename), "wb") as f:
            f.write(file.file.read())
    except: pass

    # ── history.json에도 등록 → 개정 탭의 규정 목록에 노출되도록 ──
    # 이미 같은 title이 history에 있으면 추가하지 않음 (중복 방지)
    try:
        history = _load_history_json()
        rule_title_clean = os.path.splitext(filename)[0]  # 확장자 제거
        already = any((r.get("title") or "").strip() == rule_title_clean.strip() for r in history)
        if not already:
            # 새 seq 할당 — 기존 최대 seq + 1 (충돌 방지)
            max_seq = max((int(r.get("seq", 0)) for r in history), default=0)
            new_seq = max_seq + 1
            today = datetime.now().strftime("%Y-%m-%d")
            history.append({
                "seq":        new_seq,
                "title":      rule_title_clean,
                "department": f"업로드 규정 (제{chap_num}편)",
                "chapter":    int(chap_num),
                "category":   f"제{chap_num}편 {CHAP_MAP_UP.get(chap_num, '학사행정')}",
                "url_latest": upload_tag,
                "version_count": 1,
                "versions": [{
                    "seq_history":   new_seq * 10000,   # 임시 hist id
                    "revision_date": today,
                    "revision_type": "신규",
                    "revision_label": "업로드 등록",
                    "is_latest":     True,
                    "content":       text,
                    "url":           upload_tag,
                    "department":    f"업로드 규정 (제{chap_num}편)",
                    "attachments":   [],
                }],
                "revision_history_table": [],
                "_uploaded": True,    # 업로드 출처 표시 (구분용)
            })
            _save_history_json(history)
            print(f"[UPLOAD] history.json에 등록: seq={new_seq}, title={rule_title_clean!r}")

            # DB 청크에도 seq를 채워줌 → 개정 시 _reindex가 찾을 수 있도록
            try:
                conn = psycopg2.connect(DB_URL); cur = conn.cursor()
                cur.execute("UPDATE rule_chunks SET seq=%s WHERE url=%s",
                            (str(new_seq), upload_tag))
                conn.commit(); conn.close()
            except Exception as e:
                print(f"[UPLOAD] DB seq 갱신 실패(무시): {e}")
    except Exception as e:
        # history 등록 실패해도 업로드 자체는 성공으로 처리 (그래야 사용자가 알아챔)
        print(f"[UPLOAD] history.json 등록 실패(무시): {e}")

    return {"success": True, "filename": filename, "chunks": inserted}


# ── 업로드 규정 목록 / 삭제 ───────────────────────────────────────
@app.get("/uploaded-rules")
def list_uploaded_rules(payload: dict = Depends(verify_token)):
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("""
            SELECT url, COUNT(*) as cnt, MIN(id) as first_id
            FROM rule_chunks WHERE url LIKE 'upload://%'
            GROUP BY url ORDER BY first_id DESC
        """)
        rows = cur.fetchall(); conn.close()
    except Exception as e: raise HTTPException(500, f"DB 오류: {e}")
    return {"rules": [{"filename": r[0].replace("upload://",""), "chunks": r[1], "tag": r[0]} for r in rows]}

@app.delete("/uploaded-rules/{filename:path}")
def delete_uploaded_rule(filename: str, payload: dict = Depends(verify_token)):
    tag = f"upload://{filename}"
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("DELETE FROM rule_chunks WHERE url = %s", (tag,))
        deleted = cur.rowcount; conn.commit(); conn.close()
    except Exception as e: raise HTTPException(500, f"DB 삭제 실패: {e}")
    if deleted == 0: raise HTTPException(404, "해당 규정을 찾을 수 없습니다.")
    local = os.path.join(BASE_DIR, "uploads", filename)
    if os.path.exists(local): os.remove(local)
    return {"success": True, "filename": filename, "deleted_chunks": deleted}


# ── 날짜 파싱 헬퍼 ────────────────────────────────────────────────
def _parse_article_date(article: str, content: str = "", url: str = "") -> tuple[str, str]:
    """
    article 필드에서 날짜 추출 → (art_name, date_str)
    우선순위:
      1) article 끝 (YYYY-MM-DD)
      2) article 끝 (YYYY.MM.DD.)
      3) content 안 (2021.02.12.) 형태 — 규정 본문에 박힌 개정일
      4) URL SEQ_HISTORY fallback
    """
    # 1) YYYY-MM-DD
    m = _re.search(r'\((\d{4}-\d{2}-\d{2})\)$', article)
    if m:
        return article[:m.start()].strip(), m.group(1)

    # 2) YYYY.MM.DD.
    m = _re.search(r'\((\d{4}\.\d{2}\.\d{2})\.?\)$', article)
    if m:
        return article[:m.start()].strip(), m.group(1).replace('.', '-')

    art_name = _re.sub(r'\(\s*\)$', '', article).strip()

    # 3) content 안 개정/제정 날짜 — (2021.02.12.) 패턴
    if content:
        # "개정" 또는 "제정" 바로 뒤 날짜 우선
        mc = _re.search(r'(?:개정|제정|시행)[^\d(]*\((\d{4}[.\-]\d{2}[.\-]\d{2})\.?\)', content)
        if not mc:
            # 괄호 안 날짜 아무거나
            mc = _re.search(r'\((\d{4}[.\-]\d{2}[.\-]\d{2})\.?\)', content)
        if mc:
            raw = mc.group(1).replace('.', '-')
            # YYYY-MM-DD 형식 유효성 체크
            if _re.match(r'\d{4}-\d{2}-\d{2}', raw):
                return art_name, raw

    # 4) URL SEQ_HISTORY fallback
    hist_m = _re.search(r'SEQ_HISTORY=(\d+)', url or '')
    date_str = f"v{hist_m.group(1).zfill(6)}" if hist_m else "날짜미상"
    return art_name, date_str


# ── 버전 diff 생성 ────────────────────────────────────────────────
def get_revision_history(rule_title: str) -> str:
    """개정이력 청크 직접 조회 — 날짜 목록 반환"""
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        # 1) 개정이력 전용 청크
        cur.execute("""
            SELECT content FROM rule_chunks
            WHERE rule_title = %s AND article = '개정이력'
            LIMIT 1
        """, (rule_title,))
        row = cur.fetchone()
        if row:
            conn.close()
            return row[0]

        # 2) 없으면 article에서 날짜 수집해서 직접 만들기
        cur.execute("""
            SELECT DISTINCT article, url FROM rule_chunks
            WHERE rule_title = %s AND article != '개정이력'
            ORDER BY article
        """, (rule_title,))
        rows = cur.fetchall(); conn.close()

        dates = set()
        for article, url in rows:
            _, date_str = _parse_article_date(article, "", url)
            if _re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                dates.add(date_str)

        if not dates:
            return ""

        sorted_dates = sorted(dates, reverse=True)
        lines = [f"  - {d}" for d in sorted_dates]
        return f"[{rule_title}] 개정 이력\n총 {len(sorted_dates)}회\n" + "\n".join(lines)

    except Exception:
        return ""

def get_version_diff(rule_title: str, question: str = "") -> str:
    """
    버전 간 글자 단위 diff를 생성.
    1차: hansung_rules_history.json 의 versions 배열에서 직접 비교 (가장 정확).
    2차: 위가 실패하면 DB의 rule_chunks 에서 시도 (백업).
    """
    # ── 1차: JSON 원본에서 비교 ──
    try:
        with open(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "hansung_rules_history.json"), encoding="utf-8") as f:
            history = _json.load(f)
        reg = next((r for r in history if r.get("title", "").strip() == rule_title.strip()), None)
        if reg:
            versions = reg.get("versions", [])
            # 버전 식별자: revision_date + seq_history (날짜만 같으면 hist로 보조 정렬)
            #             content가 의미 있는(>30자) 버전만 사용
            usable = [v for v in versions if len((v.get("content") or "")) > 30]
            if len(usable) >= 2:
                # 최신순 정렬 (revision_date 내림차순, 같으면 seq_history 내림차순)
                usable.sort(key=lambda v: (v.get("revision_date", ""),
                                            v.get("seq_history", "")), reverse=True)
                # 질문에서 조 번호 추출
                art_keywords = _re.findall(r'제?\s*(\d+)\s*(?:장|조|항|절)', question)

                # 인접한 두 버전을 비교
                blocks = []
                for i in range(len(usable) - 1):
                    new_v = usable[i]
                    old_v = usable[i + 1]
                    new_c = (new_v.get("content") or "").strip()
                    old_c = (old_v.get("content") or "").strip()
                    if new_c == old_c:
                        continue

                    # 조 단위로 쪼개 비교 (제 N 조 패턴)
                    def split_articles(text):
                        # '\n제 N 조' 기준 분리
                        parts = _re.split(r'(?=^제\s*\d+\s*조)', text, flags=_re.M)
                        result = {}
                        for p in parts:
                            m = _re.match(r'^제\s*(\d+)\s*조', p)
                            if m: result[m.group(1)] = p.strip()
                        return result

                    new_arts = split_articles(new_c)
                    old_arts = split_articles(old_c)

                    # 비교 대상 조 선택: 질문에 조 번호 있으면 그것만, 없으면 변경된 모든 조
                    if art_keywords:
                        target_keys = [k for k in new_arts if k in art_keywords]
                    else:
                        target_keys = [k for k in new_arts
                                       if new_arts.get(k, "") != old_arts.get(k, "")]
                    target_keys = target_keys[:6]  # 너무 많으면 자름

                    new_date = new_v.get("revision_date") or "최신"
                    old_date = old_v.get("revision_date") or "이전"

                    for k in target_keys:
                        old_text = old_arts.get(k, "")
                        new_text = new_arts.get(k, "")
                        if not new_text:
                            continue
                        if old_text == new_text:
                            continue

                        # difflib 으로 글자 단위 diff
                        sm = difflib.SequenceMatcher(None, old_text, new_text, autojunk=False)
                        changes = []
                        for tag, i1, i2, j1, j2 in sm.get_opcodes():
                            if tag == 'equal':
                                continue
                            o_seg = old_text[i1:i2].strip()
                            n_seg = new_text[j1:j2].strip()
                            if not o_seg and not n_seg:
                                continue
                            if tag == 'replace':
                                changes.append(f"  '{o_seg}' → '{n_seg}'")
                            elif tag == 'delete':
                                changes.append(f"  삭제: '{o_seg}'")
                            elif tag == 'insert':
                                changes.append(f"  추가: '{n_seg}'")
                        if not changes:
                            continue

                        block  = f"[제{k}조] {old_date} → {new_date}\n"
                        block += f"▼ 변경 전:\n{old_text[:600]}\n\n"
                        block += f"▲ 변경 후:\n{new_text[:600]}\n\n"
                        block += "변경 부분:\n" + "\n".join(changes[:20])
                        blocks.append(block)

                    if len(blocks) >= 8:
                        break

                if blocks:
                    return "\n\n".join(blocks[:10])
    except Exception as e:
        print(f"[get_version_diff JSON 비교 실패] {e}")

    # ── 2차: DB 청크 기반 (기존 로직) ──
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("""
            SELECT article, content, url FROM rule_chunks
            WHERE rule_title = %s AND article != '개정이력'
            ORDER BY article
        """, (rule_title,))
        rows = cur.fetchall(); conn.close()
    except Exception:
        return ""

    from collections import defaultdict
    art_versions = defaultdict(list)

    for article, content, url in rows:
        art_name, date_str = _parse_article_date(article, content, url)
        art_versions[art_name].append((date_str, content))

    art_keywords = _re.findall(r'제?\s*(\d+)\s*(?:장|조|항|절)', question)

    diff_results = []
    for art_name, ver_list in art_versions.items():
        if len(ver_list) < 2:
            continue
        if art_keywords:
            if not any(kw in art_name for kw in art_keywords):
                continue

        ver_list.sort(key=lambda x: x[0])

        for i in range(len(ver_list) - 1):
            date_old, content_old = ver_list[i]
            date_new, content_new = ver_list[i + 1]
            if content_old.strip() == content_new.strip():
                continue

            old_lines = content_old.strip().splitlines()
            new_lines = content_new.strip().splitlines()
            deleted = [l for l in old_lines if l.strip() and l.strip() not in {n.strip() for n in new_lines}]
            added   = [l for l in new_lines if l.strip() and l.strip() not in {o.strip() for o in old_lines}]

            block  = f"[{art_name}] {date_old} → {date_new} 변경\n"
            block += f"▼ 변경 전 전문:\n{content_old.strip()}\n\n"
            block += f"▲ 변경 후 전문:\n{content_new.strip()}\n\n"
            if deleted:
                block += "❌ 삭제된 내용:\n" + "\n".join(f"  - {l}" for l in deleted) + "\n"
            if added:
                block += "✅ 추가된 내용:\n" + "\n".join(f"  + {l}" for l in added) + "\n"
            diff_results.append(block)

    return "\n\n".join(diff_results[:10]) if diff_results else ""


# ── /diff 엔드포인트 (프론트 버전 비교 모달용) ────────────────────
@app.post("/diff")
def get_diff(req: Q):
    title = req.question.strip()
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("SELECT DISTINCT rule_title FROM rule_chunks WHERE rule_title ILIKE %s LIMIT 1", (f"%{title}%",))
        row = cur.fetchone()
        if not row: raise HTTPException(404, "규정을 찾을 수 없습니다.")
        matched = row[0]
        cur.execute("""
            SELECT article, content FROM rule_chunks
            WHERE rule_title = %s AND article != '개정이력' ORDER BY article
        """, (matched,))
        rows = cur.fetchall(); conn.close()
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, str(e))

    from collections import defaultdict
    versions = defaultdict(dict)
    for article, content in rows:
        art_name, date_str = _parse_article_date(article, content)
        versions[date_str][art_name] = content

    sorted_dates = sorted(
        versions.keys(),
        key=lambda d: (0, d) if _re.match(r'\d{4}-\d{2}-\d{2}', d) else (1, d)
    )
    all_articles = sorted({a for v in versions.values() for a in v})

    pairs = []
    for i in range(len(sorted_dates) - 1):
        od, nd = sorted_dates[i], sorted_dates[i + 1]
        articles = []
        for art in all_articles:
            oc = versions[od].get(art, "")
            nc = versions[nd].get(art, "")
            if oc != nc:
                articles.append({"name": art, "old": oc, "new": nc})
        if articles:
            pairs.append({"from_date": od, "to_date": nd, "articles": articles})

    return {"title": matched, "pairs": pairs, "dates": sorted_dates}


# ── 규정 질의 ─────────────────────────────────────────────────────
@app.post("/query", response_model=A)
def query(req: Q, request: Request):
    q = req.question.strip()
    if not q:
        raise HTTPException(400, "Empty question")

    # ── 1) 키워드 확장 (Claude, 빠르게) ────────────────────────────
    extra_keywords = []
    try:
        kw_out = claude_chat(
            messages=[{"role": "user", "content":
                f"다음 질문에서 한국 대학 규정 검색에 쓸 핵심 키워드를 추출하세요. "
                f"동의어·약어·관련 법령 용어도 포함. 쉼표로 구분해 단어만 나열 (5~10개).\n\n"
                f"질문: {q}\n키워드:"}],
            max_tokens=120,
            temperature=0,
        )
        extra_keywords = [k.strip() for k in kw_out.split(',') if k.strip()][:10]
    except Exception:
        pass

    search_text = q + ' ' + ' '.join(extra_keywords)

    # ── 2) Upstage로 질문 임베딩 (4096차원) ────────────────────────
    try:
        qemb = upstage_embed_query(search_text)
    except Exception as e:
        raise HTTPException(500, f"Embedding error: {e}")

    # ── 3) DB 벡터 검색: 후보 30개 + 키워드 보강 ──────────────────
    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("""
            SELECT id, rule_title, article, department, url, content,
                   1 - (embedding <=> %s::vector) AS score
            FROM rule_chunks ORDER BY embedding <=> %s::vector LIMIT %s;
        """, (qemb, qemb, TOP_K_VECTOR))
        rows = list(cur.fetchall())

        # 키워드 LIKE 보강 (조항명·본문에 키워드 들어간 청크)
        raw_keywords = [w for w in (q + ' ' + ' '.join(extra_keywords)).replace("?","").split() if len(w) >= 3]
        keywords = list(raw_keywords)
        for w in raw_keywords:
            if len(w) >= 4:
                for i in range(0, len(w)-2, 2):
                    sub = w[i:i+3]
                    if sub not in keywords: keywords.append(sub)

        if keywords:
            existing_ids = {r[0] for r in rows}
            for kw in keywords[:8]:   # 너무 많이 안 보강 (성능)
                cur.execute("SELECT id,rule_title,article,department,url,content,0.6 AS score FROM rule_chunks WHERE article LIKE %s LIMIT 5", (f"%{kw}%",))
                for r in cur.fetchall():
                    if r[0] not in existing_ids: rows.append(r); existing_ids.add(r[0])
                cur.execute("SELECT id,rule_title,article,department,url,content,0.5 AS score FROM rule_chunks WHERE content LIKE %s LIMIT 5", (f"%{kw}%",))
                for r in cur.fetchall():
                    if r[0] not in existing_ids: rows.append(r); existing_ids.add(r[0])

        conn.close()
    except Exception as e:
        raise HTTPException(500, f"DB error: {e}")

    if not rows:
        return A(answer="해당 내용을 규정에서 찾지 못했습니다. 담당 부서에 문의하세요.", sources=[], found=False)

    # ── 4) Voyage rerank: 후보 중 진짜 관련 깊은 8개만 ────────────
    docs_for_rerank = [
        f"[{r[1]} {r[2]}]\n{(r[5] or '')[:1500]}"
        for r in rows
    ]
    try:
        ranked = voyage_rerank(q, docs_for_rerank, top_k=TOP_K_RERANK)
        # ranked: [(원본 인덱스, 관련도)] — 관련도 순 정렬됨
        rerank_filtered = []
        for orig_idx, score in ranked:
            row = list(rows[orig_idx])
            row[6] = float(score)   # 7번째 컬럼(score)을 rerank score로 덮어씀
            rerank_filtered.append(tuple(row))
        rows = rerank_filtered
    except Exception as e:
        print(f"⚠️ rerank 실패, 임베딩 점수만 사용: {e}")
        rows = sorted(rows, key=lambda x: x[6], reverse=True)[:TOP_K_RERANK]

    # rerank score가 너무 낮으면 못 찾은 걸로
    if not rows or rows[0][6] < 0.20:
        return A(answer="해당 내용을 규정에서 찾지 못했습니다. 담당 부서에 문의하세요.", sources=[], found=False)

    # ── 5) 개정 비교 질문 감지 ────────────────────────────────────
    DIFF_KEYWORDS = [
        "뭐가 바뀌", "무엇이 바뀌", "어떻게 바뀌",
        "개정 내용", "개정사항", "개정 정보", "개정 이력", "개정 비교",
        "개정일자", "개정날짜", "개정일", "언제 개정", "몇 번 개정",
        "변경 내용", "변경사항", "달라진", "수정된", "차이",
        "언제 바뀌", "어떤 부분이", "뭐가 달라"
    ]
    is_diff_question = any(kw in q for kw in DIFF_KEYWORDS)

    diff_ctx = ""
    if is_diff_question and rows:
        from collections import Counter

        # 질문에 명시적으로 규정명이 있으면 그걸 우선 사용
        q_no_space = q.replace(" ", "")
        candidate_titles = list(set(r[1] for r in rows))
        best_title = None
        best_match_len = 0
        for t in candidate_titles:
            t_clean = t.replace(" ", "")
            for i in range(len(t_clean) - 4):
                for j in range(len(t_clean), i + 4, -1):
                    chunk = t_clean[i:j]
                    if len(chunk) >= 5 and chunk in q_no_space and len(chunk) > best_match_len:
                        best_title = t
                        best_match_len = len(chunk)
                        break

        if best_title and best_match_len >= 5:
            top_title = best_title
        else:
            top_title = Counter(r[1] for r in rows).most_common(1)[0][0]

        DATE_ONLY_KW = ["개정일자", "개정날짜", "개정일", "언제 개정", "몇 번 개정", "개정 이력", "개정 정보"]
        is_date_only = any(kw in q for kw in DATE_ONLY_KW)

        if is_date_only:
            hist = get_revision_history(top_title)
            if hist:
                diff_ctx = f"\n\n[개정 이력 데이터 — 이 규정({top_title})의 실제 데이터입니다. 반드시 이 데이터만 사용하세요.]\n{hist}"
        else:
            diff = get_version_diff(top_title, question=q)
            if diff:
                diff_ctx = f"\n\n[버전별 변경 diff — 이 규정({top_title})의 데이터입니다.]\n{diff}"

    # ── 6) 부서 매핑 ──────────────────────────────────────────────
    title_to_dept = _get_title_to_dept_map()

    def _resolve_dept(title: str, db_dept: str) -> str:
        if db_dept and db_dept.strip():
            return db_dept.strip()
        return title_to_dept.get((title or "").strip(), "")

    # ── 7) 참고 컨텍스트 ──────────────────────────────────────────
    ctx_lines = []
    for i, r in enumerate(rows):
        title, article, db_dept, content = r[1], r[2], r[3], r[5]
        dept = _resolve_dept(title, db_dept)
        head = f"[조항 {i+1}] {title} {article}"
        if dept:
            head += f"  (담당부서: {dept})"
        ctx_lines.append(f"{head}\n{content}")
    ctx = "\n\n".join(ctx_lines)
    ctx += diff_ctx

    # ── 8) 멀티턴 + Claude 답변 ───────────────────────────────────
    messages = []
    history = [m for m in req.messages if m.get("role") in ("user", "assistant")]
    prev = history[:-1][-6:] if len(history) > 1 else []
    for m in prev:
        messages.append({"role": m["role"], "content": m["content"]})
    messages.append({"role": "user", "content": f"[참고 규정 조항]\n{ctx}\n\n[질문]\n{q}"})

    try:
        answer = claude_chat(messages=messages, system=SYSTEM, max_tokens=1500, temperature=0.3)
    except Exception as e:
        raise HTTPException(500, f"Generation error: {e}")

    sources = [{
        "title": r[1],
        "article": r[2],
        "department": _resolve_dept(r[1], r[3]),
        "url": _to_viewable_url(r[4], request),
        "score": round(r[6], 3),
    } for r in rows]

    import re as _re2
    dept = ""; dept_phone = ""; followups = []

    # 부서 추출 — AI가 어떤 표기로 적든 잡도록 패턴 확장
    dept_patterns = [
        r'📞\s*(?:담당|관련)?\s*부서\s*[:：\-]?\s*([^\n(（]+)',
        r'(?:^|\n)\s*(?:담당|관련)\s*부서\s*[:：\-]\s*([^\n(（]+)',
        r'(?:^|\n)\s*부서\s*[:：]\s*([^\n(（]+)',
    ]
    for pat in dept_patterns:
        m = _re2.search(pat, answer)
        if m:
            cand = m.group(1).strip().rstrip('.,·-')
            cand = cand.split(',')[0].split('·')[0].strip()
            if cand and len(cand) <= 40:
                dept = cand
                break

    # source 후보 부서들 — 점수 높은 순
    source_depts = []
    for s in sources:
        d = (s.get("department") or "").strip()
        if not d:
            continue
        # 콤마/구두점으로 묶인 다부서는 첫 부서만
        d = d.split(',')[0].split('·')[0].split('、')[0].strip()
        if d and d not in source_depts:
            source_depts.append(d)

    # AI 응답에 부서가 없으면 source 첫 부서로
    if not dept and source_depts:
        dept = source_depts[0]

    # ⚠️ AI가 임의 부서명을 만들었는지 확인 — DEPT_PHONE에도, source에도 없는 부서는 임의 부서로 간주
    if dept and dept not in DEPT_PHONE:
        # 부분 매칭으로 사전에 있는지 확인
        in_dict = any(k in dept or dept in k for k in DEPT_PHONE.keys())
        # source에 같은 부서가 있는지 확인
        in_source = any(d == dept or d in dept or dept in d for d in source_depts)
        if not in_dict and not in_source and source_depts:
            # AI가 만들어낸 부서명 → source의 실제 부서명으로 교체
            dept = source_depts[0]

    # 부서 → 전화번호 매핑 (어떤 경우든 dept가 있으면 무조건 번호 매겨줌)
    if dept:
        # 정확 매칭 우선
        if dept in DEPT_PHONE:
            dept_phone = DEPT_PHONE[dept]
        else:
            # 부분 매칭 (긴 키 먼저 시도)
            for k in sorted(DEPT_PHONE.keys(), key=len, reverse=True):
                if k in dept or dept in k:
                    dept_phone = DEPT_PHONE[k]
                    break
        # 그래도 못 찾으면 대표번호
        if not dept_phone:
            dept_phone = DEFAULT_PHONE

    # 연관 질문 추출 — 다양한 형식 모두 매칭
    #   - **관련 질문** / 💡 관련 질문 / 관련 질문: / 연관 질문:
    fq_block = _re2.search(
        r'(?:\*\*\s*관련\s*질문\s*\*\*|💡[^\n]*|관련\s*질문\s*[:：]?|연관\s*질문\s*[:：]?)'
        r'\s*\n([\s\S]+?)(?=\n\n\*\*|\n\n[가-힣]|\Z)',
        answer
    )
    if fq_block:
        followups = []
        for line in fq_block.group(1).splitlines():
            stripped = line.strip()
            if not stripped: continue
            # 볼드(**) 마커, 번호/불릿/Q: 제거
            text = _re2.sub(r'\*\*', '', stripped)
            text = _re2.sub(r'^[-*•\d]+[.)]\s*', '', text).strip()
            text = _re2.sub(r'^Q:\s*', '', text, flags=_re2.IGNORECASE).strip()
            if _re2.match(r'^A:\s*', text, flags=_re2.IGNORECASE): continue
            # 너무 길거나 짧으면 제외 (질문 문장이 아닐 가능성)
            if len(text) < 5 or len(text) > 80: continue
            # 물음표 없으면 추가
            if not text.endswith('?') and not text.endswith('?'):
                text = text + '?'
            followups.append(text)
        followups = followups[:4]

    # AI가 4개 미만으로 줬으면 일반 fallback 으로 채우기
    _fallback_fu = [
        "관련 규정 원문은 어디서 볼 수 있나요?",
        "담당 부서에 직접 문의하려면 어떻게 하나요?",
        "비슷한 사례에는 어떤 게 있나요?",
        "예외 조항은 없나요?",
    ]
    for f in _fallback_fu:
        if len(followups) >= 4: break
        if f not in followups: followups.append(f)
    followups = followups[:4]

    clean_answer = _re2.sub(r'\n*📞[^\n]*?(?:담당|관련)?\s*부서[^\n]*', '', answer)
    # 이모지 없는 형태도 제거 (관련 부서: 교무처 / 담당부서 - 교무처 등)
    clean_answer = _re2.sub(r'(?m)^\s*(?:담당|관련)\s*부서\s*[:：\-][^\n]*', '', clean_answer)
    clean_answer = _re2.sub(r'(?m)^\s*부서\s*[:：][^\n]*', '', clean_answer)
    # AI가 답변에 박은 '(...대표번호 ... 요청)' 같은 안내 문구 제거
    clean_answer = _re2.sub(r'\([^()\n]*대표번호[^()\n]*\)', '', clean_answer)
    clean_answer = _re2.sub(r'\(\s*\d{2,3}-\d{3,4}-\d{4}[^()\n]*\)', '', clean_answer)
    clean_answer = _re2.sub(r'\n*(?:\*\*\s*관련\s*질문\s*\*\*|💡[^\n]*|관련\s*질문\s*[:：]?|연관\s*질문\s*[:：]?)[\s\S]*', '', clean_answer)
    clean_answer = _re2.sub(r'(?m)^\s*\d+\s*[.\)]\s*(제목|본문|출처|담당부서|관련\s*질문)[^\n]*\n?', '', clean_answer)
    # '**출처**' 또는 '출처:' 라벨 행과 그 아래 출처 목록(- 또는 • 시작)을 모두 제거
    #   참조 조항 카드로 따로 표시되므로 본문에서 중복 제거
    clean_answer = _re2.sub(
        r'(?m)^\s*(?:\*\*\s*)?출처\s*(?:\*\*)?\s*[:：]?\s*\n((?:\s*[-•*][^\n]*\n?)+)',
        '', clean_answer)
    clean_answer = _re2.sub(r'(?m)^\s*(?:\*\*\s*)?출처\s*(?:\*\*)?\s*[:：][^\n]*\n?', '', clean_answer)
    clean_answer = _re2.sub(r'(?m)^(제목|본문)\s*[:：]\s*', '', clean_answer)
    clean_answer = _re2.sub(r'(?m)^#{1,6}\s*(.+)$', r'\1', clean_answer)
    clean_answer = _re2.sub(r'\*{1,3}([^*\n]+)\*{1,3}', r'\1', clean_answer)
    clean_answer = _re2.sub(r'_{1,2}([^_\n]+)_{1,2}', r'\1', clean_answer)
    clean_answer = _re2.sub(r'(?m)^[-*_]{3,}\s*$', '', clean_answer)
    clean_answer = _re2.sub(r'(?m)^>\s*', '', clean_answer)
    clean_answer = _re2.sub(r'`+([^`]*)`+', r'\1', clean_answer)
    clean_answer = _re2.sub(r'\n{3,}', '\n\n', clean_answer).strip()

    return A(answer=clean_answer, sources=sources, found=True,
             dept=dept, dept_phone=dept_phone, followups=followups)


# ── 본문 정리 함수 (스트리밍 끝나고 부르는 후처리) ────────────────
def _post_process_answer(answer: str) -> tuple[str, list[str]]:
    """모델 raw 답변 → (clean_answer, followups[])
    /query 안의 정리 로직을 함수로 추출. /query-stream에서 재사용."""
    followups: list[str] = []
    fq_block = _re2.search(
        r'(?:\*\*\s*관련\s*질문\s*\*\*|💡[^\n]*|관련\s*질문\s*[:：]?|연관\s*질문\s*[:：]?)'
        r'\s*\n([\s\S]+?)(?=\n\n\*\*|\n\n[가-힣]|\Z)',
        answer
    )
    if fq_block:
        for line in fq_block.group(1).splitlines():
            stripped = line.strip()
            if not stripped: continue
            text = _re2.sub(r'\*\*', '', stripped)
            text = _re2.sub(r'^[-*•\d]+[.)]\s*', '', text).strip()
            text = _re2.sub(r'^Q:\s*', '', text, flags=_re2.IGNORECASE).strip()
            if _re2.match(r'^A:\s*', text, flags=_re2.IGNORECASE): continue
            if len(text) < 5 or len(text) > 80: continue
            if not text.endswith('?') and not text.endswith('?'):
                text = text + '?'
            followups.append(text)
        followups = followups[:4]

    _fb = [
        "관련 규정 원문은 어디서 볼 수 있나요?",
        "담당 부서에 직접 문의하려면 어떻게 하나요?",
        "비슷한 사례에는 어떤 게 있나요?",
        "예외 조항은 없나요?",
    ]
    for f in _fb:
        if len(followups) >= 4: break
        if f not in followups: followups.append(f)
    followups = followups[:4]

    ca = _re2.sub(r'\n*📞[^\n]*?(?:담당|관련)?\s*부서[^\n]*', '', answer)
    ca = _re2.sub(r'(?m)^\s*(?:담당|관련)\s*부서\s*[:：\-][^\n]*', '', ca)
    ca = _re2.sub(r'(?m)^\s*부서\s*[:：][^\n]*', '', ca)
    ca = _re2.sub(r'\([^()\n]*대표번호[^()\n]*\)', '', ca)
    ca = _re2.sub(r'\(\s*\d{2,3}-\d{3,4}-\d{4}[^()\n]*\)', '', ca)
    ca = _re2.sub(r'\n*(?:\*\*\s*관련\s*질문\s*\*\*|💡[^\n]*|관련\s*질문\s*[:：]?|연관\s*질문\s*[:：]?)[\s\S]*', '', ca)
    ca = _re2.sub(r'(?m)^\s*\d+\s*[.\)]\s*(제목|본문|출처|담당부서|관련\s*질문)[^\n]*\n?', '', ca)
    ca = _re2.sub(
        r'(?m)^\s*(?:\*\*\s*)?출처\s*(?:\*\*)?\s*[:：]?\s*\n((?:\s*[-•*][^\n]*\n?)+)',
        '', ca)
    ca = _re2.sub(r'(?m)^\s*(?:\*\*\s*)?출처\s*(?:\*\*)?\s*[:：][^\n]*\n?', '', ca)
    ca = _re2.sub(r'(?m)^(제목|본문)\s*[:：]\s*', '', ca)
    ca = _re2.sub(r'(?m)^#{1,6}\s*(.+)$', r'\1', ca)
    ca = _re2.sub(r'\*{1,3}([^*\n]+)\*{1,3}', r'\1', ca)
    ca = _re2.sub(r'_{1,2}([^_\n]+)_{1,2}', r'\1', ca)
    ca = _re2.sub(r'(?m)^[-*_]{3,}\s*$', '', ca)
    ca = _re2.sub(r'(?m)^>\s*', '', ca)
    ca = _re2.sub(r'`+([^`]*)`+', r'\1', ca)
    ca = _re2.sub(r'\n{3,}', '\n\n', ca).strip()

    return ca, followups


# ── POST /query-stream — SSE 스트리밍 답변 ────────────────────────
@app.post("/query-stream")
def query_stream(req: Q, request: Request):
    """/query와 동일한 검색·rerank를 거치되, Claude 답변을 SSE로 토큰 스트리밍.
    이벤트 타입:
      meta  : 검색 결과(sources/dept/dept_phone) 즉시 전송
      token : Claude가 내보내는 텍스트 청크
      done  : 최종 정리된 본문 + followups
      error : 오류 메시지
    """
    q = req.question.strip()
    if not q:
        raise HTTPException(400, "Empty question")

    def event_gen():
        import re as _re2_local

        try:
            # 1) Upstage 임베딩
            qemb = upstage_embed_query(q)

            # 2) DB 후보 + 키워드 LIKE 보강
            conn = psycopg2.connect(DB_URL); cur = conn.cursor()
            cur.execute("""
                SELECT id, rule_title, article, department, url, content,
                       1 - (embedding <=> %s::vector) AS score
                FROM rule_chunks ORDER BY embedding <=> %s::vector LIMIT %s;
            """, (qemb, qemb, TOP_K_VECTOR))
            rows = list(cur.fetchall())

            keywords = [w for w in q.replace("?","").split() if len(w) >= 3][:6]
            if keywords:
                existing = {r[0] for r in rows}
                for kw in keywords:
                    cur.execute("SELECT id,rule_title,article,department,url,content,0.5 AS score FROM rule_chunks WHERE content LIKE %s LIMIT 4", (f"%{kw}%",))
                    for r in cur.fetchall():
                        if r[0] not in existing: rows.append(r); existing.add(r[0])
            conn.close()

            if not rows:
                yield f'event: done\ndata: {_json.dumps({"answer": "해당 내용을 규정에서 찾지 못했습니다. 담당 부서에 문의하세요.", "sources": [], "found": False, "followups": []}, ensure_ascii=False)}\n\n'
                return

            # 3) Voyage rerank
            docs = [f"[{r[1]} {r[2]}]\n{(r[5] or '')[:1500]}" for r in rows]
            try:
                ranked = voyage_rerank(q, docs, top_k=TOP_K_RERANK)
                rerank_filtered = []
                for orig_idx, score in ranked:
                    row = list(rows[orig_idx]); row[6] = float(score)
                    rerank_filtered.append(tuple(row))
                rows = rerank_filtered
            except Exception:
                rows = sorted(rows, key=lambda x: x[6], reverse=True)[:TOP_K_RERANK]

            if not rows or rows[0][6] < 0.20:
                yield f'event: done\ndata: {_json.dumps({"answer": "해당 내용을 규정에서 찾지 못했습니다. 담당 부서에 문의하세요.", "sources": [], "found": False, "followups": []}, ensure_ascii=False)}\n\n'
                return

            # 4) 부서 매핑 + sources 메타
            title_to_dept = _get_title_to_dept_map()
            def _resolve(title, db_dept):
                if db_dept and db_dept.strip(): return db_dept.strip()
                return title_to_dept.get((title or "").strip(), "")

            sources = [{
                "title": r[1], "article": r[2],
                "department": _resolve(r[1], r[3]),
                "url": _to_viewable_url(r[4], request), "score": round(r[6], 3),
            } for r in rows]

            # 부서 추정 — 다부서 분리 + 최빈값
            def _split_depts(raw):
                """'A, B · C' → ['A', 'B', 'C']"""
                if not raw: return []
                out = []
                for piece in _re2_local.split(r'[,·、/]+', raw):
                    p = piece.strip().rstrip('.,·-')
                    if p and len(p) <= 40 and p not in out:
                        out.append(p)
                return out

            from collections import Counter
            all_dept_candidates = []
            for s in sources:
                for d in _split_depts(s["department"]):
                    all_dept_candidates.append(d)

            # 상위 3개 source(가장 관련도 높음)의 부서에 가중치 2배
            weighted = []
            for i, s in enumerate(sources[:3]):
                for d in _split_depts(s["department"]):
                    weighted.extend([d, d])   # 가중치 2배
            weighted.extend(all_dept_candidates)

            dept = ""; dept_phone = ""
            if weighted:
                # 최빈값 (동률이면 위 source의 부서 우선)
                most_common = Counter(weighted).most_common()
                # DEPT_PHONE에 존재하는 부서 우선
                for cand, _ in most_common:
                    if cand in DEPT_PHONE:
                        dept = cand
                        dept_phone = DEPT_PHONE[cand]
                        break
                # 못 찾으면 부분 매칭으로
                if not dept:
                    for cand, _ in most_common:
                        for k in sorted(DEPT_PHONE.keys(), key=len, reverse=True):
                            if k in cand or cand in k:
                                dept = cand
                                dept_phone = DEPT_PHONE[k]
                                break
                        if dept: break
                # 그래도 못 찾으면 1순위 candidate + 대표번호
                if not dept:
                    dept = most_common[0][0]
                    dept_phone = DEFAULT_PHONE

            # 메타 즉시 전송 (참조 카드/부서 카드를 답변 시작 전부터 노출 가능)
            meta = {"sources": sources, "dept": dept, "dept_phone": dept_phone, "found": True}
            yield f'event: meta\ndata: {_json.dumps(meta, ensure_ascii=False)}\n\n'

            # 5) 컨텍스트 구성
            ctx_lines = []
            for i, r in enumerate(rows):
                head = f"[조항 {i+1}] {r[1]} {r[2]}"
                d = _resolve(r[1], r[3])
                if d: head += f"  (담당부서: {d})"
                ctx_lines.append(f"{head}\n{r[5]}")
            ctx = "\n\n".join(ctx_lines)

            # 6) 메시지 구성
            messages = []
            history = [m for m in req.messages if m.get("role") in ("user", "assistant")]
            prev = history[:-1][-6:] if len(history) > 1 else []
            for m in prev:
                messages.append({"role": m["role"], "content": m["content"]})
            messages.append({"role": "user", "content": f"[참고 규정 조항]\n{ctx}\n\n[질문]\n{q}"})

            # 7) Claude 스트리밍 — 토큰 받자마자 즉시 전송
            buffer = []
            try:
                for chunk in claude_stream(messages=messages, system=SYSTEM,
                                           max_tokens=1500, temperature=0.3):
                    buffer.append(chunk)
                    yield f'event: token\ndata: {_json.dumps({"t": chunk}, ensure_ascii=False)}\n\n'
            except Exception as e:
                yield f'event: error\ndata: {_json.dumps({"error": f"Claude 스트리밍 실패: {e}"}, ensure_ascii=False)}\n\n'
                return

            # 8) 정리 + 최종 전송
            raw_answer = "".join(buffer)
            clean_answer, followups = _post_process_answer(raw_answer)

            # 부서 보정 — AI 답변 raw 텍스트에 명시된 부서명 우선
            #   AI가 "📞 담당부서: 글로컬상생홍보팀" 식으로 답하면 그걸 신뢰
            ai_dept = ""
            for pat in [
                r'📞\s*(?:담당|관련)?\s*부서\s*[:：\-]?\s*([^\n(（]+)',
                r'(?:^|\n)\s*(?:담당|관련)\s*부서\s*[:：\-]\s*([^\n(（]+)',
            ]:
                mm = _re2_local.search(pat, raw_answer)
                if mm:
                    cand = mm.group(1).strip().rstrip('.,·-')
                    cand = cand.split(',')[0].split('·')[0].strip()
                    if cand and len(cand) <= 40:
                        ai_dept = cand
                        break

            # AI가 적은 부서가 사전에 있고 다르면 그걸 사용 (단, source에 있는 부서여야 안전)
            if ai_dept and ai_dept != dept:
                # source의 모든 부서 후보들 (다부서 분리 포함)
                all_src_depts = set()
                for s in sources:
                    for d in _split_depts(s["department"]):
                        all_src_depts.add(d)
                # ai_dept가 source에도 있으면 신뢰
                if ai_dept in all_src_depts or any(ai_dept in s or s in ai_dept for s in all_src_depts):
                    dept = ai_dept
                    if dept in DEPT_PHONE:
                        dept_phone = DEPT_PHONE[dept]
                    else:
                        dept_phone = DEFAULT_PHONE
                        for k in sorted(DEPT_PHONE.keys(), key=len, reverse=True):
                            if k in dept or dept in k:
                                dept_phone = DEPT_PHONE[k]; break

            done = {"answer": clean_answer, "followups": followups,
                    "sources": sources, "dept": dept, "dept_phone": dept_phone,
                    "found": True}
            yield f'event: done\ndata: {_json.dumps(done, ensure_ascii=False)}\n\n'
        except Exception as e:
            yield f'event: error\ndata: {_json.dumps({"error": str(e)}, ensure_ascii=False)}\n\n'

    return StreamingResponse(event_gen(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── 규정 목록 ─────────────────────────────────────────────────────
@app.get("/rules")
def get_rules():
    CHAP_MAP = {"1":"학교법인","2":"학칙","3":"학사행정","4":"부속기관",
                "5":"부설기관","6":"위원회","7":"산학협력단","8":"학생군사교육단"}
    chapters = _defaultdict(list)

    try:
        json_path = _Path(BASE_DIR) / "hansung_rules.json"
        if json_path.exists():
            rules = _json.loads(json_path.read_text(encoding="utf-8"))
            for r in rules:
                try:
                    title    = r.get("title", "제목 없음")
                    raw_code = r.get("rule_code", "")
                    if not raw_code:
                        code_m  = _re.search(r'(\d+)-(\d+)-(\d+)', r.get("content", ""))
                        raw_code = f"{code_m.group(1)}-{code_m.group(2)}-{code_m.group(3)}" if code_m else ""

                    # 편 번호 결정 우선순위:
                    # ① chapter 필드 (crawler가 트리에서 가져온 정확한 값)
                    # ② raw_code의 첫 숫자
                    chapter_val = r.get("chapter", 0)
                    if chapter_val and 1 <= int(chapter_val) <= 8:
                        chap_num = str(int(chapter_val))
                    else:
                        chap_num = raw_code.split("-")[0] if raw_code else "0"

                    chap_name = CHAP_MAP.get(chap_num, "기타")
                    chap_key  = f"제{chap_num}편 {chap_name}" if chap_num != "0" else "기타"
                    chapters[chap_key].append({
                        "seq": r.get("seq", 0), "code": raw_code,
                        "name": title, "dept": r.get("department", ""),
                        "url": _to_viewable_url(r.get("url", "")), "uploaded": False
                    })
                except: continue
    except: pass

    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("SELECT DISTINCT rule_title,url,department FROM rule_chunks WHERE url LIKE 'upload://%' ORDER BY rule_title")
        for row in cur.fetchall():
            rule_title, url, department = row
            m_chap = _re.search(r'업로드:([1-8])', department or '')
            chap_num  = m_chap.group(1) if m_chap else "3"
            chap_name = CHAP_MAP.get(chap_num, "학사행정")
            chap_key  = f"제{chap_num}편 {chap_name}"
            chapters[chap_key].append({
                "seq": 9999, "code": f"{chap_num}-upload",
                "name": f"📎 {rule_title}", "dept": "업로드 규정",
                "url": _to_viewable_url(url), "uploaded": True  # ← 변환된 /uploads/... 경로
            })
        conn.close()
    except: pass

    result = []
    for key in sorted(chapters.keys(), key=lambda x: (
        int(_re.search(r'제(\d+)편', x).group(1)) if _re.search(r'제(\d+)편', x) else
        (98 if x == "기타" else 99)
    )):
        result.append({"chapter": key, "rules": sorted(chapters[key], key=lambda x: x["code"])})
    return {"chapters": result, "total": sum(len(c["rules"]) for c in result)}


# ── 규정 키워드 검색 ──────────────────────────────────────────────
@app.post("/search-rules")
def search_rules(req: Q):
    q = req.question.strip()
    if not q: raise HTTPException(400, "Empty query")

    keywords = [w for w in q.split() if len(w) >= 2] or [q]

    try:
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        rows = []; seen = set()
        for kw in keywords[:5]:
            cur.execute("SELECT rule_title,article,department,url,content FROM rule_chunks WHERE rule_title LIKE %s LIMIT 20", (f"%{kw}%",))
            for r in cur.fetchall():
                if (r[0],r[1]) not in seen: rows.append(r); seen.add((r[0],r[1]))
            cur.execute("SELECT rule_title,article,department,url,content FROM rule_chunks WHERE content LIKE %s OR article LIKE %s LIMIT 30", (f"%{kw}%",f"%{kw}%"))
            for r in cur.fetchall():
                if (r[0],r[1]) not in seen: rows.append(r); seen.add((r[0],r[1]))
        conn.close()
    except Exception as e: raise HTTPException(500, f"DB error: {e}")

    if not rows: return {"results": [], "query": q}

    from collections import defaultdict as _dd
    grouped = _dd(list)
    for r in rows:
        grouped[r[0]].append({"article":r[1],"department":r[2],"url":_to_viewable_url(r[3] or ""),"snippet":r[4][:200].strip()})

    results = [{"title":t,"department":c[0]["department"],"url":c[0]["url"],"chunks":c[:3]}
               for t,c in grouped.items()]
    return {"results": results[:20], "query": q}




# ══════════════════════════════════════════════════════════════════
# 규정 개정 (Revision) — 기존 규정 내용 변경 + 되돌리기
# ══════════════════════════════════════════════════════════════════
RULES_JSON_PATH = os.path.join(BASE_DIR, "hansung_rules.json")
HISTORY_JSON_PATH = os.path.join(BASE_DIR, "hansung_rules_history.json")
REVISION_BACKUP_DIR = os.path.join(BASE_DIR, ".revision_backup")


def _load_rules_json() -> list:
    try:
        with open(RULES_JSON_PATH, "r", encoding="utf-8") as f:
            return _json.load(f)
    except Exception as e:
        raise HTTPException(500, f"hansung_rules.json 로드 실패: {e}")


def _save_rules_json(data: list):
    try:
        with open(RULES_JSON_PATH, "w", encoding="utf-8") as f:
            _json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        raise HTTPException(500, f"hansung_rules.json 저장 실패: {e}")


def _load_history_json() -> list:
    """hansung_rules_history.json 로드 (개정 기능 전용 — 버전 정보 포함)."""
    try:
        with open(HISTORY_JSON_PATH, "r", encoding="utf-8") as f:
            return _json.load(f)
    except Exception as e:
        raise HTTPException(500, f"hansung_rules_history.json 로드 실패: {e}")


def _save_history_json(data: list):
    try:
        with open(HISTORY_JSON_PATH, "w", encoding="utf-8") as f:
            _json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        raise HTTPException(500, f"hansung_rules_history.json 저장 실패: {e}")


def _sync_uploaded_rules_to_history() -> int:
    """DB의 'upload://...' 규정 중 history.json에 누락된 것을 자동 등록한다.
    /upload 가 history 등록 코드를 갖기 전에 업로드된 규정은 매칭 후보에서 빠지므로,
    이 함수가 회의록 분석 직전에 한 번씩 호출되어 백필(backfill) 한다.
    반환값: 새로 등록한 규정 개수."""
    try:
        history = _load_history_json()
    except Exception:
        return 0
    existing_urls = set()
    existing_titles = set()
    for r in history:
        u = (r.get("url_latest") or "")
        if u.startswith("upload://"):
            existing_urls.add(u)
        t = (r.get("title") or "").strip()
        if t:
            existing_titles.add(t)

    try:
        conn = psycopg2.connect(DB_URL)
        cur = conn.cursor()
        # 청크들을 id 순서로 모아 본문을 복원 (조 단위 chunk가 id 오름차순)
        cur.execute("""
            SELECT url, MIN(id) AS first_id, COUNT(*) AS cnt,
                   STRING_AGG(content, E'\n\n' ORDER BY id) AS full_content
            FROM rule_chunks
            WHERE url LIKE 'upload://%'
            GROUP BY url
        """)
        rows = cur.fetchall()
        conn.close()
    except Exception as e:
        print(f"[SYNC] DB 조회 실패(무시): {e}")
        return 0

    max_seq = max((int(r.get("seq", 0)) for r in history), default=0)
    added = 0
    from datetime import datetime as _dt
    for url, first_id, cnt, full_content in rows:
        if url in existing_urls:
            continue
        filename = url.replace("upload://", "")
        rule_title_clean = os.path.splitext(filename)[0]
        if rule_title_clean.strip() in existing_titles:
            continue   # 이름 충돌 회피
        max_seq += 1
        try:
            ts = _dt.fromtimestamp(int(first_id) / 1000.0).strftime("%Y-%m-%d")
        except Exception:
            ts = _dt.now().strftime("%Y-%m-%d")
        history.append({
            "seq":        max_seq,
            "title":      rule_title_clean,
            "department": "업로드 규정",
            "chapter":    3,                # 기본값 — 사용자 수정 가능
            "category":   "제3편 학사행정",
            "url_latest": url,
            "version_count": 1,
            "versions": [{
                "seq_history":    max_seq * 10000,
                "revision_date":  ts,
                "revision_type":  "신규",
                "revision_label": "업로드 등록 (자동 백필)",
                "is_latest":      True,
                "content":        full_content or "",
                "url":            url,
                "department":     "업로드 규정",
                "attachments":    [],
            }],
            "revision_history_table": [],
            "_uploaded":         True,
            "_backfilled":       True,
        })
        # DB 청크에 seq도 채워주면 개정 후 _reindex가 잘 찾음
        try:
            conn = psycopg2.connect(DB_URL); cur = conn.cursor()
            cur.execute("UPDATE rule_chunks SET seq=%s WHERE url=%s", (str(max_seq), url))
            conn.commit(); conn.close()
        except Exception:
            pass
        added += 1

    if added > 0:
        try:
            _save_history_json(history)
            print(f"[SYNC] 업로드 규정 {added}건을 history.json에 자동 등록(백필)")
        except Exception as e:
            print(f"[SYNC] 저장 실패: {e}")
            return 0
    return added


def _get_latest_version(reg: dict) -> dict | None:
    """history JSON의 한 규정에서 최신 버전(is_latest=True)을 반환. 없으면 첫 버전."""
    versions = reg.get("versions", []) or []
    if not versions:
        return None
    for v in versions:
        if v.get("is_latest"):
            return v
    return versions[0]


def _get_latest_version_index(reg: dict) -> int:
    """최신 버전의 versions 배열 내 인덱스. 못 찾으면 0."""
    versions = reg.get("versions", []) or []
    for i, v in enumerate(versions):
        if v.get("is_latest"):
            return i
    return 0


# -- 개정 가능한 규정 목록 (검색용) --------------------------------
# ── 어절 단위 + 글자 단위 하이브리드 diff ───────────────────────────
# 글자 단위 diff만 쓰면 "(2026.4.10.)" vs "(2026.0.00.)" 같은 텍스트가 글자 하나씩
# 매칭돼서 "(2026.40.100.)" 식의 혼란스러운 표시가 된다. 그래서:
#   1단계: 공백 기준 어절(token) 단위로 diff
#   2단계: replace된 어절끼리만 다시 글자 단위 diff
# → 큰 변경은 단어 통째로 빨강/파랑, 미세한 변경은 글자별로 보임.
def _word_then_char_diff(cur_text: str, new_text: str) -> list:
    import difflib
    # 공백 묶음 vs 비공백 묶음으로 분리 (공백 보존)
    cur_tokens = _re.findall(r'\s+|\S+', cur_text)
    new_tokens = _re.findall(r'\s+|\S+', new_text)
    sm_w = difflib.SequenceMatcher(None, cur_tokens, new_tokens, autojunk=False)
    segments = []
    for tag, i1, i2, j1, j2 in sm_w.get_opcodes():
        if tag == "equal":
            chunk = "".join(cur_tokens[i1:i2])
            if chunk:
                segments.append({"text": chunk, "kind": "equal"})
        elif tag == "delete":
            chunk = "".join(cur_tokens[i1:i2])
            if chunk:
                segments.append({"text": chunk, "kind": "delete"})
        elif tag == "insert":
            chunk = "".join(new_tokens[j1:j2])
            if chunk:
                segments.append({"text": chunk, "kind": "insert"})
        elif tag == "replace":
            cur_chunk = "".join(cur_tokens[i1:i2])
            new_chunk = "".join(new_tokens[j1:j2])
            # 숫자·구두점이 섞인 토큰(날짜, 일자, 코드 등)은 통째로 처리해야 직관적.
            # "(2026.4.10.)" ↔ "(2026.0.00.)" 같은 변경은 글자 단위로 잘리면 혼란스러움.
            has_nonword = bool(_re.search(r'[\d\W]', cur_chunk + new_chunk))
            # 한글만으로 된 짧은 변경(조사 변화 등)은 어절 안에서 글자 단위로 봐야
            # "촉진시키는→촉진하는"에서 "시키"↔"하"가 보임.
            if (not has_nonword) and len(cur_chunk) <= 24 and len(new_chunk) <= 24:
                sm_c = difflib.SequenceMatcher(None, cur_chunk, new_chunk, autojunk=False)
                for ct, ci1, ci2, cj1, cj2 in sm_c.get_opcodes():
                    if ct == "equal":
                        s = cur_chunk[ci1:ci2]
                        if s: segments.append({"text": s, "kind": "equal"})
                    elif ct == "delete":
                        s = cur_chunk[ci1:ci2]
                        if s: segments.append({"text": s, "kind": "delete"})
                    elif ct == "insert":
                        s = new_chunk[cj1:cj2]
                        if s: segments.append({"text": s, "kind": "insert"})
                    elif ct == "replace":
                        sd = cur_chunk[ci1:ci2]
                        si = new_chunk[cj1:cj2]
                        if sd: segments.append({"text": sd, "kind": "delete"})
                        if si: segments.append({"text": si, "kind": "insert"})
            else:
                # 통째로 빨강/파랑
                if cur_chunk:
                    segments.append({"text": cur_chunk, "kind": "delete"})
                if new_chunk:
                    segments.append({"text": new_chunk, "kind": "insert"})
    # 인접 same-kind 세그먼트 병합 (렌더링 시 span 수 줄임)
    merged = []
    for seg in segments:
        if merged and merged[-1]["kind"] == seg["kind"]:
            merged[-1]["text"] += seg["text"]
        else:
            merged.append(dict(seg))
    return merged


@app.get("/revisable-rules")
def revisable_rules(payload: dict = Depends(verify_token)):
    """hansung_rules_history.json 안의 규정 목록 — 개정 대상 선택용 (최신 버전 기준)"""
    history = _load_history_json()
    out = []
    for r in history:
        try:
            latest = _get_latest_version(r)
            if not latest:
                continue
            out.append({
                "seq":         r.get("seq", 0),
                "title":       r.get("title", "제목 없음"),
                "department":  r.get("department", ""),
                "chapter":     "",
                "code":        "",
                "revised_at":  latest.get("revision_date", ""),
                "content_len": len(latest.get("content", "") or ""),
            })
        except Exception:
            continue
    out.sort(key=lambda x: x["seq"])
    return {"rules": out, "total": len(out)}


# -- 특정 규정 현재 내용 조회 (최신 버전) --------------------------
@app.get("/revisable-rules/{seq}")
def revisable_rule_detail(seq: int, payload: dict = Depends(verify_token)):
    history = _load_history_json()
    for r in history:
        if int(r.get("seq", -1)) == seq:
            latest = _get_latest_version(r)
            if not latest:
                raise HTTPException(404, "최신 버전을 찾을 수 없습니다.")
            return {
                "seq":          r.get("seq"),
                "title":        r.get("title", ""),
                "department":   r.get("department", ""),
                "content":      latest.get("content", ""),
                "revised_at":   latest.get("revision_date", ""),
                "version_count": len(r.get("versions", [])),
            }
    raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")


def _reindex_rule_chunks(seq, title: str, dept: str, url: str, content: str) -> int:
    """DB에서 해당 seq의 본문 청크(개정이력/업로드 제외)를 새 content로 재생성"""
    chunks = _chunk_text(content)
    if not chunks:
        raise HTTPException(400, "개정 내용에서 청크를 만들 수 없습니다.")
    try:
        embeddings = upstage_embed_passage(chunks)
    except Exception as e:
        raise HTTPException(500, f"임베딩 실패: {e}")

    conn = psycopg2.connect(DB_URL)
    cur = conn.cursor()
    try:
        cur.execute("""
            DELETE FROM rule_chunks
            WHERE seq = %s AND article != '개정이력'
              AND (url IS NULL OR url NOT LIKE 'upload://%%')
        """, (str(seq),))
        import time as _time
        base_id = int(_time.time() * 1000)
        inserted = 0
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            cur.execute("""
                INSERT INTO rule_chunks (id,rule_title,seq,article,department,url,content,embedding)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s::vector)
            """, (f"r{seq}_rev_{base_id+i}", title, str(seq),
                  _extract_article_title(chunk, i), dept, url, chunk, str(emb)))
            inserted += 1
        conn.commit()
        return inserted
    except HTTPException:
        conn.rollback()
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(500, f"DB 재색인 실패: {e}")
    finally:
        conn.close()


# -- 규정 개정 미리보기 (파일/텍스트 입력 → 현행과 비교만, 적용 X) -----
@app.post("/revise-preview")
async def revise_preview(
    seq: int = Form(...),
    text: str = Form(None),
    file: UploadFile = File(None),
    payload: dict = Depends(verify_token),
):
    """파일 또는 텍스트를 받아 '현행 본문' vs '개정안' 미리보기를 반환한다.
    실제 DB는 건드리지 않는다. 사용자가 비교 후 확인하면 별도로 /revise-regulation 호출."""

    # ① 개정안 본문 추출
    new_content = ""
    extract_method = ""
    if file is not None and (file.filename or ""):
        try:
            new_content = _extract_text(file)
            extract_method = f"파일({file.filename})"
        except HTTPException as e:
            # 추출 실패 시 상세 에러 그대로 전달
            raise
        except Exception as e:
            raise HTTPException(400, f"파일 읽기 실패: {e}")
    elif text and text.strip():
        new_content = text.strip()
        extract_method = "직접 입력"
    else:
        raise HTTPException(400, "개정안(파일 또는 텍스트)이 필요합니다.")

    if not new_content or len(new_content.strip()) < 10:
        raise HTTPException(400, "개정안 내용이 너무 짧습니다.")

    # ② 현행 규정 본문 로드
    history = _load_history_json()
    target = None
    for r in history:
        if int(r.get("seq", -1)) == seq:
            target = r
            break
    if not target:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")

    latest_idx = _get_latest_version_index(target)
    latest = target["versions"][latest_idx]
    current_content = latest.get("content", "") or ""

    # ③ 하이브리드 diff (어절 → 글자) — 자세한 설명은 _word_then_char_diff 참고
    import difflib
    segments = _word_then_char_diff(current_content, new_content)
    similarity = difflib.SequenceMatcher(None, current_content, new_content).ratio()

    return {
        "seq": seq,
        "title": target.get("title", ""),
        "current_content": current_content,
        "current_length": len(current_content),
        "current_revision_date": latest.get("revision_date", ""),
        "new_content": new_content,
        "new_length": len(new_content),
        "extract_method": extract_method,
        "segments": segments,
        "similarity": round(similarity * 100, 1),
    }


# -- 조 단위 개정 미리보기 (DB 적용 X, 비교만) -----------------------
@app.post("/revise-article-preview")
async def revise_article_preview(
    seq: int = Form(...),
    index: int = Form(...),
    text: str = Form(...),
    payload: dict = Depends(verify_token),
):
    """조 단위 개정 미리보기 — 현행 조항 vs 개정안 조항을 비교한다.
    실제 적용 시(/revise-article) 머리글에 '(개정 YYYY-MM-DD)'가 부착되므로,
    미리보기에서도 동일한 표시를 부착한 상태로 비교한다."""
    new_text = (text or "").strip()
    if len(new_text) < 5:
        raise HTTPException(400, "개정할 조항 내용이 너무 짧습니다.")

    history = _load_history_json()
    target = None
    for r in history:
        if int(r.get("seq", -1)) == seq:
            target = r
            break
    if target is None:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")

    latest_idx = _get_latest_version_index(target)
    latest = target["versions"][latest_idx]
    content = latest.get("content", "") or ""

    arts = _split_articles(content)
    if index < 0 or index >= len(arts):
        raise HTTPException(404, "해당 조를 찾을 수 없습니다.")
    a = arts[index]

    # 현행 조항 텍스트
    current_article_text = content[a["start"]:a["end"]].rstrip("\n")

    # 개정안 조항 텍스트 — 실제 적용과 동일하게 (개정 YYYY-MM-DD) 부착
    today = datetime.now().strftime("%Y-%m-%d")
    nt_lines = new_text.split("\n")
    head_re = _re.compile(r'^(제\s*\d+\s*조(?:의\s*\d+)?\s*(?:\([^)]*\))?)')
    if nt_lines and head_re.match(nt_lines[0].strip()):
        first = nt_lines[0].rstrip()
        first = _re.sub(r'\s*\(개정\s*\d{4}-\d{2}-\d{2}\)\s*$', '', first)
        nt_lines[0] = f"{first} (개정 {today})"
        new_article_text = "\n".join(nt_lines)
    else:
        new_article_text = f"{a['head']} (개정 {today})\n{new_text}"

    # 하이브리드 diff (어절 → 글자)
    import difflib
    segments = _word_then_char_diff(current_article_text, new_article_text)
    similarity = difflib.SequenceMatcher(None, current_article_text, new_article_text).ratio()

    return {
        "seq": seq,
        "index": index,
        "title": (target.get("title", "") + " — " + a["head"] + a["dup_label"]),
        "current_content": current_article_text,
        "current_length": len(current_article_text),
        "current_revision_date": latest.get("revision_date", ""),
        "new_content": new_article_text,
        "new_length": len(new_article_text),
        "extract_method": "직접 입력 (조 단위)",
        "segments": segments,
        "similarity": round(similarity * 100, 1),
    }


# -- 규정 개정 (전체) ----------------------------------------------
@app.post("/revise-regulation")
async def revise_regulation(
    seq: int = Form(...),
    text: str = Form(None),
    file: UploadFile = File(None),
    payload: dict = Depends(verify_token),
):
    """
    기존 규정(seq)을 새 내용으로 개정한다.
    - file 또는 text 중 하나로 새 본문을 받는다.
    - history.json의 최신 버전 content 를 새 내용으로 덮어쓴다.
    - 그 규정에 '이 날짜에 바뀜' 한 줄만 revisions[] 에 기록한다.
    - 개정 전 상태를 .revision_backup/ 에 백업한다 (되돌리기용).
    - DB 본문 청크도 재색인한다.
    """
    new_content = ""
    if file is not None and (file.filename or ""):
        try:
            new_content = _extract_text(file)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, f"파일 읽기 실패: {e}")
    elif text and text.strip():
        new_content = text.strip()
    else:
        raise HTTPException(400, "개정할 내용(파일 또는 텍스트)이 필요합니다.")

    if not new_content or len(new_content.strip()) < 10:
        raise HTTPException(400, "개정 내용이 너무 짧습니다.")

    history = _load_history_json()
    target_idx = None
    for i, r in enumerate(history):
        if int(r.get("seq", -1)) == seq:
            target_idx = i
            break
    if target_idx is None:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")

    target = history[target_idx]
    latest_idx = _get_latest_version_index(target)
    latest = target["versions"][latest_idx]

    title = target.get("title", "")
    dept = target.get("department", "")
    url = latest.get("url") or target.get("url_latest", "")
    today = datetime.now().strftime("%Y-%m-%d")

    # 백업
    os.makedirs(REVISION_BACKUP_DIR, exist_ok=True)
    backup_id = f"{seq}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    try:
        conn = psycopg2.connect(DB_URL)
        cur = conn.cursor()
        cur.execute("""
            SELECT id,rule_title,seq,article,department,url,content
            FROM rule_chunks
            WHERE seq = %s AND article != '개정이력'
              AND (url IS NULL OR url NOT LIKE 'upload://%%')
        """, (str(seq),))
        db_rows = [
            {"id": c[0], "rule_title": c[1], "seq": c[2], "article": c[3],
             "department": c[4], "url": c[5], "content": c[6]}
            for c in cur.fetchall()
        ]
        conn.close()
    except Exception as e:
        raise HTTPException(500, f"백업용 DB 조회 실패: {e}")

    backup = {
        "backup_id":     backup_id,
        "seq":           seq,
        "title":         title,
        "backed_up_at":  datetime.now().isoformat(timespec="seconds"),
        "kind":          "regulation_history",
        "history_item":  target,        # 개정 전 history 항목 통째로
        "db_chunks":     db_rows,
    }
    with open(os.path.join(REVISION_BACKUP_DIR, f"{backup_id}.json"),
              "w", encoding="utf-8") as f:
        _json.dump(backup, f, ensure_ascii=False, indent=2)

    # history.json 갱신 — 최신 버전 content 덮어쓰기 + 날짜 기록
    target["versions"][latest_idx]["content"] = new_content
    target["versions"][latest_idx]["revision_date"] = today

    revisions = target.get("revisions", [])
    revisions.append({
        "revised_at":       today,
        "prev_content_len": len(latest.get("content", "") or ""),
        "new_content_len":  len(new_content),
        "backup_id":        backup_id,
    })
    target["revisions"] = revisions
    history[target_idx] = target
    _save_history_json(history)

    # DB 재색인
    new_chunks = _reindex_rule_chunks(seq, title, dept, url, new_content)

    return {
        "success": True, "seq": seq, "title": title,
        "revised_at": today, "backup_id": backup_id,
        "new_chunks": new_chunks,
    }


# -- 되돌리기: 백업 목록 -------------------------------------------
@app.get("/revision-backups")
def revision_backups(payload: dict = Depends(verify_token)):
    """활동 로그 — 되돌릴 수 있는 개정 백업 + 업로드된 규정을 한 피드로 반환 (최신순)"""
    items = []
    # ── 1) 개정/치환 백업 ──
    if os.path.isdir(REVISION_BACKUP_DIR):
        for fn in os.listdir(REVISION_BACKUP_DIR):
            if not fn.endswith(".json"):
                continue
            try:
                with open(os.path.join(REVISION_BACKUP_DIR, fn), "r", encoding="utf-8") as f:
                    b = _json.load(f)
                items.append({
                    "backup_id":    b.get("backup_id"),
                    "seq":          b.get("seq"),
                    "title":        b.get("title"),
                    "backed_up_at": b.get("backed_up_at"),
                    "kind":         b.get("kind", ""),
                    "extras":       b.get("extras", {}),
                })
            except Exception:
                continue
    # ── 2) 업로드된 규정 — DB의 first_id에서 unix-ms 타임스탬프 복원 ──
    try:
        from datetime import datetime as _dt
        conn = psycopg2.connect(DB_URL); cur = conn.cursor()
        cur.execute("""
            SELECT url, COUNT(*) as cnt, MIN(id) as first_id
            FROM rule_chunks WHERE url LIKE 'upload://%'
            GROUP BY url
        """)
        for r in cur.fetchall():
            url, cnt, first_id = r[0], r[1], r[2]
            filename = url.replace("upload://", "")
            try:
                ts = _dt.fromtimestamp(int(first_id) / 1000.0).strftime("%Y-%m-%dT%H:%M:%S")
            except Exception:
                ts = ""
            items.append({
                "backup_id":    None,        # 되돌리기 백업 ID 없음 — 삭제로 대신
                "seq":          None,
                "title":        filename,
                "backed_up_at": ts,
                "kind":         "upload",
                "extras":       {"filename": filename, "chunks": cnt},
            })
        conn.close()
    except Exception:
        pass  # DB 오류 시 업로드 부분만 빠짐, 백업은 그대로 반환
    items.sort(key=lambda x: x.get("backed_up_at") or "", reverse=True)
    return {"backups": items}


# -- 되돌리기 실행 -------------------------------------------------
def _restore_one(seq, json_item: dict, db_chunks: list, rules: list) -> bool:
    """규정 1개를 JSON + DB 모두 개정 전 상태로 복원. (JSON은 rules 리스트를 직접 수정)"""
    restored = False
    for i, r in enumerate(rules):
        if int(r.get("seq", -1)) == int(seq):
            rules[i] = json_item
            restored = True
            break
    if not restored:
        return False

    conn = psycopg2.connect(DB_URL)
    cur = conn.cursor()
    try:
        cur.execute("""
            DELETE FROM rule_chunks
            WHERE seq = %s AND article != '개정이력'
              AND (url IS NULL OR url NOT LIKE 'upload://%%')
        """, (str(seq),))
        for c in db_chunks:
            emb = upstage_embed_passage([c["content"]])[0]
            cur.execute("""
                INSERT INTO rule_chunks (id,rule_title,seq,article,department,url,content,embedding)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s::vector)
            """, (c["id"], c["rule_title"], c["seq"], c["article"],
                  c["department"], c["url"], c["content"], str(emb)))
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()
    return True


@app.post("/revision-rollback")
def revision_rollback(backup_id: str = Form(...), payload: dict = Depends(verify_token)):
    """백업을 사용해 JSON + DB를 개정 전 상태로 되돌린다.
       단일/일괄 + rules.json 기반/history.json 기반 백업 모두 지원."""
    backup_path = os.path.join(REVISION_BACKUP_DIR, f"{backup_id}.json")
    if not os.path.exists(backup_path):
        raise HTTPException(404, "해당 백업을 찾을 수 없습니다.")

    with open(backup_path, "r", encoding="utf-8") as f:
        backup = _json.load(f)

    kind = backup.get("kind", "")

    try:
        # ── history.json 기반 백업 (전체개정 / 조 단위 개정) ──
        if kind in ("regulation_history", "article_history"):
            history = _load_history_json()
            seq = backup["seq"]
            saved_target = backup.get("history_item")
            if not saved_target:
                raise HTTPException(400, "백업에 history_item 이 없습니다.")
            # history 안에서 해당 seq 찾아서 통째로 교체
            replaced = False
            for i, r in enumerate(history):
                if int(r.get("seq", -1)) == int(seq):
                    history[i] = saved_target
                    replaced = True
                    break
            if not replaced:
                raise HTTPException(404, f"seq={seq} 규정이 현재 history.json에 없습니다.")
            _save_history_json(history)

            # DB도 백업 청크로 복원
            conn = psycopg2.connect(DB_URL); cur = conn.cursor()
            try:
                cur.execute("""
                    DELETE FROM rule_chunks
                    WHERE seq = %s AND article != '개정이력'
                      AND (url IS NULL OR url NOT LIKE 'upload://%%')
                """, (str(seq),))
                for c in backup.get("db_chunks", []):
                    emb = upstage_embed_passage([c["content"]])[0]
                    cur.execute("""
                        INSERT INTO rule_chunks (id,rule_title,seq,article,department,url,content,embedding)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s::vector)
                    """, (c["id"], c["rule_title"], c["seq"], c["article"],
                          c["department"], c["url"], c["content"], str(emb)))
                conn.commit()
            finally:
                conn.close()
            result = {"success": True, "backup_id": backup_id, "kind": kind,
                      "seq": seq, "restored_chunks": len(backup.get("db_chunks", []))}

        # ── 일괄 치환 (rules.json 기반) ──
        elif kind == "bulk_replace":
            rules = _load_rules_json()
            items = backup.get("items", [])
            restored_cnt = 0
            for it in items:
                ok = _restore_one(it["seq"], it["json_item"],
                                  it.get("db_chunks", []), rules)
                if ok:
                    restored_cnt += 1
            _save_rules_json(rules)
            result = {"success": True, "backup_id": backup_id,
                      "kind": "bulk_replace", "restored_rules": restored_cnt}

        # ── 옛 단일 규정 백업 (rules.json 기반) ──
        else:
            rules = _load_rules_json()
            ok = _restore_one(backup["seq"], backup["json_item"],
                              backup.get("db_chunks", []), rules)
            if not ok:
                raise HTTPException(404, f"seq={backup.get('seq')} 규정이 현재 JSON에 없습니다.")
            _save_rules_json(rules)
            result = {"success": True, "backup_id": backup_id,
                      "seq": backup["seq"],
                      "restored_chunks": len(backup.get("db_chunks", []))}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"되돌리기 실패: {e}")

    try:
        os.remove(backup_path)
    except Exception:
        pass

    return result




# ══════════════════════════════════════════════════════════════════
# 조(條) 단위 개정 — 규정 본문을 조 단위로 분할 / 특정 조만 개정
# ══════════════════════════════════════════════════════════════════
def _split_articles(content: str) -> list:
    """
    규정 본문을 조(條) 단위로 분할.
    줄 시작이 '제 N 조' 인 줄만 조 머리글로 인정 (본문 인용 '제2조' 제외).
    반환: [{index, head, body, start, end}]  (start/end = content 내 문자 오프셋)
    """
    head_re = _re.compile(r'^제\s*\d+\s*조(?:의\s*\d+)?\s*(?:\([^)]*\))?')
    arts = []
    pos = 0
    cur = None
    for ln in content.split("\n"):
        line_len = len(ln) + 1  # +1 = '\n'
        s = ln.strip()
        if head_re.match(s):
            if cur is not None:
                cur["end"] = pos
                arts.append(cur)
            cur = {"head": s, "start": pos, "end": None}
        pos += line_len
    if cur is not None:
        cur["end"] = len(content)
        arts.append(cur)

    # 중복 조 번호에 [n/m] 라벨 부여
    num_re = _re.compile(r'^제\s*(\d+)\s*조(?:의\s*(\d+))?')
    from collections import Counter as _C
    keys = []
    for a in arts:
        m = num_re.match(a["head"])
        keys.append((m.group(1), m.group(2)) if m else (a["head"], None))
    total = _C(keys)
    seen = {}
    out = []
    for i, a in enumerate(arts):
        k = keys[i]
        body = content[a["start"]:a["end"]]
        dup_label = ""
        if total[k] > 1:
            seen[k] = seen.get(k, 0) + 1
            dup_label = f" [{seen[k]}/{total[k]}]"
        out.append({
            "index":     i,
            "head":      a["head"],
            "dup_label": dup_label,
            "start":     a["start"],
            "end":       a["end"],
            "body_len":  len(body.strip()),
        })
    return out


# -- 특정 규정의 조 목록 (최신 버전 기준) ---------------------------
@app.get("/rule-articles/{seq}")
def rule_articles(seq: int, payload: dict = Depends(verify_token)):
    history = _load_history_json()
    target = None
    for r in history:
        if int(r.get("seq", -1)) == seq:
            target = r
            break
    if target is None:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")
    latest = _get_latest_version(target)
    if not latest:
        raise HTTPException(404, "최신 버전을 찾을 수 없습니다.")
    content = latest.get("content", "") or ""
    arts = _split_articles(content)
    return {
        "seq":      seq,
        "title":    target.get("title", ""),
        "articles": arts,
        "count":    len(arts),
    }


# -- 특정 조 1개의 현재 본문 (최신 버전 기준) -----------------------
@app.get("/rule-articles/{seq}/{index}")
def rule_article_detail(seq: int, index: int, payload: dict = Depends(verify_token)):
    history = _load_history_json()
    target = None
    for r in history:
        if int(r.get("seq", -1)) == seq:
            target = r
            break
    if target is None:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")
    latest = _get_latest_version(target)
    if not latest:
        raise HTTPException(404, "최신 버전을 찾을 수 없습니다.")
    content = latest.get("content", "") or ""
    arts = _split_articles(content)
    if index < 0 or index >= len(arts):
        raise HTTPException(404, "해당 조를 찾을 수 없습니다.")
    a = arts[index]
    return {
        "seq":   seq,
        "index": index,
        "head":  a["head"] + a["dup_label"],
        "text":  content[a["start"]:a["end"]],
    }


# -- 조 단위 개정 --------------------------------------------------
@app.post("/revise-article")
async def revise_article(
    seq: int = Form(...),
    index: int = Form(...),
    text: str = Form(...),
    payload: dict = Depends(verify_token),
):
    """
    규정(seq) 안의 index번째 조항만 새 내용으로 교체한다 (history.json 최신 버전 기준).
    - 개정된 조 머리글 옆에 (개정 YYYY-MM-DD) 표시를 붙인다.
    - 개정 전 상태를 .revision_backup/ 에 백업한다.
    - DB는 해당 규정 본문 청크를 통째로 재색인한다.
    - revisions[] 에 '이 조가 이 날 바뀜' 한 줄 기록.
    """
    new_text = (text or "").strip()
    if len(new_text) < 5:
        raise HTTPException(400, "개정할 조항 내용이 너무 짧습니다.")

    history = _load_history_json()
    target_idx = None
    for i, r in enumerate(history):
        if int(r.get("seq", -1)) == seq:
            target_idx = i
            break
    if target_idx is None:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")

    target = history[target_idx]
    latest_idx = _get_latest_version_index(target)
    latest = target["versions"][latest_idx]

    title = target.get("title", "")
    dept = target.get("department", "")
    url = latest.get("url") or target.get("url_latest", "")
    content = latest.get("content", "") or ""

    arts = _split_articles(content)
    if index < 0 or index >= len(arts):
        raise HTTPException(404, "해당 조를 찾을 수 없습니다.")
    a = arts[index]
    today = datetime.now().strftime("%Y-%m-%d")

    # 새 조항 본문에 (개정 YYYY-MM-DD) 표시 — 머리글 줄 끝에 부착
    nt_lines = new_text.split("\n")
    head_re = _re.compile(r'^(제\s*\d+\s*조(?:의\s*\d+)?\s*(?:\([^)]*\))?)')
    if nt_lines and head_re.match(nt_lines[0].strip()):
        first = nt_lines[0].rstrip()
        # 기존 (개정 ...) 표시 있으면 제거 후 새로 부착
        first = _re.sub(r'\s*\(개정\s*\d{4}-\d{2}-\d{2}\)\s*$', '', first)
        nt_lines[0] = f"{first} (개정 {today})"
        new_article = "\n".join(nt_lines)
    else:
        # 머리글이 없으면 원래 머리글을 살려 붙임
        new_article = f"{a['head']} (개정 {today})\n{new_text}"

    new_content = content[:a["start"]] + new_article
    if not new_content.endswith("\n"):
        new_content += "\n"
    new_content += content[a["end"]:]

    # 백업 (history 항목 통째로 + DB 본문 청크)
    os.makedirs(REVISION_BACKUP_DIR, exist_ok=True)
    backup_id = f"{seq}_art{index}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    try:
        conn = psycopg2.connect(DB_URL)
        cur = conn.cursor()
        cur.execute("""
            SELECT id,rule_title,seq,article,department,url,content
            FROM rule_chunks
            WHERE seq = %s AND article != '개정이력'
              AND (url IS NULL OR url NOT LIKE 'upload://%%')
        """, (str(seq),))
        db_rows = [
            {"id": c[0], "rule_title": c[1], "seq": c[2], "article": c[3],
             "department": c[4], "url": c[5], "content": c[6]}
            for c in cur.fetchall()
        ]
        conn.close()
    except Exception as e:
        raise HTTPException(500, f"백업용 DB 조회 실패: {e}")

    backup = {
        "backup_id":     backup_id,
        "seq":           seq,
        "title":         f"{title} — {a['head']}{a['dup_label']}",
        "backed_up_at":  datetime.now().isoformat(timespec="seconds"),
        "kind":          "article_history",
        "history_item":  target,
        "db_chunks":     db_rows,
    }
    with open(os.path.join(REVISION_BACKUP_DIR, f"{backup_id}.json"),
              "w", encoding="utf-8") as f:
        _json.dump(backup, f, ensure_ascii=False, indent=2)

    # history.json 갱신 — 최신 버전 content 만 덮어쓰기
    target["versions"][latest_idx]["content"] = new_content
    target["versions"][latest_idx]["revision_date"] = today

    # revisions[]에 '어떤 조가 언제 바뀌었는지' 기록 (나중에 '언제 바뀌었나' 질문에 답할 수 있도록)
    revisions = target.get("revisions", [])
    revisions.append({
        "kind":         "article",
        "article_head": a["head"] + a["dup_label"],
        "revised_at":   today,
        "backup_id":    backup_id,
    })
    target["revisions"] = revisions
    history[target_idx] = target
    _save_history_json(history)

    # DB 본문 청크 재색인
    new_chunks = _reindex_rule_chunks(seq, title, dept, url, new_content)

    return {
        "success": True, "seq": seq, "index": index,
        "article": a["head"] + a["dup_label"],
        "revised_at": today, "backup_id": backup_id,
        "new_chunks": new_chunks,
    }




# ══════════════════════════════════════════════════════════════════
# 조 단위 개정 — 파일 업로드 + AI 개정안 추출
# ══════════════════════════════════════════════════════════════════
@app.post("/extract-article-revision")
async def extract_article_revision(
    seq: int = Form(...),
    index: int = Form(...),
    file: UploadFile = File(...),
    payload: dict = Depends(verify_token),
):
    """
    개정 문서 파일을 받아:
      1) 파일에서 전체 텍스트를 추출한다 (원본 그대로 반환).
      2) AI가 그 문서에서 '대상 조항의 개정안' 부분을 찾아 추출한다.
    실제 개정은 하지 않는다. 프론트의 편집칸을 채우는 용도.
    """
    # 1) 대상 규정 / 조항 확인 (history.json 최신 버전 기준)
    history = _load_history_json()
    target = None
    for r in history:
        if int(r.get("seq", -1)) == seq:
            target = r
            break
    if target is None:
        raise HTTPException(404, f"seq={seq} 규정을 찾을 수 없습니다.")

    latest = _get_latest_version(target)
    if not latest:
        raise HTTPException(404, "최신 버전을 찾을 수 없습니다.")
    content = latest.get("content", "") or ""
    arts = _split_articles(content)
    if index < 0 or index >= len(arts):
        raise HTTPException(404, "해당 조를 찾을 수 없습니다.")
    a = arts[index]
    art_head = a["head"] + a["dup_label"]
    current_text = content[a["start"]:a["end"]].strip()

    # 2) 파일 텍스트 추출
    try:
        raw_text = _extract_text(file)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"파일 읽기 실패: {e}")

    if not raw_text or len(raw_text.strip()) < 10:
        raise HTTPException(400, "파일에서 텍스트를 추출하지 못했습니다.")

    raw_text = raw_text.strip()

    # 3) AI 개정안 추출
    #    문서가 길면 앞부분만 (모델 토큰 보호)
    doc_for_ai = raw_text[:6000]
    ai_text = ""
    ai_note = ""
    try:
        prompt = f"""너는 대학 규정 개정 문서를 분석하는 도우미다.

아래는 어떤 규정 개정 관련 문서에서 추출한 텍스트다. 이 문서에는 회의 안건 양식,
개정 사유, 현행규정과 개정(안)이 섞여 있을 수 있다.

[목표]
'{target.get("title","")}' 규정의 다음 조항에 해당하는 '개정(안)' 본문만 정확히 추출하라.

[대상 조항]
{art_head}

[현재 규정상의 해당 조항 원문 (참고용)]
{current_text[:800]}

[분석할 문서 텍스트]
{doc_for_ai}

[규칙]
- 문서에 '현행규정'과 '개정(안)'이 함께 있으면, 반드시 '개정(안)' 쪽 내용만 골라라.
- 대상 조항({art_head})에 해당하는 부분만 추출하고, 다른 조항은 포함하지 마라.
- 회의 양식, 개정 사유, 협의 여부 같은 행정 문구는 제외하라.
- 조 머리글(제 N 조 (제목))부터 그 조항 본문 끝까지 그대로 출력하라.
- 문서에서 해당 조항의 개정안을 찾을 수 없으면 정확히 'NOT_FOUND' 한 단어만 출력하라.
- 설명·머리말 없이 추출한 본문만 출력하라."""
        ai_text = claude_chat(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0,
        )
        ai_text = (ai_text or "").strip()
        # 코드블록 마크다운 제거
        ai_text = _re.sub(r'^```[a-zA-Z]*\s*', '', ai_text)
        ai_text = _re.sub(r'\s*```$', '', ai_text).strip()

        if ai_text.upper().replace(" ", "") in ("NOT_FOUND", "NOTFOUND") or len(ai_text) < 5:
            ai_text = ""
            ai_note = "AI가 이 문서에서 해당 조항의 개정안을 찾지 못했습니다. 원본 텍스트를 보고 직접 입력하세요."
        else:
            ai_note = "AI가 추출한 개정안입니다. 반드시 원본과 대조해 확인·수정 후 개정하세요."
    except Exception as e:
        ai_text = ""
        ai_note = f"AI 추출에 실패했습니다({e}). 원본 텍스트를 보고 직접 입력하세요."

    return {
        "seq":          seq,
        "index":        index,
        "article_head": art_head,
        "ai_text":      ai_text,       # AI가 뽑은 개정안 (없으면 빈 문자열)
        "ai_note":      ai_note,       # 안내 문구
        "raw_text":     raw_text,      # 파일 원본 전체 텍스트
        "raw_truncated": len(raw_text) > 6000,
    }


# ══════════════════════════════════════════════════════════════════
# 회의록 기반 AI 개정 — 회의자료/회의록 업로드 → AI가 개정 항목 자동 추출
# ══════════════════════════════════════════════════════════════════
@app.post("/extract-meeting-revisions")
async def extract_meeting_revisions(
    file: UploadFile = File(...),
    payload: dict = Depends(verify_token),
):
    """
    규정 개정 관련 회의자료/회의록을 받아:
      1) (HWP인 경우) 표 구조 직접 파싱 — 현행/개정안 컬럼 분리
      2) 텍스트 추출 (폴백)
      3) Claude가 회의록에서 '어떤 규정의 어떤 조항을 어떻게 개정하는지' 자동 분석
      4) 항목 리스트로 반환 — 프론트가 각 항목을 글자 단위 diff 미리보기로 띄움
    """
    # 1) 파일 바이트 읽기 (표 파싱과 텍스트 추출 둘 다 사용)
    data = await file.read()
    fname_lower = (file.filename or "").lower()
    # _extract_text가 다시 읽을 수 있도록 stream 재구성
    from io import BytesIO
    file.file = BytesIO(data)

    # 2) HWP면 표 구조에서 현행/개정안 직접 분리 시도 (HTML 변환)
    parsed_tables = []
    if fname_lower.endswith(".hwp"):
        try:
            parsed_tables = _extract_hwp_tables_two_column(data)
        except Exception as e:
            print(f"[MEETING] HTML 표 파싱 실패(무시): {e}")
            parsed_tables = []

    # 3) 전체 텍스트 추출
    try:
        raw_text = _extract_text(file)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"파일 읽기 실패: {e}")
    if not raw_text or len(raw_text.strip()) < 10:
        raise HTTPException(400, "파일에서 텍스트를 추출하지 못했습니다.")
    raw_text = raw_text.strip()

    # 4) HTML 파싱이 빈약/실패하면 텍스트 기반 폴백 — '제 N 조' 두 번 등장 패턴
    if not parsed_tables:
        try:
            parsed_tables = _extract_meeting_tables_from_text(raw_text)
            if parsed_tables:
                print(f"[MEETING] 텍스트 파서로 안건 {len(parsed_tables)}건 분리 성공")
        except Exception as e:
            print(f"[MEETING] 텍스트 파서 실패(무시): {e}")
            parsed_tables = []

    return _analyze_meeting_minutes(raw_text, target_seq=-1, parsed_tables=parsed_tables)


# -- 회의록 직접 입력 → AI 분석 ----------------------------------
@app.post("/extract-meeting-revisions-text")
async def extract_meeting_revisions_text(
    text: str = Form(...),
    target_seq: int = Form(-1),
    payload: dict = Depends(verify_token),
):
    """
    회의록 본문을 텍스트로 직접 받아 AI 분석.
    target_seq를 지정하면 해당 규정의 안건만 추출하도록 AI에게 강하게 지시한다 (선택).
    """
    txt = (text or "").strip()
    if len(txt) < 10:
        raise HTTPException(400, "회의록 내용이 너무 짧습니다.")
    # 텍스트 파서로 표 자동 분리 시도 — '제 N 조' 두 번 등장 패턴
    parsed_tables = []
    try:
        parsed_tables = _extract_meeting_tables_from_text(txt)
        if parsed_tables:
            print(f"[MEETING] 직접 입력 텍스트 파서로 안건 {len(parsed_tables)}건 분리")
    except Exception as e:
        print(f"[MEETING] 텍스트 파서 실패(무시): {e}")
    return _analyze_meeting_minutes(txt, target_seq=target_seq, parsed_tables=parsed_tables)


# ══════════════════════════════════════════════════════════════════
# 회의록 분석 — 2단계 결정론적 파이프라인
#   1단계: 텍스트 → 안건별 [헤더 + 현행 + 개정안] 분리 (규칙 기반, AI 없음)
#   2단계: 개정안을 조 단위로 분리 + 시스템 규정 fuzzy 매칭 (AI 호출 안 함)
# AI는 회의록 본문 추출에 절대 관여하지 않음 — 개정안을 현행으로 잘못 잡는 실수 원천 차단
# ══════════════════════════════════════════════════════════════════
def _extract_rule_name_from_heading(heading: str) -> str:
    """회의록 안건 헤더에서 규정명만 추출.
    예: '20. [5-0-2] 인문과학연구원 규정 개정(안)' → '인문과학연구원 규정'"""
    if not heading:
        return ""
    h = heading.replace("\n", " ").replace("\r", " ")
    m = _re.search(r'\d+\s*\.\s*(?:\[[\d\-]+\]\s*)?(.+?)\s*개정\s*\(?\s*안\s*\)?', h)
    if m:
        return m.group(1).strip()
    return ""


def _extract_reason_from_heading(heading: str) -> str:
    """헤더에서 '개정 사유:' 줄을 추출"""
    if not heading:
        return ""
    m = _re.search(r'개정\s*사유\s*:\s*([^\n]+)', heading)
    return m.group(1).strip() if m else ""


def _match_rule_by_title(history: list, name: str) -> dict:
    """fuzzy match: 회의록 헤더의 규정명과 history.json의 title 매칭.
    1) 공백 무시 정확 매칭 → 2) substring 양방향 → 3) 키워드 토큰 매칭"""
    if not name:
        return None
    n_name = _re.sub(r'\s+', '', name)
    # 1) 정확 매칭
    for r in history:
        t = (r.get("title") or "").strip()
        if t and _re.sub(r'\s+', '', t) == n_name:
            return r
    # 2) substring 양방향 — 더 긴 매칭 우선
    candidates = []
    for r in history:
        t = (r.get("title") or "").strip()
        if not t:
            continue
        n_t = _re.sub(r'\s+', '', t)
        if (n_t and n_name) and (n_t in n_name or n_name in n_t):
            candidates.append((r, len(n_t)))
    if candidates:
        candidates.sort(key=lambda x: -x[1])
        return candidates[0][0]
    # 3) 핵심 키워드(첫 6자 이상 한글) 토큰 매칭
    head_token = _re.sub(r'[^가-힣]', '', name)[:8]
    if len(head_token) >= 4:
        for r in history:
            t = (r.get("title") or "").strip()
            if head_token in _re.sub(r'\s+', '', t):
                return r
    return None


def _split_text_to_articles(text: str) -> list:
    """텍스트를 '제 N 조 (...)' 또는 '부 칙' 머리글 기준으로 조 단위 분리.
    반환: [{"head": "제 5 조 (학점 부여)", "body": "제 5 조 ... 전체 본문"}, ...]"""
    if not text:
        return []
    head_pat = _re.compile(
        r'(?:^|\n)\s*(제\s*\d+\s*조(?:\s*의\s*\d+)?(?:\s*\([^)]*\))?|부\s*칙)',
        _re.MULTILINE
    )
    matches = list(head_pat.finditer(text))
    if not matches:
        return []
    out = []
    for i, m in enumerate(matches):
        head_raw = m.group(1).strip()
        # 공백 정리: "제  5  조 ( 학점 부여 )" → "제 5 조 (학점 부여)"
        head = _re.sub(r'\s+', ' ', head_raw)
        head = _re.sub(r'\(\s+', '(', head)
        head = _re.sub(r'\s+\)', ')', head)
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        body = text[m.start():end].rstrip()
        if body:
            out.append({"head": head, "body": body})
    return out








# ══════════════════════════════════════════════════════════════════
# 회의록 분석 — AI 주도 4단계 파이프라인
#   ① AI가 회의록 양식 판단 (표 vs 줄글) → 파싱 방식 결정
#   ② 1차: 어떤 규정인지 판별 / 2차: 개정안 본문만 추출 (★ 원문은 절대 추출 X)
#   ③ 시스템에서 history.json의 현행 본문을 자동으로 가져옴
#   ④ 결과: [시스템 현행] vs [AI가 뽑은 개정안] 으로 미리보기 비교
# "AI가 회의록의 현행을 잘못 가져오는 문제"는 AI에게 현행 추출 작업 자체를
# 시키지 않음으로 원천 차단. 현행의 정답은 시스템에 이미 있다.
# 보조 자료: parsed_tables(서버가 HWP 표 또는 텍스트 패턴으로 분리한 좌/우)
# → AI가 양식 판단할 때 참고만. AI의 작업은 '개정안 추출'에 집중.
# ══════════════════════════════════════════════════════════════════
def _analyze_meeting_minutes(raw_text: str, target_seq: int = -1, parsed_tables: list = None) -> dict:
    parsed_tables = parsed_tables or []

    # 업로드 규정 자동 백필 — 매칭 후보에 포함되도록
    try:
        added = _sync_uploaded_rules_to_history()
        if added > 0:
            print(f"[ANALYZE] 회의록 분석 직전 업로드 규정 {added}건 백필됨")
    except Exception as e:
        print(f"[ANALYZE] 백필 실패(무시): {e}")

    history = _load_history_json()
    rule_titles = []
    for r in history[:800]:
        t = (r.get("title") or "").strip()
        if t:
            rule_titles.append({
                "seq": int(r.get("seq", -1)),
                "title": t,
                "uploaded": bool(r.get("_uploaded")),
            })

    # 사용자가 특정 규정을 콕 찍은 경우
    target_focus_block = ""
    if target_seq is not None and target_seq >= 0:
        focus = next((r for r in history if int(r.get("seq", -1)) == int(target_seq)), None)
        if focus:
            target_focus_block = (
                f"\n[★ 사용자가 지정한 대상 규정]\n"
                f"- seq={target_seq}, 제목=\"{focus.get('title','')}\"\n"
                f"  → 이 규정의 안건만 추출하라. rule_seq는 반드시 {target_seq}.\n"
            )

    # 보조 자료: 서버가 미리 분리한 표 좌/우 (AI가 양식 판단할 때 참고)
    tables_hint = ""
    if parsed_tables:
        chunks = []
        for i, t in enumerate(parsed_tables, 1):
            chunks.append(
                f"\n══ [참고용 표 {i}] 헤더 컨텍스트 ══\n{(t.get('heading_before') or '')[:400]}\n"
                f"══ [참고용 표 {i}] 좌측 셀(=현행 — 너는 절대 추출 X) ══\n{(t.get('current') or '')[:3000]}\n"
                f"══ [참고용 표 {i}] 우측 셀(=개정안 — 여기서만 추출) ══\n{(t.get('new') or '')[:4000]}\n"
            )
        tables_hint = (
            "\n[★ 서버가 미리 분리한 좌/우 셀 — 양식 판단 보조 자료]\n"
            "표 양식일 때 좌측=현행, 우측=개정안. 너는 우측에서만 본문을 가져와라.\n"
            + "".join(chunks)
        )

    titles_for_ai = "\n".join([
        f"- [seq={x['seq']}] {x['title']}" + (" (업로드)" if x.get('uploaded') else "")
        for x in rule_titles
    ])
    doc_for_ai = raw_text[:16000]

    import json as _pyjson
    prompt = f"""너는 대학 규정 개정 회의록을 분석하는 도우미다.

[너의 임무 — 4단계 정확히 수행]

(1) 회의록 양식 판단
    먼저 회의록이 어떤 양식인지 판단하라:
    - 표 양식: "현행규정 | 개정(안)" 두 컬럼이 표로 나란히
    - 줄글 양식: "현행:" 단락 다음 "개정(안):" 단락 식으로 순차
    이 판단에 따라 파싱 방식이 달라진다. format_detected에 'table' 또는 'prose' 적어라.

(2) 1차 — 규정 판별
    회의록에서 다루는 규정을 [시스템 규정 목록]에서 찾아 rule_seq에 적어라.
    안건이 여러 개면 각각 다른 rule_seq를 가질 수 있다. 매칭 애매하면 match_confidence='low'.

(3) 2차 — 개정안 본문만 추출 ★★★ 가장 중요
    ★ 원문(현행)은 절대 추출 X. 시스템에 이미 있어서 우리가 알아서 가져온다.
    ★ 오직 '개정(안)' 컬럼/섹션의 본문만 new_text에 채워라.
    - 표 양식: 우측 컬럼(개정안)의 본문만
    - 줄글 양식: "개정(안):" 헤더 다음 본문만
    회의록에 같은 조 번호가 두 번 등장하면(첫=현행, 둘=개정안), 반드시 두 번째 것만.
    new_text가 회의록의 현행 본문과 비슷하면 잘못 뽑은 것 — 다시 확인하라.

(4) 조 단위 분리
    한 안건에 여러 조 개정이 있으면 반드시 조마다 별개 항목으로.
    현행과 개정안이 완전히 동일한 조는 빼라.
    부 칙도 별개 항목 (article_head="부 칙", scope="article").

(5) 텍스트 정돈 — new_text를 깔끔하게 다듬어라
    PDF/HWP에서 추출된 텍스트는 줄바꿈과 공백이 이상하게 끊겨 있는 경우가 많다.
    추출한 new_text는 사람이 읽기 좋게 다음 규칙으로 정리하라:
    - 어절(단어) 가운데에서 끊긴 줄바꿈은 제거하라. 예: "협동을 통해 종합적\\n연구를" → "협동을 통해 종합적 연구를"
    - 한 문장 안의 부적절한 줄바꿈은 공백으로 바꾼다.
    - 항목 번호("1. ", "2. ", "①" 등)와 새 조("제 N 조") 앞에는 줄바꿈을 넣어 가독성 확보.
    - 연속된 공백은 한 칸으로.
    - 의미가 깨지지 않는 선에서, 원문에 없는 내용은 절대 추가하지 마라 (정돈만 하라).
{target_focus_block}{tables_hint}

[시스템 규정 목록 — 제목으로 매칭]
{titles_for_ai}

[회의록 텍스트 — 분석 대상]
{doc_for_ai}

[출력 JSON 스키마 — JSON 배열만, 코드블록·머리말 금지]
[
  {{
    "rule_seq":         정수 (시스템 규정 목록의 seq 그대로),
    "rule_title":       "매칭된 규정명",
    "article_head":     "제 5 조 (학점 부여)" 또는 "부 칙",
    "scope":            "article" (기본) | "regulation" (전체 다시 쓴 경우만),
    "new_text":         "★ 개정안 본문만. 현행 절대 포함 X. 조 머리글부터 끝까지.",
    "reason":           "개정 사유 한 줄",
    "match_confidence": "high" | "medium" | "low",
    "format_detected":  "table" | "prose"
  }}
]

위 형식 그대로 JSON 배열만 출력하라."""

    try:
        ai_raw = claude_chat(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=6000,
            temperature=0,
        )
    except Exception as e:
        raise HTTPException(500, f"AI 분석 실패: {e}")
    ai_raw = (ai_raw or "").strip()
    if ai_raw.startswith("```"):
        ai_raw = _re.sub(r'^```(?:json)?\s*', '', ai_raw)
        ai_raw = _re.sub(r'```\s*$', '', ai_raw).strip()

    try:
        ai_items = _pyjson.loads(ai_raw)
        if not isinstance(ai_items, list):
            raise ValueError("배열이 아님")
    except Exception as e:
        return {
            "items": [],
            "raw_text": raw_text,
            "raw_truncated": len(raw_text) > 16000,
            "table_count": len(parsed_tables),
            "ai_error": f"AI 응답 파싱 실패: {e}",
            "ai_raw": ai_raw[:2000],
            "stage1": [],
            "format_detected": "",
        }

    # ── 시스템에서 현행 본문 자동 매칭 + 채움 ──
    enriched = []
    stage1_debug = []
    detected_formats = set()
    for it in ai_items:
        try:
            seq = int(it.get("rule_seq", -1))
        except Exception:
            seq = -1
        rule_title_ai = (it.get("rule_title") or "").strip()
        article_head = (it.get("article_head") or "").strip()
        scope = (it.get("scope") or "article").lower().strip()
        new_text = (it.get("new_text") or "").strip()
        # 가벼운 후처리 — AI가 놓친 줄바꿈/공백 마무리
        if new_text:
            # 연속 공백을 한 칸으로 (탭 포함, 줄바꿈은 보존)
            new_text = _re.sub(r'[ \t]+', ' ', new_text)
            # 줄 끝 공백 제거
            new_text = _re.sub(r' +\n', '\n', new_text)
            # 줄바꿈이 3개 이상이면 2개로 (단락 구분만 유지)
            new_text = _re.sub(r'\n{3,}', '\n\n', new_text)
            new_text = new_text.strip()
        reason = (it.get("reason") or "").strip()
        confidence = (it.get("match_confidence") or "medium").lower().strip()
        fmt = (it.get("format_detected") or "").lower().strip()
        if fmt:
            detected_formats.add(fmt)

        # seq로 규정 찾기 (seq 매칭이 안 되면 제목 fuzzy)
        target = next((r for r in history if int(r.get("seq", -1)) == seq), None)
        if target is None and rule_title_ai:
            target = _match_rule_by_title(history, rule_title_ai)
            if target is not None:
                seq = int(target.get("seq", -1))

        article_index = -1
        current_text = ""
        if target is not None:
            latest = _get_latest_version(target)
            content = (latest or {}).get("content", "") or ""
            if scope == "regulation":
                current_text = content
            else:
                sys_arts = _split_articles(content)
                norm_head = _re.sub(r'\s+', '', article_head)
                for ai_idx, sa in enumerate(sys_arts):
                    sa_head = sa["head"] + sa.get("dup_label", "")
                    if _re.sub(r'\s+', '', sa_head) == norm_head:
                        article_index = ai_idx
                        current_text = content[sa["start"]:sa["end"]].rstrip()
                        break
                if article_index == -1:
                    m_a = _re.search(r'제\s*(\d+)\s*조(?:\s*의\s*(\d+))?', article_head)
                    if m_a:
                        wa, wb = m_a.group(1), m_a.group(2) or ""
                        for ai_idx, sa in enumerate(sys_arts):
                            m_s = _re.search(r'제\s*(\d+)\s*조(?:\s*의\s*(\d+))?', sa["head"])
                            if m_s and m_s.group(1) == wa and (m_s.group(2) or "") == wb:
                                article_index = ai_idx
                                current_text = content[sa["start"]:sa["end"]].rstrip()
                                break

        new_ok = bool(new_text) and len(new_text) >= 5
        applyable = (target is not None and new_ok
                     and (scope == "regulation" or article_index >= 0))

        note = ""
        if target is None:
            note = "시스템에 등록된 규정과 매칭되지 않음"
        elif not new_ok:
            note = "AI가 개정안 본문을 추출하지 못함"
        elif scope == "article" and article_index < 0 and article_head != "부 칙":
            note = "현행에 같은 조가 없음 — 신설일 수 있습니다"

        enriched.append({
            "rule_title":       (target.get("title") if target else rule_title_ai) or "(규정 미특정)",
            "rule_seq":         seq if target else -1,
            "scope":            scope,
            "article_head":     article_head,
            "article_index":    article_index,
            "current_text":     current_text,
            "new_text":         new_text,
            "reason":           reason,
            "match_confidence": confidence,
            "format_detected":  fmt,
            "note":             note,
            "matched":          applyable,
        })

        stage1_debug.append({
            "rule_seq":        seq,
            "rule_title":      rule_title_ai,
            "matched_title":   (target.get("title") if target else None),
            "article_head":    article_head,
            "format_detected": fmt,
            "new_preview":     new_text[:200],
            "new_length":      len(new_text),
        })

    # ── 조 단위 항목을 안건(rule_seq) 단위로 통합 ──
    # 사용자가 카드 여러 개 보면 정신없으니 한 안건 = 한 카드로 합친다.
    # new_text = 시스템 현행에서 변경된 조만 교체 + 신설 조(부 칙 등) 끝에 추가
    from collections import OrderedDict
    groups = OrderedDict()
    unmatched_items = []
    for it in enriched:
        seq = int(it.get("rule_seq", -1))
        if seq < 0:
            unmatched_items.append(it)
            continue
        groups.setdefault(seq, []).append(it)

    consolidated = []
    for seq, group in groups.items():
        target = next((r for r in history if int(r.get("seq", -1)) == seq), None)
        if target is None:
            consolidated.extend(group)
            continue
        latest = _get_latest_version(target)
        sys_content = (latest or {}).get("content", "") or ""
        sys_articles = _split_articles(sys_content)

        used_idx = set()
        changed_heads = []   # 표시용: 변경된 조 머리글들
        new_parts = []
        # 시스템의 각 조를 순회: 그룹에 같은 조 있으면 그것으로 교체, 없으면 그대로
        for sa in sys_articles:
            sa_head = (sa["head"] + sa.get("dup_label", "")).strip()
            sa_norm = _re.sub(r'\s+', '', sa_head)
            replaced = None
            for gi, gitem in enumerate(group):
                if gi in used_idx:
                    continue
                gh = _re.sub(r'\s+', '', (gitem.get("article_head") or ""))
                if gh == sa_norm:
                    replaced = gitem
                    used_idx.add(gi)
                    break
            # 조 번호만 같은지 fallback
            if replaced is None:
                m_s = _re.search(r'제\s*(\d+)\s*조(?:\s*의\s*(\d+))?', sa_head)
                if m_s:
                    wa, wb = m_s.group(1), m_s.group(2) or ""
                    for gi, gitem in enumerate(group):
                        if gi in used_idx:
                            continue
                        m_g = _re.search(r'제\s*(\d+)\s*조(?:\s*의\s*(\d+))?', (gitem.get("article_head") or ""))
                        if m_g and m_g.group(1) == wa and (m_g.group(2) or "") == wb:
                            replaced = gitem
                            used_idx.add(gi)
                            break
            if replaced is not None and (replaced.get("new_text") or "").strip():
                new_parts.append(replaced["new_text"].rstrip())
                changed_heads.append(sa_head)
            else:
                new_parts.append(sys_content[sa["start"]:sa["end"]].rstrip())

        # 그룹에 남은 항목 = 신설 조/부 칙 → 끝에 추가
        for gi, gitem in enumerate(group):
            if gi in used_idx:
                continue
            body = (gitem.get("new_text") or "").strip()
            if body:
                new_parts.append(body)
                head = (gitem.get("article_head") or "(신설)").strip()
                changed_heads.append(head + " (신설)")

        new_full_text = "\n\n".join(p for p in new_parts if p)

        # 통합 카드 생성
        first = group[0]
        consolidated.append({
            "rule_title":       target.get("title", first.get("rule_title", "")),
            "rule_seq":         seq,
            "scope":            "regulation",
            "article_head":     "",
            "article_index":    -1,
            "current_text":     sys_content,
            "new_text":         new_full_text,
            "reason":           first.get("reason", ""),
            "match_confidence": first.get("match_confidence", "medium"),
            "format_detected":  first.get("format_detected", ""),
            "note":             ("변경/신설: " + ", ".join(changed_heads)) if changed_heads else "변경 사항 없음",
            "matched":          bool(new_full_text.strip()) and len(new_full_text) >= 10,
            "changed_articles": changed_heads,
        })

    # 매칭 안 된 항목은 그대로 (목록 끝에 별도 표시용)
    consolidated.extend(unmatched_items)

    return {
        "items":           consolidated,
        "raw_text":        raw_text,
        "raw_truncated":   len(raw_text) > 16000,
        "table_count":     len(parsed_tables),
        "format_detected": ",".join(sorted(detected_formats)) if detected_formats else "",
        "stage1":          stage1_debug,
    }


# ══════════════════════════════════════════════════════════════════
# 단어 검색 / 일괄 치환 (Bulk word replace)
# ══════════════════════════════════════════════════════════════════

def _find_word_hits(content: str, word: str, ctx: int = 25) -> list:
    """content 안에서 word가 나오는 위치를 모두 찾아 앞뒤 문맥과 함께 반환"""
    hits = []
    if not word:
        return hits
    start = 0
    while True:
        idx = content.find(word, start)
        if idx == -1:
            break
        a = max(0, idx - ctx)
        b = min(len(content), idx + len(word) + ctx)
        hits.append({
            "pos":    idx,
            "before": content[a:idx],
            "match":  content[idx:idx+len(word)],
            "after":  content[idx+len(word):b],
        })
        start = idx + len(word)
    return hits


# -- 단어가 들어간 규정 검색 (챗봇/관리자 공용) ---------------------
class WordQ(BaseModel):
    word: str


@app.post("/word-search")
def word_search(req: WordQ):
    """본문에 특정 단어가 들어간 모든 규정을 찾는다. (인증 불필요 — 검색 전용)"""
    word = (req.word or "").strip()
    if not word:
        raise HTTPException(400, "검색할 단어를 입력하세요.")

    rules = _load_rules_json()
    results = []
    total_hits = 0
    for r in rules:
        content = r.get("content", "") or ""
        cnt = content.count(word)
        if cnt > 0:
            total_hits += cnt
            sample = _find_word_hits(content, word)[:3]
            results.append({
                "seq":     r.get("seq"),
                "title":   r.get("title", ""),
                "department": r.get("department", ""),
                "count":   cnt,
                "samples": sample,
            })
    results.sort(key=lambda x: -x["count"])
    return {"word": word, "rule_count": len(results),
            "total_hits": total_hits, "results": results}


# -- 일괄 치환 미리보기 (관리자) -----------------------------------
class ReplacePreviewQ(BaseModel):
    old_word: str
    new_word: str


@app.post("/bulk-replace/preview")
def bulk_replace_preview(req: ReplacePreviewQ, payload: dict = Depends(verify_token)):
    """old_word → new_word 치환 시 어떤 규정의 어디가 바뀌는지 미리 보여준다."""
    old = (req.old_word or "").strip()
    new = req.new_word if req.new_word is not None else ""
    if not old:
        raise HTTPException(400, "바꿀 단어(old_word)를 입력하세요.")
    if old == new:
        raise HTTPException(400, "바꿀 단어와 새 단어가 같습니다.")

    rules = _load_rules_json()
    results = []
    total_hits = 0
    for r in rules:
        content = r.get("content", "") or ""
        hits = _find_word_hits(content, old)
        if hits:
            total_hits += len(hits)
            results.append({
                "seq":   r.get("seq"),
                "title": r.get("title", ""),
                "department": r.get("department", ""),
                "count": len(hits),
                "hits":  hits[:20],   # 미리보기는 규정당 최대 20곳
            })
    results.sort(key=lambda x: -x["count"])
    return {
        "old_word": old, "new_word": new,
        "rule_count": len(results), "total_hits": total_hits,
        "results": results,
    }


# -- 일괄 치환 실행 (관리자) ---------------------------------------
class BulkReplaceQ(BaseModel):
    old_word: str
    new_word: str
    seqs: list[int]          # 실제로 치환할 규정 seq 목록 (미리보기에서 체크한 것)


@app.post("/bulk-replace/apply")
def bulk_replace_apply(req: BulkReplaceQ, payload: dict = Depends(verify_token)):
    """
    선택된 규정들에서 old_word → new_word 일괄 치환.
    - 전체 작업을 하나의 batch 백업으로 저장 (한 번에 되돌리기 가능)
    - 각 규정 JSON content 갱신 + DB 본문 청크 재색인
    """
    old = (req.old_word or "").strip()
    new = req.new_word if req.new_word is not None else ""
    if not old:
        raise HTTPException(400, "바꿀 단어를 입력하세요.")
    if old == new:
        raise HTTPException(400, "바꿀 단어와 새 단어가 같습니다.")
    if not req.seqs:
        raise HTTPException(400, "치환할 규정을 1개 이상 선택하세요.")

    rules = _load_rules_json()
    seq_set = set(req.seqs)
    today = datetime.now().strftime("%Y-%m-%d")

    # 백업 준비
    os.makedirs(REVISION_BACKUP_DIR, exist_ok=True)
    batch_id = f"bulk_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    backup_items = []   # 규정별 개정 전 스냅샷

    changed = []
    for idx, r in enumerate(rules):
        seq = r.get("seq")
        if seq not in seq_set:
            continue
        content = r.get("content", "") or ""
        cnt = content.count(old)
        if cnt == 0:
            continue

        title = r.get("title", "")
        dept  = r.get("department", "")
        url   = r.get("url", "")

        # 개정 전 DB 청크 백업
        try:
            conn = psycopg2.connect(DB_URL); cur = conn.cursor()
            cur.execute("""
                SELECT id,rule_title,seq,article,department,url,content
                FROM rule_chunks
                WHERE seq = %s AND article != '개정이력'
                  AND (url IS NULL OR url NOT LIKE 'upload://%%')
            """, (str(seq),))
            db_rows = [
                {"id": c[0], "rule_title": c[1], "seq": c[2], "article": c[3],
                 "department": c[4], "url": c[5], "content": c[6]}
                for c in cur.fetchall()
            ]
            conn.close()
        except Exception as e:
            raise HTTPException(500, f"백업용 DB 조회 실패(seq {seq}): {e}")

        backup_items.append({
            "seq": seq, "title": title,
            "json_item": dict(r),       # 개정 전 JSON 원본
            "db_chunks": db_rows,
        })

        # 치환
        new_content = content.replace(old, new)
        revisions = r.get("revisions", [])
        revisions.append({
            "kind":       "bulk_replace",
            "old_word":   old, "new_word": new,
            "replaced":   cnt,
            "revised_at": today,
            "backup_id":  batch_id,
        })
        r["content"]   = new_content
        r["revisions"] = revisions
        rules[idx] = r

        # DB 재색인
        new_chunks = _reindex_rule_chunks(seq, title, dept, url, new_content)
        changed.append({"seq": seq, "title": title,
                         "replaced": cnt, "new_chunks": new_chunks})

    if not changed:
        raise HTTPException(400, "선택한 규정에서 해당 단어를 찾지 못했습니다.")

    # JSON 저장
    _save_rules_json(rules)

    # batch 백업 파일 저장 (한 번에 되돌리기용)
    backup = {
        "backup_id":    batch_id,
        "kind":         "bulk_replace",
        "title":        f"일괄 치환: '{old}' → '{new}' ({len(changed)}개 규정)",
        "old_word":     old, "new_word": new,
        "backed_up_at": datetime.now().isoformat(timespec="seconds"),
        "items":        backup_items,
    }
    with open(os.path.join(REVISION_BACKUP_DIR, f"{batch_id}.json"),
              "w", encoding="utf-8") as f:
        _json.dump(backup, f, ensure_ascii=False, indent=2)

    return {
        "success": True, "batch_id": batch_id,
        "old_word": old, "new_word": new,
        "rule_count": len(changed),
        "total_replaced": sum(c["replaced"] for c in changed),
        "changed": changed,
    }




# ══════════════════════════════════════════════════════════════════
# AI 형식/표현 검사 — 규정 본문이 한성대 표준 형식에 맞는지 검토
# ══════════════════════════════════════════════════════════════════

# 한성대 규정 표준 형식 (크롤링한 원문 1,600여 버전에서 파악한 규칙)
_HSU_RULE_FORMAT = """[한성대학교 규정 표준 형식 — 반드시 이 규칙을 따라야 함]

■ 조 (條)
  · 형식: '제 N 조 (제목)' — '제' '숫자' '조' 사이를 반드시 띄움, 괄호 안에 조 제목
  · 예) '제 1 조 (목적)', '제 12 조 (재택근무)', '제 30 조의2 (관련 규정)'
  · 잘못된 예) '제1조(목적)' '제 1조 (목적)' 'Article 1' — 모두 형식 위반

■ 항 (項)
  · 형식: 원문자 ① ② ③ ④ ⑤ ... 만 사용 (반드시 동그라미 숫자)
  · 잘못된 예) '1)' '(1)' '제1항' — 모두 형식 위반

■ 호 (號)
  · 형식: 아라비아 숫자 + 마침표 '1.' '2.' '3.' ...
  · 잘못된 예) '가.' '나.' '①' '(1)' '1)' — 호 자리에 쓰면 모두 형식 위반

■ 목 (目)
  · 형식: 한글 가나다 + 마침표 '가.' '나.' '다.' ...
  · 호의 하위 항목일 때만 사용. 호 위치에 쓰지 말 것.

■ 개정일 표기
  · 형식: '(YYYY. M. D. 개정)' 또는 본문 끝에 '(YYYY. M. D.)'
  · 예) '(2024. 3. 15. 개정)', '(2025. 9. 23.)'
  · 잘못된 예) '(2024-03-15)' '(2024년 3월 15일)' '2024.3.15' — 형식 위반

■ 부칙
  · 본문과 분리해 '부 칙' (한 글자씩 띄어쓰기) 표기
  · 아래로 '(시행일) 이 규정은 YYYY년 M월 D일부터 시행한다.' 형식

■ 문장 표현 (규정체)
  · 평서문, '~한다' '~하여야 한다' '~할 수 있다' 형태
  · 금지: '~해요' '~합니다' '~인 것 같다' '~로 사료된다' 같은 구어/추측 표현
  · 능동·간결 — 한 문장 한 의도. 모호한 표현(예: '적절히' '필요시' '가급적')은 가급적 회피.

■ 띄어쓰기 (한성대 표기 관례)
  · '제 N 조', '제 N 항' — 띄움
  · 학칙/규정명 인용 시 「 」 (낫표) 사용. 예) 「한성대학교 학칙」 제 24 조

■ 기타
  · 영문·한자 병기는 한글(英文/漢字) 형태로 괄호 안에
  · 표·서식 인용: '별표 1', '별지 제1호 서식' 형태
  · 신구조문대비표는 '[현 행]' / '[개 정 안]' 표기
"""


# 형식 검사용 샘플 캐시 — 서버 시작 시 1회 로드 (잘 정돈된 한성대 규정 본문 일부)
_FORMAT_SAMPLE_CACHE = {"text": "", "loaded": False}

def _load_format_samples() -> str:
    """history.json에서 형식이 깔끔한 규정 본문 일부를 샘플로 추출.
    프롬프트에 직접 첨부할 용도. 토큰 절약 위해 4,500자로 제한."""
    if _FORMAT_SAMPLE_CACHE["loaded"]:
        return _FORMAT_SAMPLE_CACHE["text"]

    # 우선순위: 학칙 → 학사운영규정 → 등록금규정 → 그 외 첫 몇 편
    priority_keywords = ["학칙", "학사운영", "등록금", "장학"]
    samples = []
    total_len = 0
    LIMIT = 4500

    try:
        hist = _load_history_json()
        # 우선순위 규정부터 모음
        ordered = []
        for kw in priority_keywords:
            for r in hist:
                title = r.get("title", "")
                if kw in title and r not in ordered:
                    ordered.append(r)
        # 나머지로 채움 (최대 8편까지 검토)
        for r in hist:
            if r not in ordered:
                ordered.append(r)
            if len(ordered) >= 8:
                break

        for r in ordered:
            latest = _get_latest_version(r)
            if not latest:
                continue
            content = (latest.get("content") or "").strip()
            if not content or len(content) < 200:
                continue
            # 조 시작부터 일부만 (보통 제1조~제3조까지)
            # '제 N 조' 패턴 위치 찾고 처음 3개 조까지
            m = list(_re.finditer(r'제\s*\d+\s*조\s*\(', content))
            if len(m) >= 2:
                end_pos = m[2].start() if len(m) >= 3 else len(content)
                excerpt = content[m[0].start():end_pos].strip()
            else:
                excerpt = content[:1500]

            if len(excerpt) < 100:
                continue
            # 너무 길면 자름
            excerpt = excerpt[:1500]

            block = f"━━ 「{r.get('title', '')}」 본문 발췌 ━━\n{excerpt}"
            if total_len + len(block) > LIMIT:
                break
            samples.append(block)
            total_len += len(block)
            if len(samples) >= 3:
                break
    except Exception:
        pass

    text = "\n\n".join(samples) if samples else ""
    _FORMAT_SAMPLE_CACHE["text"] = text
    _FORMAT_SAMPLE_CACHE["loaded"] = True
    return text


class FormatCheckQ(BaseModel):
    text: str
    scope: str = "article"   # article | full | bulk — 안내 문구용


@app.post("/format-check")
def format_check(req: FormatCheckQ, payload: dict = Depends(verify_token)):
    """
    규정 본문을 AI가 검토해 표준 형식/표현에 어긋난 부분을 찾는다.
    반환: 각 지적사항 {snippet, kind(format|expression), advice}
      - snippet: 본문에서 문제가 되는 '정확한 원문 일부' (프론트가 찾아서 하이라이트)
      - kind: format = 형식 오류 / expression = 표현 제안
      - advice: 어떻게 고치면 좋은지 한 줄 조언
    """
    text = (req.text or "").strip()
    if len(text) < 5:
        raise HTTPException(400, "검사할 내용이 너무 짧습니다.")

    doc = text[:5000]   # 토큰 보호
    sample_block = _load_format_samples()
    sample_section = f"""\n[표준 샘플 — 한성대학교 실제 규정 본문]
아래는 한성대학교의 실제 규정에서 가져온 표본이다.
검사 대상 본문이 이 형식과 어긋나는 부분을 모두 찾아라.

{sample_block}\n""" if sample_block else ""

    prompt = f"""너는 한성대학교 규정 편집을 돕는 엄격한 형식 검토 도우미다.
아래 [표준 형식 규칙]과 [표준 샘플]을 기준으로, [검사 대상 본문]에서 형식·표현 위반을 모두 찾아라.

{_HSU_RULE_FORMAT}
{sample_section}
[검사 대상 본문]
{doc}

[지시 — 반드시 지킬 것]
1. 위 [표준 샘플]의 표기 방식과 다른 부분은 모두 형식 위반으로 본다.
   예: 샘플은 '제 1 조 (목적)'인데 검사 본문이 '제1조(목적)'이면 → 형식 오류.
   예: 샘플은 항이 '①'인데 검사 본문이 '1)'이면 → 형식 오류.
2. 각 지적은 '본문에 실제로 있는 짧은 원문 조각'을 그대로 인용해야 한다.
   - 인용은 5~40자 사이. 본문에 존재하지 않는 인용은 절대 만들지 마라.
   - 인용 부분에 ... 같은 생략 표시 금지. 정확한 문자열 그대로.
3. 두루뭉술한 advice 금지:
   - 나쁜 예: "형식이 어색합니다", "수정이 필요합니다", "표현이 부적절합니다"
   - 좋은 예: "'제1조(목적)' → '제 1 조 (목적)'처럼 '제·숫자·조' 사이를 띄우고 괄호 앞도 띄어 씁니다."
   - advice는 어떻게 고치는지 **구체적인 수정안**을 포함해야 한다.
4. 표현(expression)은 명백한 구어체·추측·중복일 때만 지적. 사소한 문체 차이는 넘어가라.
5. 형식 오류(format) 우선. 항·호·목·개정일 표기·띄어쓰기·괄호 사용을 가장 깐깐히 본다.
6. 위반이 없으면 빈 배열 []. 억지로 만들지 마라.
7. 출력은 JSON 배열만. 설명·머리말·코드 블록 없이.

형식:
[
  {{"snippet": "제1조(목적)", "kind": "format", "advice": "'제 1 조 (목적)'처럼 '제·숫자·조' 사이를 띄우고 괄호 앞도 띄어 씁니다."}},
  {{"snippet": "가. 학생의 신분", "kind": "format", "advice": "호는 '1.' '2.' 순으로 씁니다. '가. 나. 다.'는 호 아래 목에서만 사용합니다."}},
  {{"snippet": "필요한 것 같다", "kind": "expression", "advice": "규정체로 '~한다' 또는 '~하여야 한다'를 씁니다. 추측 표현 '~것 같다'는 부적절합니다."}}
]

- kind: "format" 또는 "expression" 둘 중 하나
- 지적은 최대 20개"""

    try:
        raw = claude_chat(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2500,
            temperature=0,
        )
        raw = (raw or "").strip()
        raw = _re.sub(r'^```[a-zA-Z]*\s*', '', raw)
        raw = _re.sub(r'\s*```$', '', raw).strip()
        # JSON 배열만 추출
        m = _re.search(r'\[.*\]', raw, _re.S)
        if m:
            raw = m.group(0)
        issues = _json.loads(raw)
        if not isinstance(issues, list):
            issues = []
    except Exception as e:
        return {"issues": [], "error": f"AI 검토 실패: {e}", "text": text}

    # ── 후처리 ─────────────────────────────────────────────────
    # 1) snippet 이 실제 본문에 있어야 함
    # 2) advice 가 너무 짧거나 두루뭉술하면 버림
    # 3) 중복 snippet 제거
    _VAGUE_PATTERNS = [
        "필요합니다", "필요해", "어색합니다", "어색해",
        "수정이 필요", "검토가 필요", "확인이 필요",
        "부적절합니다.", "적절하지 않", "맞지 않",
    ]
    cleaned = []
    for it in issues:
        if not isinstance(it, dict):
            continue
        snip = (it.get("snippet") or "").strip()
        kind = it.get("kind", "format")
        advice = (it.get("advice") or "").strip()
        if not snip or not advice:
            continue
        if snip not in text:
            continue   # 본문에 없는 인용은 버림 (AI 환각 방지)
        if kind not in ("format", "expression"):
            kind = "format"
        # advice가 너무 짧으면 버림 (구체적 수정안 강제)
        if len(advice) < 15:
            continue
        # 두루뭉술 advice — 구체적 수정안(→, ' 같은 인용 부호, '~로') 없으면 버림
        has_concrete = ("→" in advice or "'" in advice or "\"" in advice
                        or "처럼" in advice or "씁니다" in advice
                        or "표기" in advice or "사용" in advice)
        is_vague_only = any(p in advice for p in _VAGUE_PATTERNS) and not has_concrete
        if is_vague_only:
            continue
        cleaned.append({"snippet": snip, "kind": kind, "advice": advice})

    # 중복 snippet 제거
    seen = set()
    final = []
    for it in cleaned:
        if it["snippet"] in seen:
            continue
        seen.add(it["snippet"])
        final.append(it)

    return {"issues": final, "text": text, "count": len(final)}


if __name__ == "__main__":
    import uvicorn
    # 필수 HTML 파일 존재 여부 확인 (없으면 페이지 못 띄움)
    print("=" * 60)
    print(f"[Startup] BASE_DIR: {BASE_DIR}")
    for fn in ["index.html", "login.html", "upload.html"]:
        p = os.path.join(BASE_DIR, fn)
        mark = "✓" if os.path.exists(p) else "✗ 없음!"
        print(f"[Startup] {mark}  {fn}")
    print("=" * 60)
    uvicorn.run(app, host="0.0.0.0", port=8000)